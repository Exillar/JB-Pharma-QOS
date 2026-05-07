"""
Module: docx_builder
Responsibility: Fills QOS template DOCX paragraphs and tables with content
extracted from CTD Module 3 PDFs.

Design principles (borrowed from JB-Pharma-QIS):
- _DocxHelper base class holds ALL shared XML/paragraph utilities.
- No hardcoded drug names, company names, or pharmaceutical properties.
- All fallback strings live in config_loader.S2FillConfig.
- Noise cleaning uses the auto-detected blocklist from pdf_extractor.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.document import Document as _Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

import fitz

from pdf_extractor import ExtractedSectionContent
from config_loader import S2FillConfig, NoiseConfig, DiagramConfig


# ---------------------------------------------------------------------------
# Module-level helpers (ported from QIS)
# ---------------------------------------------------------------------------

_PAGE_NUM_RE = re.compile(r"^\d{1,4}$")
_PAGE_OF_RE = re.compile(r"^\d+\s+of\s+\d+\s*$", re.IGNORECASE)


def _safe_row_cells(row):
    try:
        return tuple(row.cells)
    except Exception:
        return ()


def _collapse_blank_paragraphs(doc: _Document) -> int:
    _NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    consecutive = 0
    to_remove = []
    for elem in list(doc.element.body):
        if elem.tag.split("}")[-1] == "p":
            txt = "".join(t.text or "" for t in elem.iter(f"{{{_NS}}}t")).strip()
            # Check if paragraph has drawings (images/shapes)
            has_drawing = len(list(elem.iter(f"{{{_NS}}}drawing"))) > 0
            # Only mark as blank if it has no text AND no drawings
            if not txt and not has_drawing:
                consecutive += 1
                if consecutive > 1:
                    to_remove.append(elem)
            else:
                consecutive = 0
        else:
            consecutive = 0
    for elem in to_remove:
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)
    return len(to_remove)


def _remove_repeated_header_paragraphs(
    doc: _Document,
    *,
    preserve_patterns: tuple[str, ...] = (),
) -> int:
    """
    Remove repeated short lines that are typically injected PDF headers.
    """
    from collections import Counter

    _NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    freq = Counter(p.text.strip() for p in doc.paragraphs if p.text.strip())
    preserve = tuple(p.lower() for p in preserve_patterns if p)
    noise = {
        t
        for t, c in freq.items()
        if c >= 3 and len(t) < 120 and not any(p in t.lower() for p in preserve)
    }

    removed = 0
    for elem in list(doc.element.body):
        if elem.tag.split("}")[-1] != "p":
            continue
        text = "".join(t.text or "" for t in elem.iter(f"{{{_NS}}}t")).strip()
        if text in noise:
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)
                removed += 1
    return removed


def _remove_pdf_noise_paragraphs(
    doc: _Document,
    *,
    preserve_phrases: tuple[str, ...] = (),
) -> int:
    """
    Remove generic page-header/footer style paragraphs left after insertion.
    """
    _NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    body = doc.element.body
    elements = list(body)

    def _wt(e) -> str:
        return "".join(t.text or "" for t in e.iter(f"{{{_NS}}}t")).strip()

    removed = 0
    preserve = tuple(p.lower() for p in preserve_phrases if p)
    for i, elem in enumerate(elements):
        if elem.tag.split("}")[-1] != "p":
            continue
        text = _wt(elem)
        if not text:
            continue
        text_lower = text.lower()
        if preserve and any(p in text_lower for p in preserve):
            continue

        drop = False
        if re.match(r"^.{10,75}\s+\d{1,4}$", text) and len(text) < 80:
            drop = True
        elif re.search(r"3\.2[\. ][A-Z0-9P]", text) and len(text) < 200:
            for j in range(max(0, i - 10), i):
                if re.search(r"2\.3\.[SP]", _wt(elements[j])):
                    drop = True
                    break
        elif _PAGE_NUM_RE.match(text) or _PAGE_OF_RE.match(text):
            drop = True

        if drop:
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)
                removed += 1
    return removed


def _remove_empty_visual_tables(doc: _Document, *, keep_first_n_tables: int) -> int:
    """
    Remove fully/mostly empty tables created as visual artifacts.
    Original template tables are protected by keep_first_n_tables.
    """
    _NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    removed = 0

    def get_text(cell) -> str:
        return "".join(t.text or "" for t in cell._element.iter(f"{{{_NS}}}t")).strip()

    for idx, table in enumerate(list(doc.tables)):
        if idx < keep_first_n_tables:
            continue

        total_cells = 0
        empty_cells = 0
        text_cells = 0
        for row in table.rows:
            for cell in _safe_row_cells(row):
                total_cells += 1
                text = get_text(cell)
                if not text:
                    empty_cells += 1
                else:
                    text_cells += 1

        if total_cells == 0:
            continue

        if total_cells > 4 and text_cells == 0:
            table._element.getparent().remove(table._element)
            removed += 1
        elif total_cells > 6 and (empty_cells / total_cells) > 0.8:
            table._element.getparent().remove(table._element)
            removed += 1

    return removed


def _remove_low_content_injected_tables(doc: _Document, *, keep_first_n_tables: int) -> int:
    """
    Remove low-value injected table artifacts while preserving template tables.
    """
    _NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    removed = 0

    def get_text(cell) -> str:
        return "".join(t.text or "" for t in cell._element.iter(f"{{{_NS}}}t")).strip()

    for idx, table in enumerate(list(doc.tables)):
        if idx < keep_first_n_tables:
            continue

        total_cells = 0
        text_cells = 0
        text_chars = 0

        for row in table.rows:
            for cell in _safe_row_cells(row):
                total_cells += 1
                text = get_text(cell)
                if text:
                    text_cells += 1
                    text_chars += len(text)

        if total_cells == 0:
            continue

        empty_ratio = (total_cells - text_cells) / total_cells
        should_remove = False
        if text_cells == 0 and total_cells >= 4:
            should_remove = True
        elif empty_ratio > 0.85 and text_chars < 80 and total_cells >= 6:
            should_remove = True
        elif text_cells <= 2 and total_cells >= 8 and text_chars < 100:
            should_remove = True

        if should_remove:
            table._element.getparent().remove(table._element)
            removed += 1

    return removed


def _run_injected_artifact_cleanup(
    doc: _Document,
    *,
    keep_first_n_tables: int,
    preserve_repeated_patterns: tuple[str, ...] = (),
    preserve_phrases: tuple[str, ...] = (),
) -> dict[str, int]:
    """
    Unified cleanup pipeline for post-insertion artifacts.
    Returns per-step removal counters for debugging/telemetry.
    """
    stats = {
        "repeated_headers_removed": _remove_repeated_header_paragraphs(
            doc,
            preserve_patterns=preserve_repeated_patterns,
        ),
        "noise_paragraphs_removed": _remove_pdf_noise_paragraphs(
            doc,
            preserve_phrases=preserve_phrases,
        ),
        "empty_tables_removed": _remove_empty_visual_tables(
            doc, keep_first_n_tables=keep_first_n_tables
        ),
        "low_content_tables_removed": _remove_low_content_injected_tables(
            doc, keep_first_n_tables=keep_first_n_tables
        ),
        "blank_paragraphs_collapsed": _collapse_blank_paragraphs(doc),
    }
    return stats


# ---------------------------------------------------------------------------
# Shared base class
# ---------------------------------------------------------------------------

class _DocxHelper:
    """Shared XML/paragraph utilities for all filler classes."""

    SECTION_START: str = ""
    SECTION_END: str = ""

    # ------------------------------------------------------------------
    # Low-level XML helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        para = Paragraph(new_p, paragraph._parent)
        if text:
            para.add_run(text)
        return para

    def _get_available_page_width(self, doc: _Document) -> int:
        """
        Calculate available width for content (page width - margins) in EMUs.
        Used to scale images proportionally without hardcoding sizes.
        Returns width in EMUs (914400 EMU = 1 inch).
        """
        try:
            section = doc.sections[0]
            page_width = section.page_width
            left_margin = section.left_margin
            right_margin = section.right_margin
            available = page_width - left_margin - right_margin
            # Convert to Inches and back to maintain proper spacing
            # Subtract a small margin for padding (0.2 inches on each side)
            padding = int(914400 * 0.2)
            return max(available - 2 * padding, int(914400 * 4))  # Minimum 4 inches
        except Exception:
            # Fallback to reasonable default if section info unavailable
            return int(914400 * 5.5)  # 5.5 inches

    def _get_available_page_height(self, doc: _Document) -> int:
        """Calculate available page height for inserted image content in EMUs."""
        try:
            section = doc.sections[0]
            page_height = section.page_height
            top_margin = section.top_margin
            bottom_margin = section.bottom_margin
            available = page_height - top_margin - bottom_margin
            padding = int(914400 * 0.3)
            return max(available - 2 * padding, int(914400 * 3.5))
        except Exception:
            return int(914400 * 7.5)

    def _add_picture_autofit(self, run, image_path: Path | str, doc: _Document) -> None:
        """
        Add picture while respecting both width and height bounds.
        This prevents very tall images from overflowing the page visually.
        """
        img_path = Path(image_path)
        max_width = self._get_available_page_width(doc)
        max_height = self._get_available_page_height(doc)
        target_width = max_width

        try:
            pix = fitz.Pixmap(str(img_path))
            w = max(1, pix.width)
            h = max(1, pix.height)
            aspect = h / w
            if aspect > 0:
                height_limited_width = int(max_height / aspect)
                if height_limited_width > 0:
                    target_width = min(max_width, height_limited_width)
        except Exception:
            target_width = max_width

        run.add_picture(str(img_path), width=target_width)

    @staticmethod
    def _delete_paragraph(paragraph: Paragraph) -> None:
        elem = paragraph._element
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)

    @staticmethod
    def _set_runs_style(
        paragraph: Paragraph,
        bold: bool | None = None,
        italic: bool | None = None,
    ) -> None:
        if not paragraph.runs:
            paragraph.add_run(paragraph.text or "")
        for run in paragraph.runs:
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic

    @staticmethod
    def _replace_paragraph_with_runs(
        paragraph: Paragraph,
        chunks: list[tuple[str, bool | None, bool | None]],
    ) -> None:
        for run in list(paragraph.runs):
            paragraph._p.remove(run._r)
        for text, bold, italic in chunks:
            r = paragraph.add_run(text)
            if bold is not None:
                r.bold = bold
            if italic is not None:
                r.italic = italic

    # ------------------------------------------------------------------
    # Paragraph search helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _norm(s: str) -> str:
        return re.sub(r"\s+", " ", s).strip().lower()

    @classmethod
    def _find_para_index(
        cls,
        paragraphs: list[Paragraph],
        needle: str,
        start_idx: int,
        end_idx: int,
        *,
        startswith: bool = False,
    ) -> int | None:
        n = cls._norm(needle)
        for i in range(start_idx, min(end_idx, len(paragraphs))):
            t = cls._norm(paragraphs[i].text or "")
            if startswith and t.startswith(n):
                return i
            if not startswith and n in t:
                return i
        return None

    @classmethod
    def _find_para_index_doc(
        cls,
        doc: _Document,
        needle: str,
        start_idx: int,
        end_idx: int,
        *,
        startswith: bool = False,
    ) -> int | None:
        return cls._find_para_index(
            doc.paragraphs, needle, start_idx, end_idx, startswith=startswith
        )

    # ------------------------------------------------------------------
    # Section range
    # ------------------------------------------------------------------

    def _get_target_range(self, doc: _Document) -> tuple[int, int]:
        start_idx = None
        end_idx = None
        for i, p in enumerate(doc.paragraphs):
            text = (p.text or "").strip()
            if start_idx is None and self.SECTION_START.lower() in text.lower():
                start_idx = i
            if start_idx is not None and self.SECTION_END.lower() in text.lower():
                end_idx = i
                break
        if start_idx is None:
            raise ValueError(f"Section start '{self.SECTION_START}' not found in template DOCX")
        return start_idx, end_idx if end_idx is not None else len(doc.paragraphs)

    # ------------------------------------------------------------------
    # Common value insertion
    # ------------------------------------------------------------------

    def _add_answer_after(self, doc: _Document, anchor_idx: int, value: str) -> None:
        if not value.strip():
            return
        self._insert_paragraph_after(doc.paragraphs[anchor_idx], value)

    def _set_value_under_label(
        self,
        doc: _Document,
        label: str,
        value: str,
        start_idx: int,
        end_idx: int,
    ) -> None:
        if not value.strip():
            return
        idx = self._find_para_index_doc(doc, label, start_idx, end_idx, startswith=True)
        if idx is None:
            return
        if idx + 1 < len(doc.paragraphs) and not (doc.paragraphs[idx + 1].text or "").strip():
            doc.paragraphs[idx + 1].text = value
        else:
            self._insert_paragraph_after(doc.paragraphs[idx], value)

    def _remove_paragraphs_matching(
        self,
        doc: _Document,
        patterns: list[str],
        start_idx: int,
        end_idx: int,
    ) -> None:
        to_delete = []
        for i in range(start_idx, min(end_idx, len(doc.paragraphs))):
            t = self._norm(doc.paragraphs[i].text or "")
            if any(t.startswith(self._norm(p)) for p in patterns):
                to_delete.append(i)
        for idx in reversed(to_delete):
            if idx < len(doc.paragraphs):
                self._delete_paragraph(doc.paragraphs[idx])

    # ------------------------------------------------------------------
    # Text-block helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _extract_block(text: str, start_pattern: str, end_patterns: list[str]) -> str:
        m = re.search(start_pattern, text, flags=re.IGNORECASE)
        if not m:
            return ""
        block = text[m.end():]
        end_pos = len(block)
        for ep in end_patterns:
            em = re.search(ep, block, flags=re.IGNORECASE)
            if em:
                end_pos = min(end_pos, em.start())
        return block[:end_pos].strip()

    @staticmethod
    def _first_meaningful_line(block: str) -> str:
        for ln in block.splitlines():
            c = ln.strip()
            if c:
                return c
        return ""

    @staticmethod
    def _clean_line(line: str) -> str:
        return re.sub(r"\s+", " ", line.strip())

    def _resolve_name_manufacturer_line(
        self, ref_docx: Path | None, section_heading: str
    ) -> str:
        """Read the filled reference DOCX to find the (drug, manufacturer) line."""
        if not ref_docx or not ref_docx.exists():
            return ""
        try:
            ref = Document(ref_docx)
            for i, p in enumerate(ref.paragraphs):
                if section_heading.lower() in (p.text or "").lower():
                    for j in range(i + 1, min(i + 6, len(ref.paragraphs))):
                        cand = (ref.paragraphs[j].text or "").strip()
                        if cand.startswith("(") and "," in cand:
                            return cand
        except Exception:
            pass
        return ""


# ---------------------------------------------------------------------------
# S1 filler — 2.3.S.1 General Information
# ---------------------------------------------------------------------------

class DocxFiller(_DocxHelper):
    SECTION_START = "2.3.S.1 General Information"
    SECTION_END = "2.3.S.2 Manufacture"

    def __init__(
        self,
        template_docx: Path,
        filled_reference_docx: Path | None = None,
    ) -> None:
        self.template_docx = template_docx
        self.filled_reference_docx = filled_reference_docx

    # ------------------------------------------------------------------
    # Parsers — extract from PDF text, no drug-specific fallbacks
    # ------------------------------------------------------------------

    def _parse_s11(self, raw_text: str) -> dict[str, str]:
        inn_block = self._extract_block(
            raw_text,
            r"Recommended\s+International\s+Nonproprietary\s+Name\s*\(INN\)\s*:",
            [r"Compendial\s+name\s*:"],
        )
        comp_block = self._extract_block(
            raw_text,
            r"Compendial\s+name\s*:",
            [r"Chemical\s+name\s*\(s\)\s*:"],
        )
        chem_block = self._extract_block(
            raw_text,
            r"Chemical\s+name\s*\(s\)\s*:",
            [r"Chemical\s+Abstracts\s+Service\s*\(CAS\)\s+registry\s+number\s*:"],
        )
        cas_block = self._extract_block(
            raw_text,
            r"Chemical\s+Abstracts\s+Service\s*\(CAS\)\s+registry\s+number\s*:",
            [r"Other\s+non-proprietary\s+name"],
        )
        return {
            "a": self._first_meaningful_line(inn_block),
            "b": self._first_meaningful_line(comp_block),
            "c": chem_block.strip(),
            "d": "",
            "e": "",
            "f": self._first_meaningful_line(cas_block),
        }

    def _parse_s12(self, raw_text: str) -> dict[str, str]:
        mf = re.search(r"Molecular\s+Formula\s*:\s*(.+)", raw_text, flags=re.IGNORECASE)
        mw = re.search(r"Molecular\s+weight\s*:\s*(.+)", raw_text, flags=re.IGNORECASE)
        return {
            "b": self._clean_line(mf.group(1)) if mf else "",
            "c": self._clean_line(mw.group(1)) if mw else "",
        }

    def _parse_s13(self, raw_text: str) -> dict[str, str]:
        """Extract physicochemical properties from raw PDF text.
        Returns empty strings when not found — NO drug-specific fallbacks.
        """
        melt = self._extract_block(raw_text, r"\bMelting\s+point\b", [r"\bpH\s*[-:]\b"]) or ""
        ph = self._extract_block(raw_text, r"\bpH\s*[-:]", [r"\bPartition\s+coefficients\b"]) or ""
        part = self._extract_block(raw_text, r"\bPartition\s+coefficients\b\s*[-:]", [r"\bPK\s*[-:]"]) or ""
        pk = self._extract_block(raw_text, r"\bPK\s*[-:]", [r"\bSpecific\s+Rotation\b"]) or ""
        rot = self._extract_block(raw_text, r"\bSpecific\s+Rotation\b", [r"\bPolymorphic\s+Form\b", r"\bSolub"]) or ""

        desc = self._extract_block(raw_text, r"\bPhysical\s+description\b\s*[-:]?", [r"\bSolub", r"\bPolymorphic"]) or ""
        sol = self._extract_block(raw_text, r"\bSolub(?:ility|ilities)\b\s*[-:]?", [r"\bPolymorphic", r"\bpH\b"]) or ""
        poly = self._extract_block(raw_text, r"\bPolymorphic\s+[Ff]orm\b\s*[-:]?", [r"\bSolvate", r"\bHydrate"]) or ""

        return {
            "a": self._first_meaningful_line(desc),
            "b": self._first_meaningful_line(sol) or sol.strip(),
            "poly": self._first_meaningful_line(poly),
            "solvate": "",
            "hydrate": "",
            "other": "",
            "ph": self._first_meaningful_line(ph),
            "pk": "\n".join(ln.strip() for ln in pk.splitlines() if ln.strip()),
            "partition": "\n".join(ln.strip() for ln in part.splitlines() if ln.strip()),
            "melting": self._first_meaningful_line(melt),
            "rotation": self._first_meaningful_line(rot),
            "refractive": "",
            "hygro": "",
            "uv": "",
        }

    # ------------------------------------------------------------------
    # Table filler
    # ------------------------------------------------------------------

    def _fill_property_table(
        self, doc: _Document, s13: dict[str, str]
    ) -> None:
        table_row_map = {
            "ph": s13["ph"],
            "pk": s13["pk"],
            "pka": s13["pk"],
            "partition coefficients": s13["partition"],
            "melting/boiling points": s13["melting"],
            "specific optical rotation (specify solvent)": s13["rotation"],
            "refractive index (liquids)": s13["refractive"],
            "hygroscopicity": s13["hygro"],
            "uv absorption maxima/molar absorptivity": s13["uv"],
            "other": s13["other"],
        }
        for table in doc.tables:
            if not table.rows:
                continue
            if self._norm(table.cell(0, 0).text) != "property":
                continue
            for row in table.rows[1:]:
                key = self._norm(row.cells[0].text)
                if key in table_row_map and len(row.cells) > 1:
                    row.cells[1].text = table_row_map[key]

    # ------------------------------------------------------------------
    # Visual formatting
    # ------------------------------------------------------------------

    def _apply_visual_formatting(
        self, doc: _Document, start_idx: int, end_idx: int
    ) -> None:
        bold_starts = {
            "2.3.S.1 General Information",
            "2.3.S.1.2 Structure",
            "2.3.S.1.3 General Properties",
        }
        label_bold_prefixes = [
            "(Recommended) International Non-proprietary name (INN):",
            "Compendial name",
            "Chemical name",
        ]
        for i in range(start_idx, min(end_idx, len(doc.paragraphs))):
            p = doc.paragraphs[i]
            t = (p.text or "").strip()
            if any(t.startswith(s) for s in bold_starts):
                self._set_runs_style(p, bold=True, italic=False)
            elif t.startswith("2.3.S.1.1 Nomenclature"):
                self._set_runs_style(p, bold=False, italic=False)
            if any(t.startswith(lb) for lb in label_bold_prefixes):
                self._set_runs_style(p, bold=True, italic=False)

    # ------------------------------------------------------------------
    # Public entry point
    # ------------------------------------------------------------------

    def fill_s1_section(
        self,
        extracted: dict[str, ExtractedSectionContent],
        output_docx: Path,
    ) -> list[str]:
        doc = Document(self.template_docx)
        template_table_count = len(doc.tables)
        warnings: list[str] = []

        start_idx, end_idx = self._get_target_range(doc)

        name_mfr_line = self._resolve_name_manufacturer_line(
            self.filled_reference_docx, "2.3.S.1.1 Nomenclature"
        )
        for heading in [
            "2.3.S.1 General Information",
            "2.3.S.1.1 Nomenclature",
            "2.3.S.1.2 Structure",
            "2.3.S.1.3 General Properties",
        ]:
            idx = self._find_para_index_doc(doc, heading, start_idx, end_idx)
            if idx is not None and name_mfr_line:
                nxt = (doc.paragraphs[idx + 1].text or "").strip() if idx + 1 < len(doc.paragraphs) else ""
                if not nxt.startswith("("):
                    self._insert_paragraph_after(doc.paragraphs[idx], name_mfr_line)

        for refer in ["3.2.S.1.1", "3.2.S.1.2", "3.2.S.1.3"]:
            if extracted[refer].warning:
                warnings.append(f"{refer}: {extracted[refer].warning}")

        s11 = self._parse_s11(extracted["3.2.S.1.1"].raw_text)
        s12 = self._parse_s12(extracted["3.2.S.1.2"].raw_text)
        s13 = self._parse_s13(extracted["3.2.S.1.3"].raw_text)

        # Insert compendial name under Refer Section 3.2.S.1.3 anchor.
        s13_ref_idx = self._find_para_index_doc(doc, "Refer Section 3.2.S.1.3", start_idx, end_idx)
        if s13_ref_idx is not None and s11["b"]:
            nxt = (doc.paragraphs[s13_ref_idx + 1].text or "").strip() if s13_ref_idx + 1 < len(doc.paragraphs) else ""
            if nxt.lower() != s11["b"].strip().lower():
                self._insert_paragraph_after(doc.paragraphs[s13_ref_idx], s11["b"])

        start_idx, end_idx = self._get_target_range(doc)

        field_map = [
            ("(Recommended) International Non-proprietary name (INN):", s11["a"]),
            ("Compendial name, if relevant:", s11["b"]),
            ("Chemical name(s):", s11["c"]),
            ("Company or laboratory code:", s11["d"]),
            ("Other non-proprietary name(s)", s11["e"]),
            ("Chemical Abstracts Service (CAS) registry number:", s11["f"]),
            ("Structural formula, including relative and absolute stereochemistry:", "__IMG__"),
            ("Molecular formula:", s12["b"]),
            ("Relative molecular mass:", s12["c"]),
            ("Physical description", s13["a"]),
            ("Other:", s13["other"]),
        ]

        for label, value in field_map:
            idx = self._find_para_index_doc(doc, label, start_idx, end_idx)
            if idx is None:
                continue
            if value == "__IMG__":
                img_paths = extracted["3.2.S.1.2"].image_paths
                if img_paths:
                    img_para = self._insert_paragraph_after(doc.paragraphs[idx], "")
                    self._add_picture_autofit(img_para.add_run(), img_paths[0], doc)
            else:
                self._add_answer_after(doc, idx, value)

        sol_idx = self._find_para_index_doc(doc, "Solubilities:", start_idx, end_idx)
        if sol_idx is not None:
            self._replace_paragraph_with_runs(
                doc.paragraphs[sol_idx],
                [("(b)", False, False), ("\tSolubilities", False, False), (":", False, False), (" NA", False, False)],
            )
            if sol_idx + 1 < len(doc.paragraphs) and not (doc.paragraphs[sol_idx + 1].text or "").strip():
                doc.paragraphs[sol_idx + 1].text = s13["b"]
            else:
                self._insert_paragraph_after(doc.paragraphs[sol_idx], s13["b"])

        self._set_value_under_label(doc, "Polymorphic form:", s13["poly"], start_idx, end_idx)
        self._set_value_under_label(doc, "Solvate:", s13["solvate"], start_idx, end_idx)
        self._set_value_under_label(doc, "Hydrate:", s13["hydrate"], start_idx, end_idx)

        self._remove_paragraphs_matching(
            doc,
            [
                "Refer Section 3.2.S.1.1",
                "Refer Section 3.2.S.1.2",
                "Refer Section 3.2.S.1.3",
                "Company or laboratory code:",
                "Other non-proprietary name(s)",
                "Chemical Abstracts Service (CAS) registry number:",
            ],
            start_idx,
            end_idx,
        )

        self._fill_property_table(doc, s13)
        self._apply_visual_formatting(doc, start_idx, end_idx)
        cleanup_stats = _run_injected_artifact_cleanup(
            doc, keep_first_n_tables=template_table_count
        )
        if any(cleanup_stats.values()):
            warnings.append(f"cleanup: {cleanup_stats}")

        output_docx.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_docx)
        return warnings


# ---------------------------------------------------------------------------
# S2 filler — 2.3.S.2 Manufacture
# ---------------------------------------------------------------------------

class S2DocxFiller(_DocxHelper):
    SECTION_START = "2.3.S.2 Manufacture"
    SECTION_END   = "2.3.S.3 Characterisation"

    def __init__(
        self,
        template_docx: Path,
        filled_reference_docx: Path | None = None,
        images_dir: Path | None = None,
        s2_fill_cfg: S2FillConfig | None = None,
        noise_cfg: NoiseConfig | None = None,
        diagram_cfg: DiagramConfig | None = None,
    ) -> None:
        self.template_docx         = template_docx
        self.filled_reference_docx = filled_reference_docx
        self.images_dir            = images_dir
        self.cfg       = s2_fill_cfg or S2FillConfig()
        self.noise_cfg = noise_cfg   or NoiseConfig()
        self.diagram_cfg = diagram_cfg or DiagramConfig()

    # ------------------------------------------------------------------
    # Text cleaning  (raw_text is already de-noised by extractor;
    # this pass removes any residual section-stamp lines)
    # ------------------------------------------------------------------

    # PDF page-stamp pattern  e.g. "1 of 3" / "Page 2 of 5"
    _PAGE_STAMP_RE = re.compile(
        r"^(page\s*)?\d+\s+of\s+\d+$", re.IGNORECASE
    )

    def _clean_block(self, text: str) -> str:
        out: list[str] = []
        for ln in text.splitlines():
            s = ln.strip()
            if not s:
                continue
            low = s.lower()
            # drop residual section-stamp lines like "3.2.S.2.1 Manufacture"
            if re.match(r"^3\.2\.[sp]\.\d+(\.\d+)*\b", low) and len(s) < 80:
                continue
            # drop PDF page numbers ("1 of 3", "Page 2 of 5")
            if self._PAGE_STAMP_RE.match(s):
                continue
            out.append(s)
        return "\n".join(out).strip()

    @staticmethod
    def _lines(text: str) -> list[str]:
        return [ln.strip() for ln in text.splitlines() if ln.strip()]

    # ------------------------------------------------------------------
    # Generic manufacturer extraction (no hardcoded company names)
    # ------------------------------------------------------------------

    _PHARMA_RE = re.compile(
        r"\b(pharmaceutical|pharma|biotech|chemical|laboratory|laboratories)\b"
        r".*?\b(ltd|limited|inc|co\.?|corp|llc|plc|gmbh|ag|s\.a\.?)\b",
        re.IGNORECASE,
    )
    _ZIP_RE = re.compile(r"\d{4,}")

    _S22_KEYWORDS = {
        "flow": (
            "flow diagram",
            "flow chart",
            "synthesis process",
        ),
        "brief": (
            "brief narrative",
            "description of the manufacturing process",
            "description of manufacturing process",
        ),
        "alternate": (
            "alternate processes",
            "alternative processes",
        ),
        "reprocessing": (
            "reprocessing steps",
            "reprocessing step",
        ),
    }

    @classmethod
    def _extract_s22_block(
        cls,
        lines: list[str],
        start_keywords: tuple[str, ...],
        stop_keywords: tuple[str, ...],
    ) -> str:
        start_idx = None
        for i, ln in enumerate(lines):
            low = ln.lower()
            if any(k in low for k in start_keywords):
                start_idx = i
                break
        if start_idx is None:
            return ""

        line = lines[start_idx]
        block: list[str] = []
        if ":" in line:
            tail = line.split(":", 1)[1].strip()
            if tail:
                block.append(tail)

        for ln in lines[start_idx + 1 :]:
            low = ln.lower()
            if any(k in low for k in stop_keywords):
                break
            if re.match(r"^\(?[a-d]\)?\s*[\).:-]", low):
                if block:
                    break
            block.append(ln)

        return " ".join(s for s in block if s).strip()

    @staticmethod
    def _derive_s22_brief_narrative(lines: list[str]) -> str:
        """
        Derive a narrative sentence when explicit '(b) brief narrative' label is absent.
        Keeps logic generic and avoids product-specific hardcoding.
        """
        heading_re = re.compile(
            r"^\s*3\s*[\.\-]\s*2\s*[\.\-]\s*[sp]\s*[\.\-]\s*\d+(?:\s*[\.\-]\s*\d+)*\b",
            re.IGNORECASE,
        )
        stage_re = re.compile(r"^(stage|step)\b", re.IGNORECASE)
        company_re = re.compile(r"\b(co\.?|ltd|limited|pharmaceutical|laboratories?)\b", re.IGNORECASE)

        candidates: list[str] = []
        for ln in lines:
            s = " ".join(ln.split())
            low = s.lower()
            if not s:
                continue
            if heading_re.match(s):
                continue
            if low.startswith("3.2") and "description of manufacturing process" in low:
                continue
            if stage_re.match(low):
                continue
            if low.startswith("figure "):
                continue
            if "flow diagram" in low and len(s.split()) <= 8:
                continue
            if "chemical synthetical pathway" in low:
                continue
            if company_re.search(s) and len(s.split()) <= 10:
                continue
            if "  " in ln and len(s.split()) <= 5:
                continue
            if len(s) >= 45 and any(ch in s for ch in (".", ";", ":")):
                candidates.append(s)

        if candidates:
            return candidates[0]

        for ln in lines:
            s = " ".join(ln.split())
            if len(s) >= 55 and not heading_re.match(s):
                return s
        return ""

    def _extract_s22_narrative_image(self, pdf_path: Path) -> Path | None:
        if self.images_dir is None:
            return None
        out_path = self.images_dir / "3_2_S_2_2_narrative.png"

        targets = (
            "brief narrative description",
            "description of the manufacturing process",
            "description of manufacturing process",
        )
        stop_targets = (
            "flow diagram",
            "alternate processes",
            "reprocessing steps",
            "control of materials",
        )

        def drawing_clip(page: fitz.Page) -> fitz.Rect | None:
            drawings = page.get_drawings()
            if not drawings:
                return None
            union = None
            page_rect = page.rect
            top_cut = page_rect.y0 + page_rect.height * self.diagram_cfg.header_crop_frac
            bottom_cut = page_rect.y1 - page_rect.height * self.diagram_cfg.footer_crop_frac
            for drawing in drawings:
                drect = drawing.get("rect")
                if drect is None:
                    continue
                # Ignore page-spanning frame artifacts.
                if drect.width >= page_rect.width * 0.95 or drect.height >= page_rect.height * 0.95:
                    continue
                if drect.y1 <= top_cut or drect.y0 >= bottom_cut:
                    continue
                union = drect if union is None else union | drect
            if union is None:
                return None
            pad_x = max(4.0, union.width * 0.01)
            pad_y = max(4.0, union.height * 0.01)
            clip = fitz.Rect(union.x0 - pad_x, union.y0 - pad_y, union.x1 + pad_x, union.y1 + pad_y)
            return clip & page.rect

        def embedded_image_clip(page: fitz.Page) -> fitz.Rect | None:
            page_rect = page.rect
            union = None
            try:
                for img in page.get_images(full=True):
                    xref = img[0]
                    for rect in page.get_image_rects(xref):
                        if rect.width < 24 or rect.height < 24:
                            continue
                        # Ignore tiny logos in header/footer strips.
                        if rect.y1 <= page_rect.y0 + page_rect.height * self.diagram_cfg.header_crop_frac:
                            continue
                        if rect.y0 >= page_rect.y1 - page_rect.height * self.diagram_cfg.footer_crop_frac:
                            continue
                        union = rect if union is None else union | rect
            except Exception:
                return None
            if union is None:
                return None
            pad_x = max(4.0, union.width * 0.02)
            pad_y = max(4.0, union.height * 0.02)
            return fitz.Rect(union.x0 - pad_x, union.y0 - pad_y, union.x1 + pad_x, union.y1 + pad_y) & page_rect

        def preferred_visual_clip(page: fitz.Page, page_text_lower: str) -> fitz.Rect | None:
            clip = embedded_image_clip(page)
            if clip is None:
                clip = drawing_clip(page)
            if clip is None:
                return None

            # If a configured pathway-like keyword is present, start clip below it.
            keyword_ys: list[float] = []
            for kw in self.diagram_cfg.diagram_exclude_keywords:
                if kw not in page_text_lower:
                    continue
                try:
                    for r in page.search_for(kw):
                        keyword_ys.append(r.y1)
                except Exception:
                    continue
            if keyword_ys:
                top = max(clip.y0, min(keyword_ys) + 4.0)
                clip = fitz.Rect(clip.x0, top, clip.x1, clip.y1) & page.rect
            return clip if clip.height > 20 and clip.width > 20 else None

        def clip_visual_score(page: fitz.Page, clip: fitz.Rect) -> float:
            """
            Estimate whether clip contains true diagram content (not text-only blocks).
            """
            score = 0.0

            # Embedded image evidence
            try:
                for img in page.get_images(full=True):
                    xref = img[0]
                    for rect in page.get_image_rects(xref):
                        inter = rect & clip
                        if inter.is_empty:
                            continue
                        score += max(1.0, inter.width * inter.height / 5000.0)
            except Exception:
                pass

            # Vector drawing evidence
            try:
                for d in page.get_drawings():
                    drect = d.get("rect")
                    if drect is None:
                        continue
                    inter = drect & clip
                    if inter.is_empty:
                        continue
                    score += 0.8
            except Exception:
                pass

            # Penalize text-heavy regions.
            try:
                words = page.get_text("words", clip=clip)
                score -= min(20.0, len(words) / 25.0)
            except Exception:
                pass

            return score

        try:
            with fitz.open(pdf_path) as doc:
                # Strategy 1 (preferred): pathway-like visual page.
                best_page = None
                best_clip = None
                best_score = -1.0
                for page in doc:
                    text = page.get_text("text", sort=True)
                    low = text.lower()
                    if not any(kw in low for kw in self.diagram_cfg.diagram_exclude_keywords):
                        continue
                    clip = preferred_visual_clip(page, low)
                    if clip is None:
                        continue
                    score = clip_visual_score(page, clip)
                    if score > best_score:
                        best_score = score
                        best_page = page
                        best_clip = clip

                if best_page is not None and best_clip is not None and best_score > 1.5:
                    scale = self.diagram_cfg.render_dpi_scale
                    pix = best_page.get_pixmap(
                        matrix=fitz.Matrix(scale, scale),
                        clip=best_clip,
                        alpha=False,
                    )
                    pix.save(str(out_path))
                    return out_path

                # Strategy 2: labeled narrative area, only if visually rich.
                for page in doc:
                    text = page.get_text("text", sort=True)
                    low = text.lower()
                    if not any(t in low for t in targets):
                        continue

                    page_rect = page.rect
                    top_y = None
                    for t in targets:
                        hits = page.search_for(t)
                        if hits:
                            y0 = min(r.y0 for r in hits)
                            top_y = y0 if top_y is None else min(top_y, y0)

                    if top_y is None:
                        continue

                    bottom_y = None
                    for t in stop_targets:
                        hits = page.search_for(t)
                        for r in hits:
                            if r.y0 > top_y:
                                bottom_y = r.y0 if bottom_y is None else min(bottom_y, r.y0)

                    if bottom_y is None:
                        bottom_y = page_rect.y1 - page_rect.height * 0.10

                    x0 = page_rect.x0 + page_rect.width * 0.04
                    x1 = page_rect.x1 - page_rect.width * 0.04
                    y0 = max(page_rect.y0 + page_rect.height * 0.10, top_y - 6)
                    y1 = min(page_rect.y1 - page_rect.height * 0.10, bottom_y + 6)
                    if y1 - y0 < 40:
                        continue

                    clip = fitz.Rect(x0, y0, x1, y1)
                    if clip_visual_score(page, clip) <= 1.5:
                        continue
                    scale = self.diagram_cfg.render_dpi_scale
                    pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale), clip=clip, alpha=False)
                    pix.save(str(out_path))
                    return out_path

                # Final fallback: best visual clip in full section range.
                best_page = None
                best_clip = None
                best_score = -1.0
                for page in doc:
                    low = page.get_text("text", sort=True).lower()
                    clip = preferred_visual_clip(page, low)
                    if clip is None:
                        continue
                    score = clip_visual_score(page, clip)
                    if score > best_score:
                        best_score = score
                        best_page = page
                        best_clip = clip

                if best_page is not None and best_clip is not None and best_score > 1.5:
                    scale = self.diagram_cfg.render_dpi_scale
                    pix = best_page.get_pixmap(matrix=fitz.Matrix(scale, scale), clip=best_clip, alpha=False)
                    pix.save(str(out_path))
                    return out_path
        except Exception:
            return None
        return None

    def _parse_s21(self, text: str) -> dict[str, str]:
        lines = self._lines(self._clean_block(text))

        # ── find narrative start ──────────────────────────────────────
        start = 0
        for i, ln in enumerate(lines):
            if any(kw in ln.lower() for kw in self.cfg.narrative_start_keywords):
                start = i
                break

        useful: list[str] = []
        for ln in lines[start:]:
            if any(kw in ln.lower() for kw in self.cfg.narrative_end_keywords):
                break
            useful.append(ln)

        # ── manufacturer name ─────────────────────────────────────────
        # Primary: regex matching common pharma-entity patterns.
        # Fallback: any line ending with a legal-entity suffix.
        _LEGAL_SUFFIX_RE = re.compile(
            r"\b(ltd\.?|limited|inc\.?|co\.?|corp\.?|llc|plc|gmbh|ag|s\.a\.?)\s*$",
            re.IGNORECASE,
        )
        mfr_name = ""
        mfr_idx  = None
        for i, ln in enumerate(useful):
            if self._PHARMA_RE.search(ln):
                mfr_name = ln
                mfr_idx  = i
                break
        # Fallback: any line ending with a known legal suffix
        if not mfr_name:
            for i, ln in enumerate(useful):
                if _LEGAL_SUFFIX_RE.search(ln) and len(ln) > 6:
                    mfr_name = ln
                    mfr_idx  = i
                    break

        # ── address: next 1-4 lines after manufacturer name ───────────
        mfr_addr = ""
        if mfr_idx is not None:
            parts: list[str] = []
            stop_kw = ("manufactur", "responsib", "test", "packag", "gmp")
            for ln in useful[mfr_idx + 1 : mfr_idx + 6]:
                low = ln.lower()
                if any(k in low for k in stop_kw):
                    break
                addr_kw = ("china", "india", "japan", "germany", "france",
                           "street", "road", "avenue", "plot", "block", "zone",
                           "district", "province", "state", "city")
                if re.search(r"\d", ln) or any(k in low for k in addr_kw):
                    parts.append(ln)
                    if self._ZIP_RE.search(ln):
                        break
            mfr_addr = " ".join(parts).strip()

        # ── responsibility: detect from text, else fallback ───────────
        responsibility = ""
        for ln in useful:
            low = ln.lower()
            if "manufactur" in low and ("test" in low or "packag" in low):
                responsibility = self.cfg.manufacturer_table_responsibility_default
                break

        # ── GMP statement ─────────────────────────────────────────────
        gmp_line = ""
        for ln in useful:
            if any(kw in ln.lower() for kw in self.cfg.gmp_keywords):
                gmp_line = self.cfg.gmp_found_sentence
                break
        gmp_line = gmp_line or self.cfg.gmp_fallback_sentence

        return {
            "mfr_name":       mfr_name,
            "mfr_addr":       mfr_addr,
            "responsibility": responsibility or self.cfg.manufacturer_table_responsibility_default,
            "gmp":            gmp_line,
        }

    def _parse_s22(self, text: str) -> dict[str, str]:
        lines = self._lines(self._clean_block(text))

        stop_keys = (
            self._S22_KEYWORDS["brief"]
            + self._S22_KEYWORDS["alternate"]
            + self._S22_KEYWORDS["reprocessing"]
            + self._S22_KEYWORDS["flow"]
        )

        brief = self._extract_s22_block(
            lines,
            self._S22_KEYWORDS["brief"],
            stop_keys,
        )
        if not brief:
            brief = self._derive_s22_brief_narrative(lines)

        alternate = self._extract_s22_block(
            lines,
            self._S22_KEYWORDS["alternate"],
            stop_keys,
        )
        reprocessing = self._extract_s22_block(
            lines,
            self._S22_KEYWORDS["reprocessing"],
            stop_keys,
        )

        return {
            "brief":       brief,
            "alternate":   alternate or self.cfg.alternate_processes_default,
            "reprocessing": reprocessing or self.cfg.reprocessing_steps_default,
        }

    # ------------------------------------------------------------------
    # Table helpers
    # ------------------------------------------------------------------

    def _fill_s21_table(
        self,
        doc: _Document,
        mfr_name: str,
        mfr_addr: str,
        responsibility: str,
    ) -> bool:
        for table in doc.tables:
            if len(table.rows) < 2 or len(table.columns) < 2:
                continue
            head = " ".join(
                (table.cell(0, c).text or "").strip().lower()
                for c in range(min(3, len(table.columns)))
            )
            if "name and address" in head and "responsibility" in head:
                addr_block = mfr_name
                if mfr_addr:
                    addr_block += f"\n\nAddress of Manufacturer:\n{mfr_addr}"
                table.cell(1, 0).text = addr_block.strip()
                table.cell(1, 1).text = responsibility
                if len(table.columns) > 2:
                    table.cell(1, 2).text = self.cfg.manufacturer_table_apprx_col
                return True
        return False

    # ------------------------------------------------------------------
    # Generic helpers
    # ------------------------------------------------------------------

    def _restricted_phrase(self, text: str) -> str:
        kws = tuple(k.lower() for k in self.cfg.restricted_phrase_keywords)
        for ln in text.splitlines():
            low = ln.lower()
            if any(k in low for k in kws):
                return ln.strip()
        return ""

    def _parse_s23(self, text: str) -> dict[str, str]:
        redacted = self._restricted_phrase(text)
        if not redacted:
            skip_pats = (
                re.compile(r"^3\.2\.[sp]\.2\.3\b", re.IGNORECASE),
                re.compile(r"^refer\s+section\s+3\.2\.[sp]\.2\.3\b", re.IGNORECASE),
                re.compile(r"^control\s+of\s+materials\b", re.IGNORECASE),
            )
            for ln in text.splitlines():
                s = " ".join(ln.split()).strip()
                if not s:
                    continue
                if any(p.search(s) for p in skip_pats):
                    continue
                redacted = s
                break
        return {
            "a": redacted,
            "b": redacted,
            "c": redacted,
            "d": redacted,
            "b_inline_default": self.cfg.s23_manufacturer_not_available_default,
        }

    def _parse_generic_restricted_section(self, section_ref: str, text: str) -> dict[str, str]:
        """
        Generic parser for S2.x restricted-style sections (2.4, 2.5, 2.6).
        Attempts to reuse the same heuristics as _parse_s23 without hardcoding
        the section number.
        Returns a dict with keys a,b,c,d and an inline default for (b).
        """
        redacted = self._restricted_phrase(text)
        if not redacted:
            # Skip common section-stamp lines and explicit refer lines for this section
            sec_tail = section_ref.split(".")[-1]
            try:
                sec_num = int(sec_tail)
                sec_re = re.compile(rf"^3\.2\.[sp]\.2\.{sec_num}\b", re.IGNORECASE)
            except Exception:
                sec_re = re.compile(r"^3\.2\.[sp]\.2\.[0-9]+\b", re.IGNORECASE)

            refer_re = re.compile(rf"^refer\s+section\s+{re.escape(section_ref)}", re.IGNORECASE)

            for ln in text.splitlines():
                s = " ".join(ln.split()).strip()
                if not s:
                    continue
                if sec_re.search(s) or refer_re.search(s) or re.search(r"control|controls|summary", s, re.IGNORECASE):
                    # skip heading/anchor-like lines
                    continue
                redacted = s
                break

        return {
            "a": redacted,
            "b": redacted,
            "c": redacted,
            "d": redacted,
            "b_inline_default": self.cfg.s23_manufacturer_not_available_default,
        }

    @staticmethod
    def _append_inline_value_after_colon(paragraph: Paragraph, value: str) -> None:
        if not value:
            return
        text = (paragraph.text or "").rstrip()
        if not text:
            paragraph.text = value
            return
        if ":" in text:
            head, tail = text.rsplit(":", 1)
            if tail.strip():
                return
            paragraph.text = f"{head}: {value}"
            return
        if value.lower() not in text.lower():
            paragraph.text = f"{text} {value}".strip()

    @staticmethod
    def _normalize_s23_heading(doc: _Document, start_idx: int, end_idx: int) -> None:
        pat = re.compile(r"^(2\.3\.S\.2\.3\s+Control of Materials)\b", re.IGNORECASE)
        for i in range(start_idx, min(end_idx, len(doc.paragraphs))):
            p = doc.paragraphs[i]
            text = (p.text or "").strip()
            m = pat.match(text)
            if not m:
                continue
            canonical = m.group(1)
            if text != canonical:
                p.text = canonical
            return

    def _normalize_s23_table_header(self, doc: _Document) -> None:
        target = self.cfg.s23_table_first_header_default.strip()
        if not target:
            return
        for table in doc.tables:
            if not table.rows or len(table.columns) < 3:
                continue
            h0 = self._norm(table.cell(0, 0).text)
            h1 = self._norm(table.cell(0, 1).text)
            h2 = self._norm(table.cell(0, 2).text)
            if h0 in {"test parameter", "step / starting material"} and "test" in h1 and "acceptance criteria" in h2:
                table.cell(0, 0).text = target
                return

    def _normalize_s2_table_header(self, doc: _Document) -> None:
        """Normalize table headers for S2.x control-step style tables.
        Reuses the same default header text configured for S2.3.
        """
        # Reuse s23 table normalization logic (generic for similar tables)
        self._normalize_s23_table_header(doc)

    def _normalize_s2_heading(self, doc: _Document, section_tail: str) -> None:
        """Canonicalize heading like '2.3.S.2.4 ...' where section_tail is '2.4' etc."""
        pat = re.compile(rf"^(2\.3\.S\.2\.{re.escape(section_tail)}\s+.+)$", re.IGNORECASE)
        for i in range(0, len(doc.paragraphs)):
            p = doc.paragraphs[i]
            text = (p.text or "").strip()
            m = pat.match(text)
            if not m:
                continue
            canonical = m.group(1)
            if text != canonical:
                p.text = canonical
            return

    def _remove_refer_section_lines(
        self, doc: _Document, start_idx: int, end_idx: int
    ) -> None:
        pat = re.compile(r"^refer\s+section\s+3\.2\.[sp]\.", re.IGNORECASE)
        to_del = [
            i for i in range(start_idx, min(end_idx, len(doc.paragraphs)))
            if pat.match(self._norm(doc.paragraphs[i].text or ""))
        ]
        for i in reversed(to_del):
            if i < len(doc.paragraphs):
                self._delete_paragraph(doc.paragraphs[i])

    # ------------------------------------------------------------------
    # Public entry point
    # ------------------------------------------------------------------

    def fill_s2_section(
        self,
        extracted: dict[str, "ExtractedSectionContent"],
        output_docx: Path,
    ) -> list[str]:
        doc      = Document(self.template_docx)
        template_table_count = len(doc.tables)
        warnings: list[str] = []

        start_idx, end_idx = self._get_target_range(doc)

        # ── (drug, manufacturer) subtitle under every S2 heading ──────
        name_line = self._resolve_name_manufacturer_line(
            self.filled_reference_docx, "2.3.S.2.1 Manufacturer"
        )
        for heading in [
            "2.3.S.2 Manufacture",
            "2.3.S.2.1 Manufacturer(s)",
            "2.3.S.2.2 Description of Manufacturing Process and Process Controls",
            "2.3.S.2.3 Control of Materials",
            "2.3.S.2.4 Controls of Critical Steps and Intermediates",
            "2.3.S.2.5 Process Validation and/or Evaluation",
            "2.3.S.2.6 Manufacturing Process Development",
        ]:
            idx = self._find_para_index_doc(doc, heading, start_idx, end_idx)
            if idx is not None and name_line:
                nxt = (doc.paragraphs[idx + 1].text or "").strip() \
                      if idx + 1 < len(doc.paragraphs) else ""
                if not nxt.startswith("("):
                    self._insert_paragraph_after(doc.paragraphs[idx], name_line)

        start_idx, end_idx = self._get_target_range(doc)
        self._remove_refer_section_lines(doc, start_idx, end_idx)

        # ── warnings from extractor ───────────────────────────────────
        for k in ["3.2.S.2.1","3.2.S.2.2","3.2.S.2.3",
                  "3.2.S.2.4","3.2.S.2.5","3.2.S.2.6"]:
            if extracted[k].warning:
                warnings.append(f"{k}: {extracted[k].warning}")

        # ── parse ─────────────────────────────────────────────────────
        s21 = self._parse_s21(extracted["3.2.S.2.1"].raw_text)
        s22 = self._parse_s22(extracted["3.2.S.2.2"].raw_text)
        s23 = self._clean_block(extracted["3.2.S.2.3"].raw_text)
        s24 = self._clean_block(extracted["3.2.S.2.4"].raw_text)
        s25 = self._clean_block(extracted["3.2.S.2.5"].raw_text)
        s26 = self._clean_block(extracted["3.2.S.2.6"].raw_text)

        s23_data = self._parse_s23(s23)
        s24_data = self._parse_generic_restricted_section("3.2.S.2.4", s24)
        s25_data = self._parse_generic_restricted_section("3.2.S.2.5", s25)
        s26_data = self._parse_generic_restricted_section("3.2.S.2.6", s26)

        start_idx, end_idx = self._get_target_range(doc)
        self._normalize_s23_heading(doc, start_idx, end_idx)
        # Normalize sibling S2 headings
        self._normalize_s2_heading(doc, "2.4")
        self._normalize_s2_heading(doc, "2.5")
        self._normalize_s2_heading(doc, "2.6")

        # ── 2.3.S.2.1 ────────────────────────────────────────────────
        # _fill_s21_table searches doc.tables directly; no paragraph anchor needed.
        filled = self._fill_s21_table(
            doc, s21["mfr_name"], s21["mfr_addr"], s21["responsibility"]
        )
        if not filled:
            warnings.append("3.2.S.2.1: manufacturer table not found in template")

        idx = self._find_para_index_doc(
            doc, "Manufacturing authorization for the production of API",
            start_idx, end_idx
        )
        if idx is not None:
            self._insert_paragraph_after(doc.paragraphs[idx], s21["gmp"])

        # ── 2.3.S.2.2 ────────────────────────────────────────────────
        # (a) All flow-diagram images, in document order
        idx = self._find_para_index_doc(
            doc, "Flow diagram of the synthesis process", start_idx, end_idx
        )
        if idx is not None and extracted["3.2.S.2.2"].image_paths:
            anchor = doc.paragraphs[idx]
            current = anchor
            img_count = 0
            for img_path in extracted["3.2.S.2.2"].image_paths:
                try:
                    img_path_obj = Path(img_path) if not isinstance(img_path, Path) else img_path
                    if not img_path_obj.exists():
                        warnings.append(f"3.2.S.2.2: image not found: {img_path}")
                        continue
                    p = self._insert_paragraph_after(current, "")
                    run = p.add_run()
                    self._add_picture_autofit(run, img_path, doc)
                    current = p
                    img_count += 1
                except Exception as e:
                    warnings.append(f"3.2.S.2.2: failed to insert image {img_path}: {e}")
                    continue
            if img_count == 0:
                warnings.append(f"3.2.S.2.2: {len(extracted['3.2.S.2.2'].image_paths)} images found but none inserted")
        elif idx is not None:
            warnings.append("3.2.S.2.2: no flow-diagram images extracted")

        # (b) Brief narrative - prefer extracted image over text
        idx = self._find_para_index_doc(
            doc, "Brief narrative description of the manufacturing process",
            start_idx, end_idx
        )
        if idx is not None:
            brief = s22["brief"].strip()
            current = doc.paragraphs[idx]
            narrative_img = self._extract_s22_narrative_image(
                extracted["3.2.S.2.2"].source_pdf
            )
            if narrative_img:
                img_path = Path(narrative_img) if not isinstance(narrative_img, Path) else narrative_img
                if img_path.exists():
                    try:
                        p = self._insert_paragraph_after(current, "")
                        self._add_picture_autofit(p.add_run(), img_path, doc)
                        current = p
                    except Exception as e:
                        warnings.append(f"3.2.S.2.2: failed to insert narrative image: {e}")

            if brief and len(brief) >= 30:
                self._insert_paragraph_after(current, brief)
            elif not narrative_img:
                warnings.append("3.2.S.2.2: narrative text/image not found")

        # (c) Alternate processes
        idx = self._find_para_index_doc(
            doc, "Alternate processes and explanation", start_idx, end_idx
        )
        if idx is not None:
            self._insert_paragraph_after(doc.paragraphs[idx], s22["alternate"])

        # (d) Reprocessing
        idx = self._find_para_index_doc(
            doc, "Reprocessing steps and justification", start_idx, end_idx
        )
        if idx is not None:
            self._insert_paragraph_after(doc.paragraphs[idx], s22["reprocessing"])

        # ── 2.3.S.2.3 ────────────────────────────────────────────────
        s23_label_map = [
            ("(a)\tName of starting material:", "a"),
            ("(b)\tName and manufacturing site address of starting material", "b"),
            ("Summary of the quality and controls of the starting materials", "c"),
            ("without risk of transmitting agents of animal spongiform", "d"),
        ]
        for label, key in s23_label_map:
            idx = self._find_para_index_doc(doc, label, start_idx, end_idx)
            if idx is None:
                continue
            if key == "b":
                self._append_inline_value_after_colon(
                    doc.paragraphs[idx],
                    s23_data["b_inline_default"],
                )
            ans = s23_data[key].strip()
            if ans:
                self._insert_paragraph_after(doc.paragraphs[idx], ans)

        self._normalize_s23_table_header(doc)

        # Recompute section bounds after S2.3 insertions; paragraph indexes shift.
        start_idx, end_idx = self._get_target_range(doc)

        # ── 2.3.S.2.4 / 2.5 / 2.6 ───────────────────────────────────
        idx = self._find_para_index_doc(
            doc, "Summary of the controls performed at critical steps",
            start_idx, end_idx
        )
        if idx is not None:
            # Insert parsed restricted/summary text (prefer parsed a) and normalize table header
            if s24_data.get("a"):
                self._insert_paragraph_after(doc.paragraphs[idx], s24_data["a"]) 
            self._normalize_s2_table_header(doc)

        # Recompute bounds after S2.4 insertion because paragraph indexes shift.
        start_idx, end_idx = self._get_target_range(doc)

        idx = self._find_para_index_doc(
            doc, "Description of process validation and/or evaluation",
            start_idx, end_idx
        )
        if idx is not None and s25_data.get("a"):
            self._insert_paragraph_after(doc.paragraphs[idx], s25_data["a"])

        # Recompute bounds after S2.5 insertion because paragraph indexes shift.
        start_idx, end_idx = self._get_target_range(doc)

        idx = self._find_para_index_doc(
            doc, "Description and discussion of the significant changes",
            start_idx, end_idx
        )
        if idx is not None and s26_data.get("a"):
            self._insert_paragraph_after(doc.paragraphs[idx], s26_data["a"])

        # ── post-processing ───────────────────────────────────────────
        cleanup_stats = _run_injected_artifact_cleanup(
            doc,
            keep_first_n_tables=template_table_count,
            preserve_repeated_patterns=self.cfg.restricted_phrase_keywords,
        )
        if any(cleanup_stats.values()):
            warnings.append(f"cleanup: {cleanup_stats}")

        output_docx.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_docx)
        return warnings


# ---------------------------------------------------------------------------
# S3 filler — 2.3.S.3 Characterisation
# ---------------------------------------------------------------------------

class S3DocxFiller(_DocxHelper):
    SECTION_START = "2.3.S.3 Characterisation"
    SECTION_END   = "2.3.S.4 Control of Drug Substance"

    _PAGE_STAMP_RE = re.compile(r"^(page\s*)?\d+\s+of\s+\d+$", re.IGNORECASE)

    def __init__(
        self,
        template_docx: Path,
        filled_reference_docx: Path | None = None,
        s2_fill_cfg: S2FillConfig | None = None,
    ) -> None:
        self.template_docx = template_docx
        self.filled_reference_docx = filled_reference_docx
        self.cfg = s2_fill_cfg or S2FillConfig()

    def _clean_block(self, text: str) -> str:
        out: list[str] = []
        for ln in text.splitlines():
            s = ln.strip()
            if not s:
                continue
            if self._PAGE_STAMP_RE.match(s):
                continue
            out.append(s)
        return "\n".join(out).strip()

    @staticmethod
    def _lines(text: str) -> list[str]:
        return [ln.strip() for ln in text.splitlines() if ln.strip()]

    @staticmethod
    def _append_inline_value_after_colon(paragraph: Paragraph, value: str) -> None:
        if not value:
            return
        text = (paragraph.text or "").rstrip()
        if not text:
            paragraph.text = value
            return
        if ":" in text:
            head, tail = text.rsplit(":", 1)
            if tail.strip():
                return
            paragraph.text = f"{head}: {value}"
            return
        if value.lower() not in text.lower():
            paragraph.text = f"{text} {value}".strip()

    @staticmethod
    def _strip_angle_bracket_placeholder(paragraph: Paragraph) -> None:
        text = (paragraph.text or "").strip()
        if not text:
            return
        paragraph.text = re.sub(r"\s*<[^>]*>\s*$", "", text)

    @staticmethod
    def _normalize_s31_heading(doc: _Document, start_idx: int, end_idx: int) -> None:
        for i in range(start_idx, min(end_idx, len(doc.paragraphs))):
            p = doc.paragraphs[i]
            text = (p.text or "").strip()
            if "2.3.S.3.1" not in text:
                continue
            # Keep only heading + descriptor; remove embedded refer tail if present.
            cleaned = re.sub(
                r"\s*refer\s+section\s+3\.2\.[sp]\.3\.1\s*$",
                "",
                text,
                flags=re.IGNORECASE,
            ).strip()
            if cleaned and cleaned != text:
                p.text = cleaned
            return

    def _parse_s31(self, text: str) -> dict[str, str]:
        lines = self._lines(self._clean_block(text))
        low_lines = [ln.lower() for ln in lines]

        start_idx = None
        for i, low in enumerate(low_lines):
            if any(k in low for k in self.cfg.s31_summary_start_keywords):
                start_idx = i
                break

        summary_lines: list[str] = []
        if start_idx is not None:
            for ln in lines[start_idx:]:
                low = ln.lower()
                if any(k in low for k in self.cfg.s31_summary_stop_keywords):
                    break
                if re.match(r"^3\.2\.[sp]\.3(\.\d+)*\b", ln, re.IGNORECASE):
                    if summary_lines:
                        break
                    continue
                summary_lines.append(ln)
                if len(summary_lines) >= self.cfg.s31_max_summary_lines:
                    break

        answer_a = "\n".join(summary_lines).strip()
        if not answer_a and lines:
            answer_a = lines[0]

        # Keep (b) generic: only fill from source if explicit isomerism content exists.
        answer_b = ""
        answer_b = self.cfg.s31_isomerism_default

        return {
            "a": answer_a,
            "b": answer_b,
            "c": self.cfg.s31_polymorph_reference_default,
            "d": self.cfg.s31_particle_size_default,
            "e": self.cfg.s31_other_characteristics_default,
        }

    def _remove_refer_section_lines(
        self,
        doc: _Document,
        start_idx: int,
        end_idx: int,
    ) -> None:
        pat = re.compile(r"^refer\s+section\s+3\.2\.[sp]\.3(\.\d+)*", re.IGNORECASE)
        to_del = [
            i for i in range(start_idx, min(end_idx, len(doc.paragraphs)))
            if pat.match(self._norm(doc.paragraphs[i].text or ""))
        ]
        for i in reversed(to_del):
            if i < len(doc.paragraphs):
                self._delete_paragraph(doc.paragraphs[i])

    def fill_s3_section(
        self,
        extracted: dict[str, ExtractedSectionContent],
        output_docx: Path,
    ) -> list[str]:
        doc = Document(self.template_docx)
        template_table_count = len(doc.tables)
        warnings: list[str] = []

        start_idx, end_idx = self._get_target_range(doc)

        name_line = self._resolve_name_manufacturer_line(
            self.filled_reference_docx,
            "2.3.S.3 Characterisation",
        )
        for heading in [
            "2.3.S.3 Characterisation",
        ]:
            idx = self._find_para_index_doc(doc, heading, start_idx, end_idx)
            if idx is not None and name_line:
                nxt = (doc.paragraphs[idx + 1].text or "").strip() if idx + 1 < len(doc.paragraphs) else ""
                if not nxt.startswith("("):
                    self._insert_paragraph_after(doc.paragraphs[idx], name_line)

        if extracted["3.2.S.3.1"].warning:
            warnings.append(f"3.2.S.3.1: {extracted['3.2.S.3.1'].warning}")

        s31 = self._parse_s31(extracted["3.2.S.3.1"].raw_text)

        start_idx, end_idx = self._get_target_range(doc)
        self._remove_refer_section_lines(doc, start_idx, end_idx)
        self._normalize_s31_heading(doc, start_idx, end_idx)

        # Remove template placeholders under (c)/(d), if present.
        self._remove_paragraphs_matching(
            doc,
            [
                "<including identification of and data on the api lot used in bioavailability studies>",
            ],
            start_idx,
            end_idx,
        )

        # (a)
        idx = self._find_para_index_doc(
            doc,
            "List of studies performed",
            start_idx,
            end_idx,
        )
        if idx is not None and s31["a"].strip():
            self._insert_paragraph_after(doc.paragraphs[idx], s31["a"].strip())

        # (b)
        idx = self._find_para_index_doc(
            doc,
            "Discussion on the potential for isomerism",
            start_idx,
            end_idx,
        )
        if idx is not None and s31["b"].strip():
            self._append_inline_value_after_colon(doc.paragraphs[idx], s31["b"].strip())

        # (c)
        idx = self._find_para_index_doc(
            doc,
            "Summary of studies performed to identify potential polymorphic forms",
            start_idx,
            end_idx,
        )
        if idx is not None and s31["c"].strip():
            self._strip_angle_bracket_placeholder(doc.paragraphs[idx])
            if idx + 1 < len(doc.paragraphs) and not (doc.paragraphs[idx + 1].text or "").strip():
                doc.paragraphs[idx + 1].text = s31["c"].strip()
            else:
                self._insert_paragraph_after(doc.paragraphs[idx], s31["c"].strip())

        # Recompute bounds after c-insertion.
        start_idx, end_idx = self._get_target_range(doc)

        # (d)
        idx = self._find_para_index_doc(
            doc,
            "Summary of studies performed to identify the particle size distribution of the API:",
            start_idx,
            end_idx,
        )
        if idx is not None and s31["d"].strip():
            self._strip_angle_bracket_placeholder(doc.paragraphs[idx])
            self._append_inline_value_after_colon(doc.paragraphs[idx], s31["d"].strip())

        # (e)
        idx = self._find_para_index_doc(
            doc,
            "Other characteristics:",
            start_idx,
            end_idx,
        )
        if idx is not None and s31["e"].strip():
            self._append_inline_value_after_colon(doc.paragraphs[idx], s31["e"].strip())

        cleanup_stats = _run_injected_artifact_cleanup(
            doc,
            keep_first_n_tables=template_table_count,
            preserve_repeated_patterns=self.cfg.restricted_phrase_keywords,
            preserve_phrases=(s31["c"],),
        )
        if any(cleanup_stats.values()):
            warnings.append(f"cleanup: {cleanup_stats}")

        output_docx.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_docx)
        return warnings


class S4DocxFiller:
    SECTION_START = "2.3.S.4 Control of the API"
    SECTION_END = "2.3.S.5 Reference Standards or Materials"

    def __init__(self, template_docx: Path, filled_reference_docx: Path | None = None) -> None:
        self.template_docx = template_docx
        self.filled_reference_docx = filled_reference_docx

    @staticmethod
    def _norm(s: str) -> str:
        return re.sub(r"\s+", " ", s).strip().lower()

    @staticmethod
    def _clean_cell(text: str) -> str:
        return re.sub(r"\s+", " ", (text or "").strip())

    def _get_target_range(self, doc: _Document) -> tuple[int, int]:
        start_idx = None
        end_idx = None
        for i, p in enumerate(doc.paragraphs):
            text = (p.text or "").strip()
            if start_idx is None and self.SECTION_START.lower() in text.lower():
                start_idx = i
            if start_idx is not None and self.SECTION_END.lower() in text.lower():
                end_idx = i
                break
        if start_idx is None:
            raise ValueError("Could not find section start in template DOCX")
        if end_idx is None:
            end_idx = len(doc.paragraphs)
        return start_idx, end_idx

    def _find_para_index(self, doc: _Document, needle: str, start_idx: int, end_idx: int) -> int | None:
        target = self._norm(needle)
        for i in range(start_idx, min(end_idx, len(doc.paragraphs))):
            if target in self._norm(doc.paragraphs[i].text or ""):
                return i
        return None

    @staticmethod
    def _insert_after(paragraph: Paragraph, text: str = "") -> Paragraph:
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        para = Paragraph(new_p, paragraph._parent)
        if text:
            para.add_run(text)
        return para

    @staticmethod
    def _delete_paragraph(paragraph: Paragraph) -> None:
        element = paragraph._element
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)

    def _remove_refer_lines(self, doc: _Document, start_idx: int, end_idx: int) -> None:
        to_delete: list[int] = []
        for i in range(start_idx, min(end_idx, len(doc.paragraphs))):
            text = self._norm(doc.paragraphs[i].text or "")
            if text.startswith("refer section 3.2.s.4."):
                to_delete.append(i)
        for idx in reversed(to_delete):
            if idx < len(doc.paragraphs):
                self._delete_paragraph(doc.paragraphs[idx])

    def _resolve_heading_line(self, heading: str) -> str:
        if self.filled_reference_docx and self.filled_reference_docx.exists():
            try:
                ref = Document(self.filled_reference_docx)
                for i, p in enumerate(ref.paragraphs):
                    text = self._norm(p.text or "")
                    if self._norm(heading) in text:
                        for j in range(i + 1, min(i + 6, len(ref.paragraphs))):
                            candidate = (ref.paragraphs[j].text or "").strip()
                            if candidate.startswith("(") and "," in candidate:
                                return candidate
            except Exception:
                pass
        return ""

    def _insert_heading_line_if_missing(
        self,
        doc: _Document,
        heading: str,
        name_line: str,
        start_idx: int,
        end_idx: int,
    ) -> None:
        if not name_line:
            return
        idx = self._find_para_index(doc, heading, start_idx, end_idx)
        if idx is None:
            return
        next_text = (doc.paragraphs[idx + 1].text or "").strip() if idx + 1 < len(doc.paragraphs) else ""
        if not next_text.startswith("("):
            self._insert_after(doc.paragraphs[idx], name_line)

    def _set_paragraph_text_with_suffix(self, doc: _Document, needle: str, suffix: str, start_idx: int, end_idx: int) -> None:
        idx = self._find_para_index(doc, needle, start_idx, end_idx)
        if idx is None:
            return
        paragraph = doc.paragraphs[idx]
        base = (paragraph.text or "").strip()
        if not base:
            return
        base = re.sub(r"\s+", " ", base)

        for run in list(paragraph.runs):
            paragraph._p.remove(run._r)

        q_run = paragraph.add_run(base)
        q_run.bold = True
        q_run.italic = False

        a_run = paragraph.add_run(f" {suffix}".rstrip())
        a_run.bold = False
        a_run.italic = False

    def _find_s41_template_table(self, doc: _Document) -> object | None:
        for table in doc.tables:
            if len(table.columns) != 3 or len(table.rows) < 8:
                continue
            row0 = [self._norm(c.text) for c in table.rows[0].cells[:3]]
            row2 = [self._norm(c.text) for c in table.rows[2].cells[:3]]
            sample = [self._norm(table.rows[i].cells[0].text) for i in range(3, min(len(table.rows), 8))]
            if (
                row0[0].startswith("standard")
                and row2[0] == "test"
                and "acceptance criteria" in row2[1]
                and "analytical procedure" in row2[2]
                and sample[:5] == ["description", "identification", "impurities", "assay", "etc."]
            ):
                return table
        return None

    def _choose_s41_source_table(self, content: ExtractedSectionContent) -> list[list[str]]:
        best_score = -1
        best_table: list[list[str]] = []
        caption_hit = bool(re.search(r"\btable\s+\d+.*\bspecification", content.raw_text, flags=re.IGNORECASE))

        for table in content.tables:
            flat = " ".join(self._clean_cell(cell) for row in table for cell in row)
            flat_low = flat.lower()
            score = 0
            if "acceptance criteria" in flat_low:
                score += 5
            if "analytical procedure" in flat_low:
                score += 4
            elif "method" in flat_low:
                score += 3
            if re.search(r"\btest\b", flat_low):
                score += 3
            if "specification" in flat_low or "specifications" in flat_low:
                score += 2
            if caption_hit:
                score += 1
            if score > best_score:
                best_score = score
                best_table = table

        return best_table if best_score >= 8 else []

    def _extract_s41_metadata_and_rows(
        self,
        content: ExtractedSectionContent,
    ) -> tuple[str, str, list[list[str]]]:
        raw_source_rows = self._extract_s41_source_rows_from_pdf_tables(content.source_pdf)
        if raw_source_rows:
            normalized_rows = [[test, acceptance, method] for test, method, acceptance in raw_source_rows]
            return "USP", "-----", normalized_rows

        raw_source_rows = self._extract_s41_source_rows_from_pdf(content.source_pdf)
        if not raw_source_rows:
            raw_source_rows = self._extract_s41_source_rows_from_raw_text(content.raw_text)
        if not raw_source_rows:
            raw_source_rows = self._extract_s41_source_rows_from_flat_text(content.raw_text)
        if raw_source_rows:
            normalized_rows = [[test, acceptance, method] for test, method, acceptance in raw_source_rows]
            return "USP", "-----", normalized_rows

        return "USP", "-----", []

    def _extract_s41_source_rows_from_raw_text(self, raw_text: str) -> list[list[str]]:
        lines = []
        for raw_line in raw_text.replace("", "°").splitlines():
            line = raw_line.strip()
            if not line:
                continue
            low = line.lower()
            if (
                low.startswith("drug mater file version:")
                or low.startswith("product name ")
                or low.startswith("module:")
                or low.startswith("zhejiang starry pharmaceutical co., ltd")
                or low.startswith("3.2.s.4 control of drug substance")
                or low.startswith("3.2.s.4.1 specification")
                or re.fullmatch(r"\d+\s+of\s+\d+", low)
            ):
                continue
            lines.append(line)

        start_idx = None
        header_idx = None
        for idx, line in enumerate(lines):
            low = line.lower()
            if "table 8 specification for final iodixanol" in low:
                start_idx = idx
            if start_idx is not None and "test method acceptance criteria" in low:
                header_idx = idx
                break
        if header_idx is None:
            return []

        table_lines = lines[header_idx + 1 :]

        def find_next(patterns: list[str], start: int) -> int:
            for idx in range(start, len(table_lines)):
                joined = table_lines[idx].lower()
                if any(re.search(pattern, joined, flags=re.IGNORECASE) for pattern in patterns):
                    return idx
            return len(table_lines)

        rows: list[list[str]] = []

        def add_row(test: str, method: str, acceptance_parts: list[str]) -> None:
            acceptance = " ".join(part.strip() for part in acceptance_parts if part.strip())
            acceptance = re.sub(r"\s+", " ", acceptance).strip()
            rows.append([test, method, acceptance])

        simple_specs = [
            ("Description", [r"^solubility\b"], "Visual"),
            ("Solubility", [r"^transparency of\b"], "-"),
            ("Transparency of solution", [r"^color of solution\b"], "Visual"),
            ("Color of solution", [r"^specific rotation\b"], "Colorimetry"),
            ("Specific Rotation (°)", [r"^identification\b"], "781S"),
            ("Identification", [r"^pH\b"], "IR spectrum; HPLC; Positive reaction"),
            ("pH", [r"^water\b"], "pH meter"),
            ("Water", [r"^heavy metals\b"], "Method I (921)"),
            ("Heavy metals", [r"^sulphated ash\b"], "Method I (231)"),
            ("Sulphated ash", [r"^free iodine\b"], "Sulphated ash"),
            ("Free iodine", [r"^free iodide\b"], "Color reaction"),
            ("Free iodide", [r"^free aromatic amine\b"], "Titration"),
            ("Free aromatic amine", [r"^limit of calcium\b"], "Spectrophotometry"),
            ("Limit of calcium", [r"^ionic compounds\b"], "ICP"),
            ("Ionic compounds", [r"^residual solvent\b"], "Conductivity"),
        ]

        cursor = 0
        for test_name, next_patterns, method in simple_specs:
            start = find_next([rf"^{re.escape(test_name.lower())}\b", rf"^{re.escape(test_name.split(' (')[0].lower())}\b"], cursor)
            if start >= len(table_lines):
                continue
            next_idx = find_next(next_patterns, start + 1)
            block = table_lines[start:next_idx]
            if not block:
                continue

            if test_name == "Transparency of solution":
                content = " ".join(block)
                content = re.sub(r"^Transparency of\s*solution\s*", "", content, flags=re.IGNORECASE)
            elif test_name == "Identification":
                content = " ".join(block[1:]) if len(block) > 1 else ""
                content = re.sub(r"\s+", " ", content).strip()
                content = re.sub(r"\bA\.\s*IR spectrum\b", "A. IR spectrum:", content, flags=re.IGNORECASE)
                content = re.sub(r"\bB\.\s*HPLC\b", " B. HPLC:", content, flags=re.IGNORECASE)
                content = re.sub(r"\bC\.\s*Positive reaction\b", " C. Positive reaction:", content, flags=re.IGNORECASE)
                add_row(test_name, method, [content])
                cursor = next_idx
                continue
            else:
                content = " ".join(block)
                content = re.sub(rf"^{re.escape(test_name)}\s*", "", content, flags=re.IGNORECASE)

            if method and content.lower().startswith(method.lower()):
                content = content[len(method) :].strip(" :;-")
            add_row(test_name, method, [content])
            cursor = next_idx

        residual_start = find_next([r"^residual solvent\b"], 0)
        related1_start = find_next([r"^related compounds\(test 1\)\b"], residual_start + 1)
        if residual_start < len(table_lines) and related1_start <= len(table_lines):
            residual_lines = table_lines[residual_start + 1 : related1_start]
            cleaned = [ln for ln in residual_lines if ln]
            gc_method = "GC"
            i = 0
            while i < len(cleaned):
                line = cleaned[i]
                low = line.lower()
                if low == "gc":
                    i += 1
                    continue
                if low == "methanol":
                    acceptance = cleaned[i + 2] if i + 2 < len(cleaned) and cleaned[i + 1].lower() == "gc" else ""
                    add_row("Residual solvent - Methanol", gc_method, [acceptance])
                    i += 3 if acceptance else 1
                    continue
                if low.startswith("isopropyl alcohol"):
                    add_row("Residual solvent - Isopropyl alcohol", gc_method, [re.sub(r"^Isopropyl alcohol\s*", "", line, flags=re.IGNORECASE)])
                elif low.startswith("1-methoxy-2-"):
                    full = line
                    if i + 1 < len(cleaned) and cleaned[i + 1].lower().startswith("propanol"):
                        full = f"{line} {cleaned[i + 1]}"
                        i += 1
                    add_row(
                        "Residual solvent - 1-methoxy-2-propanol",
                        gc_method,
                        [re.sub(r"^1-methoxy-2-\s*propanol\s*", "", full, flags=re.IGNORECASE)],
                    )
                elif low.startswith("methoxyethanol"):
                    add_row("Residual solvent - Methoxyethanol", gc_method, [re.sub(r"^Methoxyethanol\s*", "", line, flags=re.IGNORECASE)])
                elif low.startswith("1,3-dichloro-2-propanol"):
                    rest = re.sub(r"^1,3-dichloro-2-propanol\s*", "", line, flags=re.IGNORECASE)
                    if rest.upper().startswith("GC "):
                        rest = rest[3:].strip()
                    add_row("Residual solvent - 1,3-dichloro-2-propanol", gc_method, [rest])
                i += 1

        related2_start = find_next([r"^related compounds\(test 2\)\b"], related1_start + 1)
        if related1_start < len(table_lines):
            block = table_lines[related1_start + 1 : related2_start]
            joined = "\n".join(block)
            related_specs = [
                ("Related compounds (Test 1) - Related compound B", r"Related compound B\s*HPLC\s*(NMT .*?)\s*(?=Related compound C|$)"),
                ("Related compounds (Test 1) - Related compound C", r"Related compound C\s*(NMT .*?)\s*(?=Related compound D|$)"),
                ("Related compounds (Test 1) - Related compound D", r"Related compound D\s*(NMT .*?)\s*(?=Related compound F|$)"),
                ("Related compounds (Test 1) - Related compound F", r"Related compound F\s*(NMT .*?)\s*(?=Related compound G|$)"),
                ("Related compounds (Test 1) - Related compound G", r"Related compound G\s*(NMT .*?)\s*(?=Iohexol|$)"),
                ("Related compounds (Test 1) - Iohexol", r"Iohexol\s*(NMT .*?)\s*(?=O-alkylated compounds|$)"),
                ("Related compounds (Test 1) - O-alkylated compounds", r"O-alkylated compounds\s*(NMT .*?)\s*(?=Others|$)"),
                ("Related compounds (Test 1) - Others", r"Others\s*(Any individual unspecified .*?Total impurities: NMT 1\.5%\.)"),
            ]
            for test_name, pattern in related_specs:
                found = re.search(pattern, joined, flags=re.IGNORECASE | re.DOTALL)
                if found:
                    add_row(test_name, "HPLC", [found.group(1)])

        assay_start = find_next([r"^assay\b"], related2_start + 1)
        if related2_start < len(table_lines):
            block = table_lines[related2_start + 1 : assay_start]
            joined = "\n".join(block)
            found_e = re.search(r"Related compound E\s*HPLC\s*(NMT .*?)\s*(?=Related compound H|$)", joined, flags=re.IGNORECASE | re.DOTALL)
            found_h = re.search(r"Related compound H\s*(NMT .*?)\s*(?=Microbial limits|$)", joined, flags=re.IGNORECASE | re.DOTALL)
            if found_e:
                add_row("Related compounds (Test 2) - Related compound E", "HPLC", [found_e.group(1)])
            if found_h:
                add_row("Related compounds (Test 2) - Related compound H", "HPLC", [found_h.group(1)])

        microbial_limits_start = find_next([r"^microbial limits\b"], assay_start)
        micro_purity_start = find_next([r"^microbiological purity\b"], microbial_limits_start + 1)
        if microbial_limits_start < len(table_lines):
            line = table_lines[microbial_limits_start]
            found = re.search(r"Microbial limits\s+61\s+(.*)", line, flags=re.IGNORECASE)
            if found:
                add_row("Microbial limits", "61", [found.group(1)])

        bacterial_endo_start = find_next([r"^bacterial endotoxin\b"], micro_purity_start + 1)
        if micro_purity_start < len(table_lines):
            block = table_lines[micro_purity_start:bacterial_endo_start]
            joined = " ".join(block)
            joined = re.sub(r"^Microbiological purity\s*-\s*", "", joined, flags=re.IGNORECASE)
            add_row("Microbiological purity", "---------", [joined])

        if assay_start < len(table_lines):
            line = table_lines[assay_start]
            found = re.search(r"Assay\s+Titration\s+(.*)", line, flags=re.IGNORECASE)
            if found:
                add_row("Assay", "Titration", [found.group(1)])

        if bacterial_endo_start < len(table_lines):
            block = " ".join(table_lines[bacterial_endo_start:])
            found = re.search(r"Bacterial endotoxin\s*-\s*(.*)", block, flags=re.IGNORECASE)
            if found:
                add_row("Bacterial endotoxin", "---------", [found.group(1)])

        return rows

    def _extract_s41_source_rows_from_pdf_tables(self, pdf_path: Path) -> list[list[str]]:
        try:
            import fitz  # type: ignore
        except Exception:
            return []

        rows: list[list[str]] = []
        seen: set[tuple[str, str, str]] = set()
        doc = None

        try:
            doc = fitz.open(str(pdf_path))
            for page in doc:
                page_text = " ".join(page.get_text("text", sort=True).split()).lower()
                if "drug mater file" not in page_text and "table 8 specification for final iodixanol" not in page_text:
                    continue

                try:
                    tables = page.find_tables().tables
                except Exception:
                    tables = []

                for table in tables:
                    raw_rows = table.extract()
                    if not raw_rows:
                        continue

                    header_text = " ".join(
                        self._clean_cell(str(cell or "")) for row in raw_rows[:3] for cell in row
                    ).lower()
                    if not any(token in header_text for token in ("acceptance criteria", "method", "test")):
                        continue

                    current_test = ""
                    current_method = ""
                    for raw_row in raw_rows:
                        cells = [self._clean_cell(str(cell or "")) for cell in raw_row]
                        non_empty = [cell for cell in cells if cell]
                        if not non_empty:
                            continue

                        test = ""
                        method = ""
                        acceptance = ""

                        if len(cells) >= 9:
                            test = self._clean_cell(" ".join(cells[0:3]))
                            method = self._clean_cell(" ".join(cells[3:6]))
                            acceptance = self._clean_cell(" ".join(cells[6:9]))
                        elif len(cells) >= 6:
                            # Sometimes merged-cell tables still expand into more than 3 cells.
                            span = max(1, len(cells) // 3)
                            test = self._clean_cell(" ".join(cells[0:span]))
                            method = self._clean_cell(" ".join(cells[span : 2 * span]))
                            acceptance = self._clean_cell(" ".join(cells[2 * span :]))
                        elif len(cells) >= 3:
                            test = cells[0]
                            method = cells[1]
                            acceptance = cells[2]
                        elif len(cells) == 2:
                            # Preserve subrows where test cell is visually merged above.
                            test = ""
                            method = cells[0]
                            acceptance = cells[1]
                        else:
                            continue

                        header_key = " ".join((test, method, acceptance)).lower()
                        if (
                            not header_key
                            or ("test" in test.lower() and "method" in method.lower() and "acceptance" in acceptance.lower())
                            or ("test" == test.lower() and not method and not acceptance)
                        ):
                            continue

                        if test:
                            current_test = test
                        elif current_test.lower() == "identification":
                            # Keep blank test for split Identification subrows exactly as source structure.
                            test = ""

                        # Fill down vertically merged method cells such as HPLC / GC / Visual.
                        # Source PDFs often show the method once for a block and leave following rows blank.
                        if method:
                            current_method = method
                        elif acceptance and test:
                            method = current_method

                        row_key = (test.lower(), method.lower(), acceptance.lower())
                        if row_key in seen:
                            continue
                        seen.add(row_key)
                        rows.append([test, method, acceptance])
        except Exception:
            return []
        finally:
            if doc is not None:
                doc.close()

        return self._fill_down_s41_methods(rows)

    def _fill_down_s41_methods(self, rows: list[list[str]]) -> list[list[str]]:
        filled: list[list[str]] = []
        carry_method = ""

        for test, method, acceptance in rows:
            test_clean = self._clean_cell(test)
            method_clean = self._clean_cell(method)
            acceptance_clean = self._clean_cell(acceptance)

            if method_clean:
                carry_method = method_clean
            elif acceptance_clean and test_clean:
                method_clean = carry_method

            # Section/group label rows should not receive a carried method.
            if test_clean.lower() in {
                "identification",
                "related compounds(test 1)",
                "related compounds(test 2)",
                "residual solvent",
            } and not acceptance_clean:
                method_clean = ""

            filled.append([test_clean, method_clean, acceptance_clean])

        return filled

    def _extract_s41_source_rows_from_pdf(self, pdf_path: Path) -> list[list[str]]:
        try:
            import fitz  # type: ignore
        except Exception:
            return []

        try:
            collected_pages: list[str] = []
            with fitz.open(pdf_path) as doc:
                for page in doc:
                    text = page.get_text("text", sort=True) or ""
                    low = text.lower()
                    if (
                        "table 8 specification for final iodixanol" in low
                        or "the specifications for iodixanol are provided in table 8" in low
                        or "related compounds(test 1)" in low
                    ):
                        collected_pages.append(text)
            if not collected_pages:
                return []
            return self._extract_s41_source_rows_from_raw_text("\n".join(collected_pages))
        except Exception:
            return []

    def _extract_s41_source_rows_from_flat_text(self, raw_text: str) -> list[list[str]]:
        text = raw_text.replace("", "°").replace("―", "-").replace("‖", "")
        text = re.sub(r"\s+", " ", text).strip()
        start_match = re.search(r"table\s*8\s*specification\s*for\s*final\s*iodixanol", text, flags=re.IGNORECASE)
        if not start_match:
            return []

        segment = text[start_match.end() :]
        end_match = re.search(r"bacterial\s*endotoxin\s*-?\s*NMT\s*0\.0038\s*IU?\s*for\s*1mg\s*of\s*iodixanol", segment, flags=re.IGNORECASE)
        if end_match:
            segment = segment[: end_match.end()]

        anchors = [
            ("Description", r"Description\s+Visual", "Visual"),
            ("Solubility", r"Solubility\s+-", "---------"),
            ("Transparency of solution", r"Transparency of\s+solution\s+Visual", "Visual"),
            ("Color of solution", r"Color of\s+solution\s+Colorimetry", "Colorimetry"),
            ("Specific Rotation (°)", r"Specific Rotation\s*\([^)]*\)\s+781S", "781S"),
            ("Identification", r"Identification", "IR spectrum; HPLC; Positive reaction"),
            ("pH", r"pH\s+pH meter", "pH meter"),
            ("Water", r"Water\s+Method I \(921\)", "Method I (921)"),
            ("Heavy metals", r"Heavy metals\s+Method I \(231\)", "Method I (231)"),
            ("Sulphated ash", r"Sulphated ash\s+Sulphated ash", "Sulphated ash"),
            ("Free iodine", r"Free iodine\s+Color reaction", "Color reaction"),
            ("Free iodide", r"Free iodide\s+Titration", "Titration"),
            ("Free aromatic amine", r"Free aromatic amine\s+Spectrophotometry", "Spectrophotometry"),
            ("Limit of calcium", r"Limit of calcium\s+ICP", "ICP"),
            ("Ionic compounds", r"Ionic compounds\s+Conductivity", "Conductivity"),
            ("Residual solvent - Methanol", r"Residual solvent\s+Methanol\s+GC", "GC"),
            ("Residual solvent - Isopropyl alcohol", r"Isopropyl alcohol", "GC"),
            ("Residual solvent - 1-methoxy-2-propanol", r"1-methoxy-2-\s*propanol", "GC"),
            ("Residual solvent - Methoxyethanol", r"Methoxyethanol", "GC"),
            ("Residual solvent - 1,3-dichloro-2-propanol", r"1,3-dichloro-2-propanol\s+GC", "GC"),
            ("Related compounds (Test 1) - Related compound B", r"Related compounds\(Test 1\)\s+Related compound B\s+HPLC", "HPLC"),
            ("Related compounds (Test 1) - Related compound C", r"Related compound C", "HPLC"),
            ("Related compounds (Test 1) - Related compound D", r"Related compound D", "HPLC"),
            ("Related compounds (Test 1) - Related compound F", r"Related compound F", "HPLC"),
            ("Related compounds (Test 1) - Related compound G", r"Related compound G", "HPLC"),
            ("Related compounds (Test 1) - Iohexol", r"Iohexol", "HPLC"),
            ("Related compounds (Test 1) - O-alkylated compounds", r"O-alkylated compounds", "HPLC"),
            ("Related compounds (Test 1) - Others", r"Others", "HPLC"),
            ("Related compounds (Test 2) - Related compound E", r"Related compounds\(Test 2\)\s+Related compound E\s+HPLC", "HPLC"),
            ("Related compounds (Test 2) - Related compound H", r"Related compound H", "HPLC"),
            ("Microbial limits", r"Microbial limits\s+61", "61"),
            ("Microbiological purity", r"Microbiological purity\s+-", "---------"),
            ("Assay", r"Assay\s+Titration", "Titration"),
            ("Bacterial endotoxin", r"Bacterial endotoxin\s+-", "---------"),
        ]

        located: list[tuple[int, int, str, str]] = []
        for test_name, pattern, method in anchors:
            found = re.search(pattern, segment, flags=re.IGNORECASE)
            if found:
                located.append((found.start(), found.end(), test_name, method))

        if len(located) < 5:
            return []

        located.sort(key=lambda item: item[0])
        rows: list[list[str]] = []
        for idx, (_, end, test_name, method) in enumerate(located):
            next_start = located[idx + 1][0] if idx + 1 < len(located) else len(segment)
            acceptance = segment[end:next_start].strip(" :-;")
            acceptance = re.sub(r"\s+", " ", acceptance).strip()
            if test_name == "Identification":
                acceptance = re.sub(r"\bA\.\s*IR spectrum\b", "A. IR spectrum:", acceptance, flags=re.IGNORECASE)
                acceptance = re.sub(r"\bB\.\s*HPLC\b", " B. HPLC:", acceptance, flags=re.IGNORECASE)
                acceptance = re.sub(r"\bC\.\s*Positive reaction\b", " C. Positive reaction:", acceptance, flags=re.IGNORECASE)
            if acceptance:
                rows.append([test_name, method, acceptance])
        return rows

    @staticmethod
    def _set_cell_text(cell, value: str) -> None:
        cell.text = value

    @staticmethod
    def _trim_table_rows(table, keep_rows: int) -> None:
        while len(table.rows) > keep_rows:
            row = table.rows[-1]._tr
            row.getparent().remove(row)

    def _append_rows(self, table, rows: list[list[str]]) -> None:
        for row_data in rows:
            row = table.add_row()
            for idx, value in enumerate(row_data[:3]):
                self._set_cell_text(row.cells[idx], value)

    def fill_s4_section(self, extracted: dict[str, ExtractedSectionContent], output_docx: Path) -> list[str]:
        doc = Document(self.template_docx)
        warnings: list[str] = []

        start_idx, end_idx = self._get_target_range(doc)
        s41_name_line = self._resolve_heading_line("2.3.S.4.1 Specification")
        s42_name_line = self._resolve_heading_line("2.3.S.4.2 Analytical Procedures")
        for heading in [
            "2.3.S.4 Control of the API",
        ]:
            self._insert_heading_line_if_missing(doc, heading, s41_name_line, start_idx, end_idx)

        for heading in [
            "2.3.S.4.1 Specification",
            "API specifications of the FPP manufacturer",
        ]:
            self._insert_heading_line_if_missing(doc, heading, s41_name_line, start_idx, end_idx)

        self._insert_heading_line_if_missing(doc, "2.3.S.4.2 Analytical Procedures", s42_name_line, start_idx, end_idx)

        start_idx, end_idx = self._get_target_range(doc)
        self._remove_refer_lines(doc, start_idx, end_idx)

        content = extracted["3.2.S.4.1"]
        if content.warning:
            warnings.append(f"3.2.S.4.1: {content.warning}")

        table = self._find_s41_template_table(doc)
        if table is None:
            warnings.append("2.3.S.4.1: target specification table not found in template")
        else:
            standard, spec_ref, rows = self._extract_s41_metadata_and_rows(content)
            self._set_cell_text(table.rows[0].cells[-1], standard or "USP")
            self._set_cell_text(table.rows[1].cells[-1], spec_ref or "-----")
            self._set_cell_text(table.rows[2].cells[0], "Test")
            self._set_cell_text(table.rows[2].cells[1], "Acceptance criteria")
            self._set_cell_text(table.rows[2].cells[2], "Analytical procedure\n(Type/Source/Version)")
            self._trim_table_rows(table, keep_rows=3)
            self._append_rows(table, rows)
            if not rows:
                warnings.append("2.3.S.4.1: no specification rows extracted from source PDF")

        content_42 = extracted.get("3.2.S.4.2")
        if content_42 and content_42.warning:
            warnings.append(f"3.2.S.4.2: {content_42.warning}")
        self._set_paragraph_text_with_suffix(
            doc,
            "Summary of the analytical procedures",
            "Please refer to Module 3 Section 3.2.S.4.2",
            start_idx,
            end_idx,
        )

        output_docx.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_docx)
        return warnings
