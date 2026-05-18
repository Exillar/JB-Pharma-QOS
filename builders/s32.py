"""builders/s32.py — S32DocxFiller: fills QOS section 2.3.S.3.2 Impurities."""
from __future__ import annotations

import re
from pathlib import Path

import fitz
from docx import Document
from docx.document import Document as _Document

from builders.base import _DocxHelper, run_artifact_cleanup
from config_loader import DiagramConfig
from pdf_extractor import ExtractedSectionContent
from ctd_utils import section_flexible_regex


class S32DocxFiller(_DocxHelper):
    SECTION_START = "2.3.S.3.2 Impurities"
    SECTION_END = "2.3.S.4 Control"

    def __init__(
        self,
        template_docx: Path,
        filled_reference_docx: Path | None = None,
        *,
        images_dir: Path | None = None,
        diagram_cfg: DiagramConfig | None = None,
        preserve_repeated_patterns: tuple[str, ...] = (),
    ) -> None:
        self.template_docx = template_docx
        self.filled_reference_docx = filled_reference_docx
        self.images_dir = images_dir
        self.diagram_cfg = diagram_cfg or DiagramConfig()
        self._preserve_repeated_patterns = preserve_repeated_patterns

    def _find_section_pages(
        self,
        doc: fitz.Document,
        refer_section: str,
        *,
        next_section: str | None,
    ) -> tuple[int, int]:
        start_re = re.compile(section_flexible_regex(refer_section), re.IGNORECASE)
        start_page = None
        for i in range(doc.page_count):
            if start_re.search(doc.load_page(i).get_text("text", sort=True)):
                start_page = i
                break
        if start_page is None:
            return (0, doc.page_count - 1)

        end_page = doc.page_count - 1
        if next_section:
            next_re = re.compile(section_flexible_regex(next_section), re.IGNORECASE)
            for i in range(start_page + 1, doc.page_count):
                if next_re.search(doc.load_page(i).get_text("text", sort=True)):
                    end_page = max(start_page, i - 1)
                    break
        return (start_page, end_page)

    @staticmethod
    def _norm_cell(s) -> str:
        if s is None:
            return ""
        return re.sub(r"\s+", " ", str(s).replace(" ", " ").strip())


    def _extract_api_impurity_table_images(
        self,
        pdf_path: Path,
        *,
        refer_section: str = "3.2.S.3.2",
        next_section: str = "3.2.S.3.3",
    ) -> tuple[list[Path], list[str]]:
        warnings: list[str] = []
        if self.images_dir is None:
            return ([], ["images_dir not configured"])
        self.images_dir.mkdir(parents=True, exist_ok=True)
        scale = float(self.diagram_cfg.render_dpi_scale or 2.0)

        def is_api_table(tab_rows: list[list[object]]) -> bool:
            window = tab_rows[:4]
            flat = " ".join(
                self._norm_cell(c) for r in window for c in (r[:3] if r else [])
            ).lower()
            return ("api-related impurity" in flat) and ("structure" in flat) and ("origin" in flat)

        images: list[Path] = []
        try:
            with fitz.open(pdf_path) as doc:
                sp, ep = self._find_section_pages(doc, refer_section, next_section=next_section)
                for pidx in range(sp, ep + 1):
                    page = doc.load_page(pidx)
                    try:
                        found = page.find_tables()
                    except Exception:
                        continue
                    for tab in found.tables:
                        if getattr(tab, "col_count", 0) != 3:
                            continue
                        try:
                            extracted = tab.extract() or []
                        except Exception:
                            extracted = []
                        if not extracted or not is_api_table(extracted):
                            continue

                        try:
                            rect = fitz.Rect(tab.bbox) & page.rect
                        except Exception:
                            rect = None
                        if rect is None or rect.is_empty:
                            continue
                        pad_x = max(6.0, rect.width * 0.02)
                        pad_y = max(6.0, rect.height * 0.02)
                        clip = fitz.Rect(
                            rect.x0 - pad_x, rect.y0 - pad_y,
                            rect.x1 + pad_x, rect.y1 + pad_y,
                        ) & page.rect
                        out = self.images_dir / f"3_2_S_3_2_api_table_p{pidx + 1}_{len(images) + 1}.png"
                        try:
                            pix = page.get_pixmap(
                                matrix=fitz.Matrix(scale, scale), clip=clip, alpha=False
                            )
                            pix.save(str(out))
                            images.append(out)
                        except Exception as img_err:
                            warnings.append(f"failed to render table image: {img_err}")
        except Exception as e:
            warnings.append(f"failed to read PDF {pdf_path.name}: {e}")

        return (images, warnings)

    def fill_s32_section(
        self,
        extracted: dict[str, ExtractedSectionContent],
        output_docx: Path,
    ) -> list[str]:
        doc = Document(self.template_docx)
        template_table_count = len(doc.tables)
        warnings: list[str] = []

        start_idx, end_idx = self._get_target_range(doc)

        name_line = self._resolve_name_manufacturer_line(
            self.filled_reference_docx, "2.3.S.3.2 Impurities"
        )
        idx = self._find_para_index_doc(doc, "2.3.S.3.2 Impurities", start_idx, end_idx)
        if idx is not None and name_line:
            nxt = (doc.paragraphs[idx + 1].text or "").strip() if idx + 1 < len(doc.paragraphs) else ""
            if not nxt.startswith("("):
                self._insert_paragraph_after(doc.paragraphs[idx], name_line)

        payload = extracted.get("3.2.S.3.2")
        if not payload:
            warnings.append("3.2.S.3.2: missing extracted payload")
            table_images: list[Path] = []
        else:
            table_images, ws = self._extract_api_impurity_table_images(payload.source_pdf)
            warnings.extend(f"3.2.S.3.2: {w}" for w in ws)

        start_idx, end_idx = self._get_target_range(doc)
        anchor_idx = self._find_para_index_doc(doc, "List of API-related impurities", start_idx, end_idx)
        if anchor_idx is None:
            warnings.append("3.2.S.3.2: API impurities anchor paragraph not found in template")
        else:
            placeholder = None
            try:
                from docx.oxml.table import CT_Tbl
                from docx.oxml.text.paragraph import CT_P
                from docx.table import Table
                from docx.text.paragraph import Paragraph as _P

                def _iter_blocks(d: _Document):
                    for child in d.element.body.iterchildren():
                        if isinstance(child, CT_P):
                            yield _P(child, d)
                        elif isinstance(child, CT_Tbl):
                            yield Table(child, d)

                blocks = list(_iter_blocks(doc))
                anchor_p = doc.paragraphs[anchor_idx]
                anchor_block_i = None
                for i, b in enumerate(blocks):
                    if hasattr(b, "_p") and getattr(b, "_p") is anchor_p._p:
                        anchor_block_i = i
                        break
                if anchor_block_i is not None:
                    for b in blocks[anchor_block_i + 1: anchor_block_i + 6]:
                        if isinstance(b, Table):
                            placeholder = b
                            break
            except Exception:
                placeholder = None

            if placeholder is None:
                warnings.append("3.2.S.3.2: API impurity placeholder table not found in template")
            else:
                if table_images:
                    anchor_p = doc.paragraphs[anchor_idx]
                    p = self._insert_paragraph_after(anchor_p, "")
                    current_p = p
                    try:
                        try:
                            from docx.enum.text import WD_BREAK
                        except Exception:
                            WD_BREAK = None

                        for idx, table_img in enumerate(table_images, start=1):
                            if not isinstance(table_img, Path) or not table_img.exists():
                                warnings.append(f"3.2.S.3.2: table image missing: {table_img}")
                                continue
                            if idx > 1 and WD_BREAK is not None:
                                run = current_p.add_run()
                                run.add_break(WD_BREAK.PAGE)
                                current_p = self._insert_paragraph_after(current_p, "")
                            run = current_p.add_run()
                            self._add_picture_autofit(run, table_img, doc)
                            if idx < len(table_images):
                                current_p = self._insert_paragraph_after(current_p, "")
                        parent = placeholder._element.getparent()
                        if parent is not None:
                            parent.remove(placeholder._element)
                    except Exception as e:
                        warnings.append(f"3.2.S.3.2: failed to insert table image: {e}")
                else:
                    warnings.append("3.2.S.3.2: table image not extracted from PDF")

        preserve_patterns = tuple(self._preserve_repeated_patterns)
        if name_line:
            preserve_patterns = preserve_patterns + (name_line,)
        cleanup_stats = run_artifact_cleanup(
            doc,
            keep_first_n_tables=template_table_count,
            preserve_repeated_patterns=preserve_patterns,
        )
        if any(cleanup_stats.values()):
            warnings.append(f"cleanup: {cleanup_stats}")

        output_docx.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_docx)
        return warnings
