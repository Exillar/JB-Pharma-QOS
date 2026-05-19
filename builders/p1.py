"""builders/p1.py — P1DocxFiller: fills QOS 2.3.P.1 Description and Composition of the FPP.

Strategy (same pattern as S5DocxFiller):
  1. Read the already-filled reference DOCX for every label value.
  2. Read the 3.2.P.1 PDF via find_tables() for the clean composition table.
  3. Fill the blank template label-by-label:
       • Remove "Refer Section 3.2.P.1" placeholder.
       • (a) Description of the FPP          → inline after colon.
       • (b) Composition / "Each ml contains" → text lines inserted after label.
       • Table 22                             → filled with PDF ingredient rows.
       • (b-ii) Composition of mixtures      → inline after colon.
       • (c) Reconstitution diluent           → inline after colon.
       • (d) Container closure system         → inline after colon.
"""
from __future__ import annotations

import re
from pathlib import Path

import fitz
from docx import Document
from docx.document import Document as _Document

from builders.base import _DocxHelper, run_artifact_cleanup
from pdf_extractor import ExtractedSectionContent


class P1DocxFiller(_DocxHelper):

    SECTION_START = "2.3.P.1 Description and Composition of the FPP"
    SECTION_END   = "2.3.P.2 Pharmaceutical Development"

    def __init__(
        self,
        template_docx: Path,
        filled_reference_docx: Path | None = None,
        *,
        preserve_repeated_patterns: tuple[str, ...] = (),
    ) -> None:
        self.template_docx = template_docx
        self.filled_reference_docx = filled_reference_docx
        self._preserve_repeated_patterns = preserve_repeated_patterns

    # ------------------------------------------------------------------
    # Step 1 — read values from the already-filled reference DOCX
    # ------------------------------------------------------------------

    def _read_reference_p1(self) -> dict[str, str]:
        """Return a dict with keys: description, mixtures, diluent, container."""
        result: dict[str, str] = {}
        if not self.filled_reference_docx or not self.filled_reference_docx.exists():
            return result

        try:
            ref = Document(self.filled_reference_docx)
        except Exception:
            return result

        paras = ref.paragraphs
        start = end = None
        for i, p in enumerate(paras):
            t = (p.text or "").strip()
            if start is None and "2.3.P.1" in t:
                start = i
            elif start is not None and "2.3.P.2" in t:
                end = i
                break
        if start is None:
            return result
        end = end or len(paras)

        def _after_colon(needle: str) -> str:
            lo = needle.lower()
            for i in range(start, end):
                t = (paras[i].text or "").strip()
                if lo not in t.lower():
                    continue
                # Value after the colon on the same paragraph
                if ":" in t:
                    val = t.split(":", 1)[1].strip()
                    if val:
                        # Truncate bleed-in from next section (e.g. "Accelerated Stability...")
                        val = re.split(
                            r"[\.\s]Accelerated\s+Stability\b",
                            val, maxsplit=1, flags=re.IGNORECASE
                        )[0].strip()
                        return val
                # Value may be on the next non-empty paragraph
                for j in range(i + 1, min(i + 6, end)):
                    nxt = (paras[j].text or "").strip()
                    if nxt and not any(k in nxt.lower() for k in [
                        "composition", "description of", "type of container", "refer section"
                    ]):
                        return nxt
            return ""

        # Description: may span two paragraphs → join them
        desc_val = _after_colon("Description of the FPP (in signed specifications)")
        if desc_val:
            # Check if the next paragraph is a sentence continuation (no leading label word)
            for i in range(start, end):
                t = (paras[i].text or "").strip()
                if "Description of the FPP (in signed specifications)" in t and ":" in t:
                    tail = t.split(":", 1)[1].strip()
                    # Collect continuation paragraphs
                    parts = [tail] if tail else []
                    for j in range(i + 1, min(i + 5, end)):
                        nxt = (paras[j].text or "").strip()
                        if not nxt:
                            continue
                        if any(k in nxt.lower() for k in ["composition", "description of", "type of"]):
                            break
                        parts.append(nxt)
                    desc_val = " ".join(parts).strip()
                    break
            result["description"] = desc_val

        mixtures = _after_colon("Composition of all components purchased as mixtures")
        if mixtures:
            result["mixtures"] = mixtures

        diluent = _after_colon("Description of accompanying reconstitution diluent")
        if diluent:
            result["diluent"] = diluent

        container = _after_colon("Type of container closure system")
        if container:
            result["container"] = container

        return result

    # ------------------------------------------------------------------
    # Step 2 — read clean composition table from PDF via find_tables()
    # ------------------------------------------------------------------

    def _read_pdf_composition_table(
        self, pdf_path: Path
    ) -> list[tuple[str, str, str]]:
        """Return list of (component_with_spec, function, amount) from the PDF table."""
        rows: list[tuple[str, str, str]] = []
        try:
            with fitz.open(pdf_path) as doc:
                for pidx in range(doc.page_count):
                    page = doc.load_page(pidx)
                    try:
                        tables = page.find_tables().tables
                    except Exception:
                        continue
                    for tab in tables:
                        try:
                            extracted = tab.extract() or []
                        except Exception:
                            continue
                        if len(extracted) < 2:
                            continue
                        header = " ".join(str(c or "") for c in extracted[0]).lower()
                        if "ingredient" not in header and "component" not in header:
                            continue
                        for row in extracted[1:]:
                            if len(row) < 4:
                                continue
                            ingredient = re.sub(r"\s+", " ", str(row[1] or "")).strip()
                            spec       = re.sub(r"\s+", " ", str(row[2] or "")).strip()
                            function   = re.sub(r"\s+", " ", str(row[3] or "")).strip()
                            amount     = re.sub(r"\s+", " ", str(row[4] or "")).strip() if len(row) > 4 else ""
                            if not ingredient and not amount:
                                continue
                            component = (
                                f"{ingredient} {spec}".strip()
                                if spec and spec not in ingredient
                                else ingredient
                            )
                            rows.append((component, function, amount))
                        if rows:
                            return rows
        except Exception:
            pass
        return rows

    def _read_pdf_strength(self, pdf_path: Path) -> str:
        try:
            with fitz.open(pdf_path) as doc:
                text = doc.load_page(0).get_text("text", sort=True)
            m = re.search(r"(\d+)\s*mg", text)
            return f"{m.group(1)}mg" if m else ""
        except Exception:
            return ""

    # ------------------------------------------------------------------
    # Step 3 — fill the composition table (Table 22 in blank template)
    # ------------------------------------------------------------------

    def _find_composition_table(self, doc: _Document):
        """Find the template composition table (right after 'Composition, i.e.' para)."""
        NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        prev_lo = ""
        for elem in doc.element.body:
            tag = elem.tag.split("}")[-1]
            if tag == "p":
                prev_lo = "".join(
                    t.text or "" for t in elem.iter(f"{{{NS}}}t")
                ).strip().lower()
            elif tag == "tbl":
                if "composition, i.e., list of all components" in prev_lo:
                    for tbl in doc.tables:
                        if tbl._tbl is elem:
                            return tbl
        # Fallback by header text
        for tbl in doc.tables:
            if tbl.rows and "component and quality standard" in " ".join(
                c.text.strip() for c in tbl.rows[0].cells
            ).lower():
                return tbl
        return None

    def _fill_composition_table(
        self,
        doc: _Document,
        ingredients: list[tuple[str, str, str]],
        strength: str,
        warnings: list[str],
    ) -> None:
        tbl = self._find_composition_table(doc)
        if tbl is None:
            warnings.append("2.3.P.1: composition table not found in template")
            return
        if not ingredients:
            warnings.append("2.3.P.1: no ingredient data for composition table")
            return

        # Update "Strength (label claim)" header (row 0)
        if strength and tbl.rows:
            seen = set()
            for cell in tbl.rows[0].cells:
                cid = id(cell._tc)
                if cid in seen:
                    continue
                seen.add(cid)
                if "strength (label claim)" in cell.text.lower():
                    cell.text = f"Strength (label claim) : {strength}"
                    break

        # Remove all data/placeholder rows (keep only the 3 header rows)
        while len(tbl.rows) > 3:
            tbl._tbl.remove(tbl.rows[-1]._tr)

        # Add one row per ingredient
        for component, function, amount in ingredients:
            row = tbl.add_row()
            cells = row.cells
            if cells:
                cells[0].text = component
            if len(cells) >= 2:
                cells[1].text = function
            if len(cells) >= 3:
                cells[2].text = amount

    # ------------------------------------------------------------------
    # Public entry point
    # ------------------------------------------------------------------

    def fill_p1_section(
        self,
        extracted: dict[str, ExtractedSectionContent],
        output_docx: Path,
    ) -> list[str]:
        doc = Document(self.template_docx)
        template_table_count = len(doc.tables)
        warnings: list[str] = []

        start_idx, end_idx = self._get_target_range(doc)

        # --- Read filled reference (primary data source) ---
        ref = self._read_reference_p1()

        # --- Read PDF composition table (clean data, not available from ref) ---
        payload = extracted.get("3.2.P.1")
        ingredients: list[tuple[str, str, str]] = []
        strength = ""
        if payload and Path(payload.source_pdf).exists():
            ingredients = self._read_pdf_composition_table(payload.source_pdf)
            strength    = self._read_pdf_strength(payload.source_pdf)
            if payload.warning:
                warnings.append(f"3.2.P.1: {payload.warning}")
        else:
            warnings.append("3.2.P.1: source PDF not found")

        # --- Name / manufacturer line under section heading ---
        name_line = self._resolve_name_manufacturer_line(
            self.filled_reference_docx, self.SECTION_START
        )
        if name_line:
            h_idx = self._find_para_index_doc(doc, self.SECTION_START, start_idx, end_idx)
            if h_idx is not None:
                nxt = (doc.paragraphs[h_idx + 1].text or "").strip() if h_idx + 1 < len(doc.paragraphs) else ""
                if not nxt.startswith("("):
                    self._insert_paragraph_after(doc.paragraphs[h_idx], name_line)
            start_idx, end_idx = self._get_target_range(doc)

        # --- Remove "Refer Section 3.2.P.1" placeholder ---
        self._remove_paragraphs_matching(
            doc, ["Refer Section 3.2.P.1", "Refer section 3.2.P.1"],
            start_idx, end_idx,
        )
        start_idx, end_idx = self._get_target_range(doc)

        # --- (a) Description of the FPP ---
        description = ref.get("description", "")
        if description:
            idx = self._find_para_index_doc(
                doc, "Description of the FPP (in signed specifications)", start_idx, end_idx
            )
            if idx is not None:
                self._append_inline_value_after_colon(doc.paragraphs[idx], description)
            else:
                warnings.append("2.3.P.1: 'Description of the FPP' label not found")

        # --- (b-i) Composition, i.e., list: insert "Each ml contains:" block ---
        # The composition ingredient text is garbled in both PDF and reference (font artefacts).
        # Build clean lines from the PDF table data instead.
        c_idx = self._find_para_index_doc(
            doc, "Composition, i.e., list of all components", start_idx, end_idx
        )
        if c_idx is not None and ingredients:
            cursor = doc.paragraphs[c_idx]
            cursor = self._insert_paragraph_after(cursor, "Each ml contains:")
            for component, _func, amount in ingredients:
                cursor = self._insert_paragraph_after(cursor, f"{component}  {amount}")
            start_idx, end_idx = self._get_target_range(doc)
        elif c_idx is None:
            warnings.append("2.3.P.1: 'Composition, i.e., list of all components' label not found")

        # --- Fill Table 22 (composition table) ---
        self._fill_composition_table(doc, ingredients, strength, warnings)

        # --- (b-ii) Composition of all components purchased as mixtures ---
        mixtures = ref.get("mixtures", "Not Applicable")
        m_idx = self._find_para_index_doc(
            doc, "Composition of all components purchased as mixtures", start_idx, end_idx
        )
        if m_idx is not None:
            self._append_inline_value_after_colon(doc.paragraphs[m_idx], mixtures)
        else:
            warnings.append("2.3.P.1: 'Composition of all components purchased as mixtures' label not found")

        # --- (c) Description of accompanying reconstitution diluent ---
        diluent = ref.get("diluent", "Not Applicable")
        d_idx = self._find_para_index_doc(
            doc, "Description of accompanying reconstitution diluent", start_idx, end_idx
        )
        if d_idx is not None:
            self._append_inline_value_after_colon(doc.paragraphs[d_idx], diluent)
        else:
            warnings.append("2.3.P.1: 'Description of accompanying reconstitution diluent' label not found")

        # --- (d) Type of container closure system ---
        container = ref.get("container", "")
        if container:
            ct_idx = self._find_para_index_doc(
                doc, "Type of container closure system", start_idx, end_idx
            )
            if ct_idx is not None:
                self._append_inline_value_after_colon(doc.paragraphs[ct_idx], container)
            else:
                warnings.append("2.3.P.1: 'Type of container closure system' label not found")

        # --- Cleanup ---
        preserve = tuple(self._preserve_repeated_patterns)
        if name_line:
            preserve = preserve + (name_line,)
        stats = run_artifact_cleanup(
            doc,
            keep_first_n_tables=template_table_count,
            preserve_repeated_patterns=preserve,
        )
        if any(stats.values()):
            warnings.append(f"cleanup: {stats}")

        output_docx.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_docx)
        return warnings
