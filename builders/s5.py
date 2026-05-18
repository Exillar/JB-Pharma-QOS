"""builders/s5.py — S5DocxFiller: fills QOS section 2.3.S.5 Reference Standards or Materials."""
from __future__ import annotations

import re
from pathlib import Path

import fitz
from docx import Document
from docx.document import Document as _Document
from builders.base import _DocxHelper, run_artifact_cleanup
from ctd_utils import section_flexible_regex
from config_loader import DiagramConfig
from pdf_extractor import ExtractedSectionContent

class S5DocxFiller(_DocxHelper):
    SECTION_START = "2.3.S.5 Reference Standards or Materials"
    SECTION_END = "2.3.S.6 Container Closure System"
    
    def __init__(
        self,
        template_docx: Path,
        filled_reference_docx: Path | None = None,
        *,
        images_dir: Path | None = None,
        diagram_cfg: DiagramConfig | None = None,
        preserve_repeated_patterns: tuple[str, ...] = (),
    ) -> None:
        self.template_docx = template_docx
        self.filled_reference_docx = filled_reference_docx
        self.images_dir = images_dir
        self.diagram_cfg = diagram_cfg or DiagramConfig()
        self._preserve_repeated_patterns = preserve_repeated_patterns

    def _find_section_page(self, doc: fitz.Document, section_id: str) -> int | None:
        pat = re.compile(section_flexible_regex(section_id), re.IGNORECASE)
        if section_id == "3.2.S.5":
            heading_pat = re.compile(
                r"\b3\s*\.\s*2\s*\.\s*S\s*\.\s*5\s+Reference\s+Standards\s+or\s+Materials\b",
                re.IGNORECASE,
            )
        else:
            heading_pat = None
        for i in range(doc.page_count):
            text = doc.load_page(i).get_text("text", sort=True)
            if heading_pat and heading_pat.search(text):
                return i
            if pat.search(text):
                return i
        return None

    def _extract_s5_narrative(self, pdf_path: Path) -> str:
        try:
            with fitz.open(pdf_path) as doc:
                page_idx = self._find_section_page(doc, "3.2.S.5")
                if page_idx is None:
                    return ""
                text = doc.load_page(page_idx).get_text("text", sort=True)
        except Exception:
            return ""

        header_pat = re.compile(
            r"^\s*(drug\s+master\s+file|product\s+name|module\s*:|version\s*:|date\s*:|open\s+part)\b",
            re.IGNORECASE,
        )
        footer_pat = re.compile(r"^.{5,75}\s+\d{1,4}$")
        page_pat = re.compile(r"^\s*\d+\s+of\s+\d+\s*$", re.IGNORECASE)
        section_pat = re.compile(section_flexible_regex("3.2.S.5"), re.IGNORECASE)

        out_lines: list[str] = []
        for raw in text.splitlines():
            line = re.sub(r"\s+", " ", raw).strip()
            if not line:
                continue
            if section_pat.search(line):
                continue
            if header_pat.match(line) or page_pat.match(line):
                continue
            if footer_pat.match(line):
                continue
            out_lines.append(line)

        paragraphs: list[str] = []
        current: list[str] = []
        for line in out_lines:
            current.append(line)
            if re.search(r"[\.;:]\s*$", line):
                paragraphs.append(" ".join(current).strip())
                current = []
        if current:
            paragraphs.append(" ".join(current).strip())

        return "\n".join(p for p in paragraphs if p).strip()

    def _extract_s5_1_narrative(self, pdf_path: Path) -> str:
        try:
            with fitz.open(pdf_path) as doc:
                page_idx = self._find_section_page(doc, "3.2.S.5.1")
                if page_idx is None:
                    return ""
                lines: list[str] = []
                for i in range(page_idx, min(page_idx + 3, doc.page_count)):
                    text = doc.load_page(i).get_text("text", sort=True)
                    lines.extend(text.splitlines())
        except Exception:
            return ""

        header_pat = re.compile(
            r"^\s*(drug\s+master\s+file|product\s+name|module\s*:|version\s*:|date\s*:|open\s+part)\b",
            re.IGNORECASE,
        )
        footer_pat = re.compile(r"^.{5,75}\s+\d{1,4}$")
        page_pat = re.compile(r"^\s*\d+\s+of\s+\d+\s*$", re.IGNORECASE)
        section_pat = re.compile(section_flexible_regex("3.2.S.5.1"), re.IGNORECASE)

        out_lines: list[str] = []
        for raw in lines:
            line = re.sub(r"\s+", " ", raw).strip()
            if not line:
                continue
            if section_pat.search(line):
                continue
            if header_pat.match(line) or page_pat.match(line):
                continue
            if footer_pat.match(line):
                continue
            out_lines.append(line)

        paragraphs: list[str] = []
        current: list[str] = []
        for line in out_lines:
            current.append(line)
            if re.search(r"[\.;:]\s*$", line):
                paragraphs.append(" ".join(current).strip())
                current = []
        if current:
            paragraphs.append(" ".join(current).strip())

        return "\n".join(p for p in paragraphs if p).strip()

    def _extract_abc_from_reference(self) -> dict[str, list[str]]:
        if not self.filled_reference_docx or not self.filled_reference_docx.exists():
            return {}
        try:
            from docx import Document
        except Exception:
            return {}

        doc = Document(self.filled_reference_docx)
        paras = doc.paragraphs

        def norm(s: str) -> str:
            return re.sub(r"\s+", " ", s or "").strip().lower()

        start_idx = end_idx = None
        for i, p in enumerate(paras):
            t = norm(p.text)
            if start_idx is None and "2.3.s.5 reference standards" in t:
                start_idx = i
            elif start_idx is not None and "2.3.s.6 container closure" in t:
                end_idx = i
                break
        if start_idx is None:
            return {}
        if end_idx is None:
            end_idx = len(paras)

        a_idx = b_idx = c_idx = None
        for i in range(start_idx, end_idx):
            t = norm(paras[i].text)
            if a_idx is None and "source (including lot number)" in t:
                a_idx = i
            elif b_idx is None and "characterization and evaluation of non-official" in t:
                b_idx = i
            elif c_idx is None and "description of the process controls of the secondary reference standard" in t:
                c_idx = i

        def collect(start: int | None, end: int | None) -> list[str]:
            if start is None:
                return []
            stop = end if end is not None else end_idx
            lines: list[str] = []
            for i in range(start + 1, min(stop, end_idx)):
                text = (paras[i].text or "").strip()
                if text:
                    lines.append(text)
            return lines

        return {
            "a": collect(a_idx, b_idx),
            "b": collect(b_idx, c_idx),
            "c": collect(c_idx, None),
        }

    def _extract_s5_table_images(self, pdf_path: Path) -> tuple[list[Path], list[str]]:
        warnings: list[str] = []
        if self.images_dir is None:
            return ([], ["images_dir not configured"])
        self.images_dir.mkdir(parents=True, exist_ok=True)
        scale = float(self.diagram_cfg.render_dpi_scale or 2.0)

        def is_working_standards_table(rows: list[list[object]]) -> bool:
            if not rows:
                return False
            head = " ".join(str(c or "") for c in rows[0]).lower()
            return (
                "chemical structure" in head
                and "batch" in head
                and ("sr" in head or "no." in head or "no" in head)
            )

        def is_continuation_table(rows: list[list[object]]) -> bool:
            if not rows:
                return False
            col_count = max((len(r) for r in rows if r), default=0)
            if col_count < 3:
                return False
            flat = " ".join(str(c or "") for r in rows for c in r).lower()
            if "drug mater file" in flat or "drug master file" in flat:
                return False
            if "module" in flat and "open part" in flat:
                return False
            if "chemical structure" in flat or "batch no" in flat:
                return True
            if re.search(r"\bws[- ]?\d+\b", flat):
                return True
            return any(k in flat for k in ("iodixanol", "iohexol", "related compound", "compound"))

        images: list[Path] = []
        seen_header = False
        try:
            with fitz.open(pdf_path) as doc:
                for pidx in range(doc.page_count):
                    page = doc.load_page(pidx)
                    try:
                        tables = page.find_tables().tables
                    except Exception:
                        tables = []
                    for tab in tables:
                        try:
                            extracted = tab.extract() or []
                        except Exception:
                            extracted = []
                        header_hit = is_working_standards_table(extracted)
                        if header_hit:
                            seen_header = True
                        elif not seen_header:
                            continue
                        elif not is_continuation_table(extracted):
                            continue
                        try:
                            rect = fitz.Rect(tab.bbox) & page.rect
                        except Exception:
                            rect = None
                        if rect is None or rect.is_empty:
                            continue
                        pad_x = max(6.0, rect.width * 0.02)
                        pad_y = max(6.0, rect.height * 0.02)
                        clip = fitz.Rect(
                            rect.x0 - pad_x,
                            rect.y0 - pad_y,
                            rect.x1 + pad_x,
                            rect.y1 + pad_y,
                        ) & page.rect
                        out = self.images_dir / f"3_2_S_5_working_standards_p{pidx + 1}_{len(images) + 1}.png"
                        try:
                            pix = page.get_pixmap(
                                matrix=fitz.Matrix(scale, scale), clip=clip, alpha=False
                            )
                            pix.save(str(out))
                            images.append(out)
                        except Exception as img_err:
                            warnings.append(f"failed to render S5 table image: {img_err}")
        except Exception as e:
            warnings.append(f"failed to read PDF {pdf_path.name}: {e}")

        return (images, warnings)

    def fill_s5_section(
        self,
        extracted: dict[str, ExtractedSectionContent],
        output_docx: Path,
    ) -> list[str]:
        doc = Document(self.template_docx)
        template_table_count = len(doc.tables)
        warnings: list[str] = []

        start_idx, end_idx = self._get_target_range(doc)

        name_line = self._resolve_name_manufacturer_line(
            self.filled_reference_docx, self.SECTION_START
        )
        idx = self._find_para_index_doc(doc, self.SECTION_START, start_idx, end_idx)
        if idx is not None and name_line:
            nxt = (doc.paragraphs[idx + 1].text or "").strip() if idx + 1 < len(doc.paragraphs) else ""
            if not nxt.startswith("("):
                self._insert_paragraph_after(doc.paragraphs[idx], name_line)

        self._remove_paragraphs_matching(
            doc,
            ["Refer Section 3.2.S.5"],
            start_idx,
            end_idx,
        )

        payload = extracted.get("3.2.S.5")
        if not payload:
            warnings.append("3.2.S.5: missing extracted payload")
            narrative = ""
            table_images: list[Path] = []
            s51_narrative = ""
        else:
            ref_abc = self._extract_abc_from_reference()
            if ref_abc:
                narrative = "\n".join(ref_abc.get("a", []))
                s51_narrative = "\n".join(ref_abc.get("b", []))
                s52_narrative = "\n".join(ref_abc.get("c", []))
            else:
                narrative = self._extract_s5_narrative(payload.source_pdf)
                s51_narrative = self._extract_s5_1_narrative(payload.source_pdf)
                s52_narrative = s51_narrative
            table_images, ws = self._extract_s5_table_images(payload.source_pdf)
            warnings.extend(f"3.2.S.5: {w}" for w in ws)

        start_idx, end_idx = self._get_target_range(doc)
        a_idx = self._find_para_index_doc(
            doc,
            "Source (including lot number) of primary reference standards",
            start_idx,
            end_idx,
        )
        if a_idx is None:
            warnings.append("2.3.S.5(a): target prompt not found in template")
        else:
            cursor = doc.paragraphs[a_idx]
            if narrative:
                for line in narrative.splitlines():
                    if line.strip():
                        cursor = self._insert_paragraph_after(cursor, line)
            if table_images:
                for idx, img in enumerate(table_images, start=1):
                    if not isinstance(img, Path) or not img.exists():
                        warnings.append(f"3.2.S.5: table image missing: {img}")
                        continue
                    cursor = self._insert_paragraph_after(cursor, "")
                    self._add_picture_autofit(cursor.add_run(), img, doc)

        b_idx = self._find_para_index_doc(
            doc,
            "Characterization and evaluation of non-official",
            start_idx,
            end_idx,
        )
        if b_idx is not None and s51_narrative:
            cursor = doc.paragraphs[b_idx]
            for line in s51_narrative.splitlines():
                if line.strip():
                    cursor = self._insert_paragraph_after(cursor, line)

        c_idx = self._find_para_index_doc(
            doc,
            "Description of the process controls of the secondary reference standard",
            start_idx,
            end_idx,
        )
        if c_idx is not None and s52_narrative:
            cursor = doc.paragraphs[c_idx]
            for line in s52_narrative.splitlines():
                if line.strip():
                    cursor = self._insert_paragraph_after(cursor, line)

        preserve_patterns = tuple(self._preserve_repeated_patterns)
        if name_line:
            preserve_patterns = preserve_patterns + (name_line,)
        cleanup_stats = run_artifact_cleanup(
            doc,
            keep_first_n_tables=template_table_count,
            preserve_repeated_patterns=preserve_patterns,
        )
        if any(cleanup_stats.values()):
            warnings.append(f"cleanup: {cleanup_stats}")

        output_docx.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_docx)
        return warnings
