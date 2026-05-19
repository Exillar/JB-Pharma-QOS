"""builders/s1.py — S1DocxFiller: fills QOS section 2.3.S.1 General Information."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.document import Document as _Document

from builders.base import _DocxHelper, run_artifact_cleanup
from config_loader import S2FillConfig
from pdf_extractor import ExtractedSectionContent


class S1DocxFiller(_DocxHelper):
    SECTION_START = "2.3.S.1 General Information"
    SECTION_END = "2.3.S.2 Manufacture"

    def __init__(
        self,
        template_docx: Path,
        filled_reference_docx: Path | None = None,
    ) -> None:
        self.template_docx = template_docx
        self.filled_reference_docx = filled_reference_docx

    # ------------------------------------------------------------------
    # Parsers
    # ------------------------------------------------------------------

    def _parse_s11(self, raw_text: str) -> dict[str, str]:
        inn_block = self._extract_block(
            raw_text,
            r"Recommended\s+International\s+Nonproprietary\s+Name\s*\(INN\)\s*:",
            [r"Compendial\s+name\s*:"],
        )
        comp_block = self._extract_block(
            raw_text,
            r"Compendial\s+name\s*:",
            [r"Chemical\s+name\s*\(s\)\s*:"],
        )
        chem_block = self._extract_block(
            raw_text,
            r"Chemical\s+name\s*\(s\)\s*:",
            [r"Chemical\s+Abstracts\s+Service\s*\(CAS\)\s+registry\s+number\s*:"],
        )
        cas_block = self._extract_block(
            raw_text,
            r"Chemical\s+Abstracts\s+Service\s*\(CAS\)\s+registry\s+number\s*:",
            [r"Other\s+non-proprietary\s+name"],
        )
        return {
            "a": self._first_meaningful_line(inn_block),
            "b": self._first_meaningful_line(comp_block),
            "c": chem_block.strip(),
            "d": "",
            "e": "",
            "f": self._first_meaningful_line(cas_block),
        }

    def _parse_s12(self, raw_text: str) -> dict[str, str]:
        mf = re.search(r"Molecular\s+Formula\s*:\s*(.+)", raw_text, flags=re.IGNORECASE)
        mw = re.search(r"Molecular\s+weight\s*:\s*(.+)", raw_text, flags=re.IGNORECASE)
        return {
            "b": self._clean_line(mf.group(1)) if mf else "",
            "c": self._clean_line(mw.group(1)) if mw else "",
        }

    def _parse_s13(self, raw_text: str) -> dict[str, str]:
        def _block_any(starts: list[str], ends: list[str]) -> str:
            for s in starts:
                block = self._extract_block(raw_text, s, ends)
                if block:
                    return block
            return ""

        melt = _block_any(
            [r"\bMelting\s+point\b"],
            [r"\bpH\s*[-:]\b", r"\bPartition\s+coefficients\b"],
        )
        ph = _block_any(
            [r"\bpH\s*[-:]"],
            [r"\bPartition\s+coefficients\b", r"\bPK\s*[-:]"]
        )
        part = _block_any(
            [r"\bPartition\s+coefficients\b\s*[-:]"],
            [r"\bPK\s*[-:]", r"\bSpecific\s+Rotation\b"],
        )
        pk = _block_any(
            [r"\bPK\s*[-:]"],
            [r"\bSpecific\s+Rotation\b", r"\bPolymorphic\s+Form\b"],
        )
        rot = _block_any(
            [r"\bSpecific\s+Rotation\b"],
            [r"\bPolymorphic\s+Form\b", r"\bSolub"],
        )
        desc = _block_any(
            [
                r"\bPhysical\s+description\b\s*[-:]?",
                r"\bDescription\b\s*[-:]?",
            ],
            [r"\bSolub", r"\bPolymorphic", r"\bMelting\s+point\b"],
        )
        sol = _block_any(
            [r"\bSolub(?:ility|ilities)\b\s*[-:]?"],
            [r"\bPolymorphic", r"\bpH\b", r"\bMelting\s+point\b"],
        )
        poly = _block_any(
            [r"\bPolymorphic\s+[Ff]orm\b\s*[-:]?", r"\bPolymeric\s+[Ff]orm\b\s*[-:]?"],
            [r"\bSolvate", r"\bHydrate", r"\bOther\b"],
        )
        solvate = _block_any(
            [r"\bSolvate\b\s*[-:]?"],
            [r"\bHydrate\b", r"\bOther\b"],
        )
        hydrate = _block_any(
            [r"\bHydrate\b\s*[-:]?"],
            [r"\bOther\b"],
        )
        other = _block_any(
            [r"\bOther\b\s*[:]?"],
            [r"\b3\.2\.S\.", r"\b2\.3\.S\."]
        )
        return {
            "a": self._first_meaningful_line(desc) or desc.strip(),
            "b": self._first_meaningful_line(sol) or sol.strip(),
            "poly": self._first_meaningful_line(poly) or poly.strip(),
            "solvate": self._first_meaningful_line(solvate) or solvate.strip(),
            "hydrate": self._first_meaningful_line(hydrate) or hydrate.strip(),
            "other": other.strip(),
            "ph": self._first_meaningful_line(ph) or ph.strip(),
            "pk": "\n".join(ln.strip() for ln in pk.splitlines() if ln.strip()),
            "partition": "\n".join(ln.strip() for ln in part.splitlines() if ln.strip()),
            "melting": self._first_meaningful_line(melt) or melt.strip(),
            "rotation": self._first_meaningful_line(rot) or rot.strip(),
            "refractive": "",
            "hygro": "",
            "uv": "",
        }

    # ------------------------------------------------------------------
    # Table filler
    # ------------------------------------------------------------------

    def _fill_property_table(self, doc: _Document, s13: dict[str, str]) -> None:
        table_row_map = {
            "ph": s13["ph"],
            "pk": s13["pk"],
            "pka": s13["pk"],
            "partition coefficients": s13["partition"],
            "melting/boiling points": s13["melting"],
            "specific optical rotation (specify solvent)": s13["rotation"],
            "refractive index (liquids)": s13["refractive"],
            "hygroscopicity": s13["hygro"],
            "uv absorption maxima/molar absorptivity": s13["uv"],
            "other": s13["other"],
        }
        for table in doc.tables:
            if not table.rows:
                continue
            if self._norm(table.cell(0, 0).text) != "property":
                continue
            for row in table.rows[1:]:
                key = self._norm(row.cells[0].text)
                if key in table_row_map and len(row.cells) > 1:
                    row.cells[1].text = table_row_map[key]

    # ------------------------------------------------------------------
    # Visual formatting
    # ------------------------------------------------------------------

    def _apply_visual_formatting(self, doc: _Document, start_idx: int, end_idx: int) -> None:
        bold_starts = {
            "2.3.S.1 General Information",
            "2.3.S.1.2 Structure",
            "2.3.S.1.3 General Properties",
        }
        label_bold_prefixes = [
            "(Recommended) International Non-proprietary name (INN):",
            "Compendial name",
            "Chemical name",
        ]
        for i in range(start_idx, min(end_idx, len(doc.paragraphs))):
            p = doc.paragraphs[i]
            t = (p.text or "").strip()
            if any(t.startswith(s) for s in bold_starts):
                self._set_runs_style(p, bold=True, italic=False)
            elif t.startswith("2.3.S.1.1 Nomenclature"):
                self._set_runs_style(p, bold=False, italic=False)
            if any(t.startswith(lb) for lb in label_bold_prefixes):
                self._set_runs_style(p, bold=True, italic=False)

    # ------------------------------------------------------------------
    # Public entry point
    # ------------------------------------------------------------------

    def fill_s1_section(
        self,
        extracted: dict[str, ExtractedSectionContent],
        output_docx: Path,
    ) -> list[str]:
        doc = Document(self.template_docx)
        template_table_count = len(doc.tables)
        warnings: list[str] = []

        start_idx, end_idx = self._get_target_range(doc)

        name_mfr_line = self._resolve_name_manufacturer_line(
            self.filled_reference_docx, "2.3.S.1.1 Nomenclature"
        )
        for heading in [
            "2.3.S.1 General Information",
            "2.3.S.1.1 Nomenclature",
            "2.3.S.1.2 Structure",
            "2.3.S.1.3 General Properties",
        ]:
            idx = self._find_para_index_doc(doc, heading, start_idx, end_idx)
            if idx is not None and name_mfr_line:
                nxt = (doc.paragraphs[idx + 1].text or "").strip() if idx + 1 < len(doc.paragraphs) else ""
                if not nxt.startswith("("):
                    self._insert_paragraph_after(doc.paragraphs[idx], name_mfr_line)

        for refer in ["3.2.S.1.1", "3.2.S.1.2", "3.2.S.1.3"]:
            if extracted[refer].warning:
                warnings.append(f"{refer}: {extracted[refer].warning}")

        s11 = self._parse_s11(extracted["3.2.S.1.1"].raw_text)
        s12 = self._parse_s12(extracted["3.2.S.1.2"].raw_text)
        s13 = self._parse_s13(extracted["3.2.S.1.3"].raw_text)

        s13_ref_idx = self._find_para_index_doc(doc, "Refer Section 3.2.S.1.3", start_idx, end_idx)
        if s13_ref_idx is not None and s11["b"]:
            nxt = (doc.paragraphs[s13_ref_idx + 1].text or "").strip() if s13_ref_idx + 1 < len(doc.paragraphs) else ""
            if nxt.lower() != s11["b"].strip().lower():
                self._insert_paragraph_after(doc.paragraphs[s13_ref_idx], s11["b"])

        start_idx, end_idx = self._get_target_range(doc)

        field_map = [
            ("(Recommended) International Non-proprietary name (INN):", s11["a"]),
            ("Compendial name, if relevant:", s11["b"]),
            ("Chemical name(s):", s11["c"]),
            ("Company or laboratory code:", s11["d"]),
            ("Other non-proprietary name(s)", s11["e"]),
            ("Chemical Abstracts Service (CAS) registry number:", s11["f"]),
            ("Structural formula, including relative and absolute stereochemistry:", "__IMG__"),
            ("Molecular formula:", s12["b"]),
            ("Relative molecular mass:", s12["c"]),
            ("Physical description", s13["a"]),
            ("Other:", s13["other"]),
        ]

        for label, value in field_map:
            idx = self._find_para_index_doc(doc, label, start_idx, end_idx)
            if idx is None:
                continue
            if value == "__IMG__":
                img_paths = extracted["3.2.S.1.2"].image_paths
                if img_paths:
                    img_para = self._insert_paragraph_after(doc.paragraphs[idx], "")
                    self._add_picture_autofit(img_para.add_run(), img_paths[0], doc)
            else:
                self._add_answer_after(doc, idx, value)

        sol_idx = self._find_para_index_doc(doc, "Solubilities:", start_idx, end_idx)
        if sol_idx is not None:
            self._replace_paragraph_with_runs(
                doc.paragraphs[sol_idx],
                [("(b)", False, False), ("\tSolubilities", False, False), (":", False, False), (" NA", False, False)],
            )
            if sol_idx + 1 < len(doc.paragraphs) and not (doc.paragraphs[sol_idx + 1].text or "").strip():
                doc.paragraphs[sol_idx + 1].text = s13["b"]
            else:
                self._insert_paragraph_after(doc.paragraphs[sol_idx], s13["b"])

        self._set_value_under_label(doc, "Polymorphic form:", s13["poly"], start_idx, end_idx)
        self._set_value_under_label(doc, "Solvate:", s13["solvate"], start_idx, end_idx)
        self._set_value_under_label(doc, "Hydrate:", s13["hydrate"], start_idx, end_idx)

        self._remove_paragraphs_matching(
            doc,
            [
                "Refer Section 3.2.S.1.1",
                "Refer Section 3.2.S.1.2",
                "Refer Section 3.2.S.1.3",
                "Company or laboratory code:",
                "Other non-proprietary name(s)",
                "Chemical Abstracts Service (CAS) registry number:",
            ],
            start_idx,
            end_idx,
        )

        self._fill_property_table(doc, s13)
        self._apply_visual_formatting(doc, start_idx, end_idx)
        preserve_patterns = tuple(p for p in (name_mfr_line,) if p)
        cleanup_stats = run_artifact_cleanup(
            doc,
            keep_first_n_tables=template_table_count,
            preserve_repeated_patterns=preserve_patterns,
        )
        if any(cleanup_stats.values()):
            warnings.append(f"cleanup: {cleanup_stats}")

        output_docx.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_docx)
        return warnings


# backward-compat alias
DocxFiller = S1DocxFiller
