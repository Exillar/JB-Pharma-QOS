"""builders/generic.py — GenericSectionFiller: template-driven filler for any CTD section.

Works for 2.3.S.5, 2.3.S.6, 2.3.S.7, 2.3.P.1–P.8, and any future sections.
Scans the QOS template for "Refer Section 3.2.X.Y" placeholder paragraphs and
replaces each one with the extracted text, tables, and images from the source PDF.
No drug-specific or section-specific hardcoding.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.document import Document as _Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from builders.base import _DocxHelper, run_artifact_cleanup
from ctd_utils import section_flexible_regex
from pdf_extractor import ExtractedSectionContent

_REFER_SECTION_RE = re.compile(
    r"^refer\s+section\s+(3\.2\.[sp](?:\.[0-9a-z]+)+)\b",
    re.IGNORECASE,
)
_QOS_HEADING_RE = re.compile(r"^2\.3\.[sp]\.", re.IGNORECASE)


class GenericSectionFiller(_DocxHelper):
    """Template-driven filler that replaces 'Refer Section X.X.X' placeholders.

    Suitable for S5, S6, S7, P1–P8 and any CTD section where the QOS template
    contains one or more lines of the form "Refer Section 3.2.X.Y".
    """

    def __init__(
        self,
        template_docx: Path,
        section_start: str,
        section_end: str,
        filled_reference_docx: Path | None = None,
        *,
        preserve_repeated_patterns: tuple[str, ...] = (),
    ) -> None:
        self.template_docx = template_docx
        self.SECTION_START = section_start
        self.SECTION_END = section_end
        self.filled_reference_docx = filled_reference_docx
        self._preserve_repeated_patterns = preserve_repeated_patterns

    # ------------------------------------------------------------------
    # Placeholder scanning
    # ------------------------------------------------------------------

    @staticmethod
    def _normalize_section_id(raw: str) -> str:
        """'3.2.s.5.1' → '3.2.S.5.1'"""
        parts = [p for p in re.split(r"[\s.]+", raw.strip()) if p]
        return ".".join(
            p.upper() if re.fullmatch(r"[sp]", p, re.IGNORECASE) else p
            for p in parts
        )

    def _scan_refer_placeholders(
        self, doc: _Document, start_idx: int, end_idx: int
    ) -> list[tuple[Paragraph, str]]:
        """Return (paragraph, normalized_section_id) for every 'Refer Section' line."""
        results: list[tuple[Paragraph, str]] = []
        for i in range(start_idx, min(end_idx, len(doc.paragraphs))):
            text = (doc.paragraphs[i].text or "").strip()
            m = _REFER_SECTION_RE.match(text)
            if m:
                results.append(
                    (doc.paragraphs[i], self._normalize_section_id(m.group(1)))
                )
        return results

    # ------------------------------------------------------------------
    # Content helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _table_has_content(rows: list[list[str]]) -> bool:
        return any((cell or "").strip() for row in rows for cell in row)

    def _extract_narrative(
        self,
        content: ExtractedSectionContent,
        section_id: str,
        table_row_texts: set[str],
    ) -> str:
        section_pat = re.compile(
            rf"^\s*{section_flexible_regex(section_id)}(?:\s+.*)?$",
            re.IGNORECASE,
        )
        refer_pat = re.compile(
            rf"^\s*refer\s+section\s+{section_flexible_regex(section_id)}\b",
            re.IGNORECASE,
        )
        page_pat = re.compile(r"^\s*(\d+\s+of\s+\d+|\d+)\s*$", re.IGNORECASE)
        header_pat = re.compile(
            r"^\s*(module\s*:|version\s*:|date\s*:|open\s+part|confidential)\b",
            re.IGNORECASE,
        )
        footer_pat = re.compile(r"^.{5,75}\s+\d{1,4}$")

        out_lines: list[str] = []
        for raw_line in (content.raw_text or "").splitlines():
            line = re.sub(r"\s+", " ", raw_line).strip()
            if not line:
                continue
            if section_pat.match(line) or refer_pat.match(line) or page_pat.match(line):
                continue
            if header_pat.match(line):
                continue
            if footer_pat.match(line) and len(line) < 80:
                continue
            if table_row_texts and self._norm(line) in table_row_texts:
                continue
            out_lines.append(line)

        paragraphs: list[str] = []
        current: list[str] = []
        for line in out_lines:
            current.append(line)
            if re.search(r"[\.;:]\s*$", line):
                paragraphs.append(" ".join(current).strip())
                current = []
        if current:
            paragraphs.append(" ".join(current).strip())

        seen: set[str] = set()
        deduped: list[str] = []
        for para in paragraphs:
            norm = self._norm(para)
            if norm and norm not in seen:
                deduped.append(para)
                seen.add(norm)

        return "\n".join(deduped).strip()

    # ------------------------------------------------------------------
    # Table/paragraph insertion
    # ------------------------------------------------------------------

    def _insert_table_after(
        self,
        doc: _Document,
        paragraph: Paragraph,
        rows: list[list[str]],
    ):
        if not rows:
            return None
        max_cols = max((len(r) for r in rows if r), default=0)
        if max_cols <= 0:
            return None
        norm_rows = [
            [re.sub(r"\s+", " ", (c or "").strip()) for c in row]
            + [""] * (max_cols - len(row))
            for row in rows
        ]
        table = doc.add_table(rows=len(norm_rows), cols=max_cols)
        for r_idx, row in enumerate(norm_rows):
            for c_idx, cell_text in enumerate(row):
                table.cell(r_idx, c_idx).text = cell_text
        paragraph._p.addnext(table._tbl)
        return table

    @staticmethod
    def _insert_paragraph_after_table(table, text: str = "") -> Paragraph:
        new_p = OxmlElement("w:p")
        table._tbl.addnext(new_p)
        para = Paragraph(new_p, table._parent)
        if text:
            para.add_run(text)
        return para

    # ------------------------------------------------------------------
    # Per-placeholder replacement
    # ------------------------------------------------------------------

    def _replace_refer_placeholder(
        self,
        doc: _Document,
        target_para: Paragraph,
        content: ExtractedSectionContent,
        section_id: str,
        warnings: list[str],
    ) -> None:
        if content.warning:
            warnings.append(f"{section_id}: {content.warning}")

        tables = [t for t in content.tables if self._table_has_content(t)]
        table_row_texts: set[str] = set()
        for tbl in tables:
            for row in tbl:
                joined = " ".join((c or "").strip() for c in row if (c or "").strip())
                if joined:
                    table_row_texts.add(self._norm(joined))

        narrative = self._extract_narrative(content, section_id, table_row_texts)
        images = [Path(p) for p in content.image_paths if Path(p).is_file()]

        if not narrative and not tables and not images:
            warnings.append(f"{section_id}: no content extracted, placeholder kept")
            return

        cursor: Paragraph = target_para
        tables_placed = False

        if narrative:
            for line in narrative.splitlines():
                if line.strip():
                    cursor = self._insert_paragraph_after(cursor, line)
                if tables and not tables_placed and re.search(
                    r"\btable\s*\d+\b", line, flags=re.IGNORECASE
                ):
                    for t_rows in tables:
                        tbl = self._insert_table_after(doc, cursor, t_rows)
                        if tbl is not None:
                            cursor = self._insert_paragraph_after_table(tbl)
                    tables_placed = True

        if tables and not tables_placed:
            for t_rows in tables:
                tbl = self._insert_table_after(doc, cursor, t_rows)
                if tbl is not None:
                    cursor = self._insert_paragraph_after_table(tbl)

        for img in images[:1]:
            try:
                cursor = self._insert_paragraph_after(cursor)
                self._add_picture_autofit(cursor.add_run(), img, doc)
            except Exception as e:
                warnings.append(f"{section_id}: failed to insert image {img.name}: {e}")

        self._delete_paragraph(target_para)

    # ------------------------------------------------------------------
    # Public entry point
    # ------------------------------------------------------------------

    def fill_section(
        self,
        extracted: dict[str, ExtractedSectionContent],
        output_docx: Path,
    ) -> list[str]:
        doc = Document(self.template_docx)
        template_table_count = len(doc.tables)
        warnings: list[str] = []

        start_idx, end_idx = self._get_target_range(doc)

        # Insert name/manufacturer line under each QOS sub-heading in this section
        name_line = self._resolve_name_manufacturer_line(
            self.filled_reference_docx, self.SECTION_START
        )
        if name_line:
            for i in range(start_idx, min(end_idx, len(doc.paragraphs))):
                text = (doc.paragraphs[i].text or "").strip()
                if _QOS_HEADING_RE.match(text):
                    nxt = (
                        (doc.paragraphs[i + 1].text or "").strip()
                        if i + 1 < len(doc.paragraphs)
                        else ""
                    )
                    if not nxt.startswith("("):
                        self._insert_paragraph_after(doc.paragraphs[i], name_line)

            start_idx, end_idx = self._get_target_range(doc)

        placeholders = self._scan_refer_placeholders(doc, start_idx, end_idx)
        for target_para, section_id in placeholders:
            content = extracted.get(section_id)
            if content is None:
                warnings.append(
                    f"{section_id}: not in extracted payload, placeholder kept"
                )
                continue
            self._replace_refer_placeholder(
                doc, target_para, content, section_id, warnings
            )

        preserve_patterns = tuple(self._preserve_repeated_patterns)
        if name_line:
            preserve_patterns = preserve_patterns + (name_line,)
        cleanup_stats = run_artifact_cleanup(
            doc,
            keep_first_n_tables=template_table_count,
            preserve_repeated_patterns=preserve_patterns,
        )
        if any(cleanup_stats.values()):
            warnings.append(f"cleanup: {cleanup_stats}")

        output_docx.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_docx)
        return warnings
