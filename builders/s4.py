"""builders/s4.py — S4DocxFiller: fills QOS section 2.3.S.4 Control of the API."""
from __future__ import annotations

import re
from pathlib import Path

import fitz
from docx import Document
from docx.document import Document as _Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from builders.base import _DocxHelper, run_artifact_cleanup
from ctd_utils import section_flexible_regex
from pdf_extractor import ExtractedSectionContent


class S4DocxFiller(_DocxHelper):
    SECTION_START = "2.3.S.4 Control of the API"
    SECTION_END = "2.3.S.5 Reference Standards or Materials"

    def __init__(
        self,
        template_docx: Path,
        filled_reference_docx: Path | None = None,
    ) -> None:
        self.template_docx = template_docx
        self.filled_reference_docx = filled_reference_docx

    @staticmethod
    def _clean_cell(text: str) -> str:
        return re.sub(r"\s+", " ", (text or "").strip())

    # ------------------------------------------------------------------
    # Template table detection
    # ------------------------------------------------------------------

    def _find_s41_template_table(self, doc: _Document) -> object | None:
        """Find the 3-column specification table in the QOS template."""
        for table in doc.tables:
            if len(table.columns) != 3 or len(table.rows) < 3:
                continue
            row0 = [self._norm(c.text) for c in table.rows[0].cells[:3]]
            row2 = [self._norm(c.text) for c in table.rows[2].cells[:3]]
            if (
                row0[0].startswith("standard")
                and row2[0] == "test"
                and "acceptance criteria" in row2[1]
                and "analytical procedure" in row2[2]
            ):
                return table
        return None

    # ------------------------------------------------------------------
    # Spec row extraction — PDF tables (generic, drug-agnostic)
    # ------------------------------------------------------------------

    def _extract_s41_source_rows_from_pdf_tables(
        self, pdf_path: Path
    ) -> list[list[str]]:
        """Extract [test, method, acceptance] rows from spec tables in the PDF.

        Page filter: only process pages containing 'acceptance criteria' or
        'analytical procedure' — generic CTD spec table column headers.
        """
        rows: list[list[str]] = []
        seen: set[tuple[str, str, str]] = set()
        try:
            with fitz.open(str(pdf_path)) as doc:
                for page in doc:
                    page_text = " ".join(page.get_text("text", sort=True).split()).lower()
                    if not any(
                        kw in page_text
                        for kw in ("acceptance criteria", "analytical procedure")
                    ):
                        continue
                    try:
                        tables = page.find_tables().tables
                    except Exception:
                        tables = []
                    for tab in tables:
                        raw_rows = tab.extract()
                        if not raw_rows:
                            continue
                        header_text = " ".join(
                            self._clean_cell(str(cell or ""))
                            for row in raw_rows[:3]
                            for cell in row
                        ).lower()
                        if not any(
                            token in header_text
                            for token in ("acceptance criteria", "method", "test")
                        ):
                            continue
                        current_test = ""
                        current_method = ""
                        for raw_row in raw_rows:
                            cells = [self._clean_cell(str(c or "")) for c in raw_row]
                            if not any(cells):
                                continue
                            test = method = acceptance = ""
                            if len(cells) >= 9:
                                test = self._clean_cell(" ".join(cells[0:3]))
                                method = self._clean_cell(" ".join(cells[3:6]))
                                acceptance = self._clean_cell(" ".join(cells[6:9]))
                            elif len(cells) >= 6:
                                span = max(1, len(cells) // 3)
                                test = self._clean_cell(" ".join(cells[0:span]))
                                method = self._clean_cell(" ".join(cells[span: 2 * span]))
                                acceptance = self._clean_cell(" ".join(cells[2 * span:]))
                            elif len(cells) >= 3:
                                test, method, acceptance = cells[0], cells[1], cells[2]
                            elif len(cells) == 2:
                                test, method, acceptance = "", cells[0], cells[1]
                            else:
                                continue
                            combined = " ".join((test, method, acceptance)).lower()
                            if (
                                not combined
                                or (
                                    "test" in test.lower()
                                    and "method" in method.lower()
                                    and "acceptance" in acceptance.lower()
                                )
                                or (test.lower() == "test" and not method and not acceptance)
                            ):
                                continue
                            if test:
                                current_test = test
                            elif current_test.lower() == "identification":
                                test = ""
                            if method:
                                current_method = method
                            elif acceptance and test:
                                method = current_method
                            row_key = (test.lower(), method.lower(), acceptance.lower())
                            if row_key in seen:
                                continue
                            seen.add(row_key)
                            rows.append([test, method, acceptance])
        except Exception:
            return []
        return self._fill_down_s41_methods(rows)

    def _fill_down_s41_methods(self, rows: list[list[str]]) -> list[list[str]]:
        """Carry method value down through sub-rows; clear method on header rows."""
        filled: list[list[str]] = []
        carry_method = ""
        for test, method, acceptance in rows:
            test_clean = self._clean_cell(test)
            method_clean = self._clean_cell(method)
            acceptance_clean = self._clean_cell(acceptance)
            if method_clean:
                carry_method = method_clean
            elif acceptance_clean and test_clean:
                method_clean = carry_method
            # Header/group-label rows: test name present but no acceptance → no method
            if test_clean and not acceptance_clean:
                method_clean = ""
            filled.append([test_clean, method_clean, acceptance_clean])
        return filled

    # ------------------------------------------------------------------
    # Fallback: score and choose best table from already-extracted content
    # ------------------------------------------------------------------

    def _choose_s41_source_table(
        self, content: ExtractedSectionContent
    ) -> list[list[str]]:
        """Score pre-extracted tables and return the best spec table, or []."""
        best_score = -1
        best_table: list[list[str]] = []
        caption_hit = bool(
            re.search(r"\btable\s+\d+.*\bspecification", content.raw_text, flags=re.IGNORECASE)
        )
        for table in content.tables:
            flat = " ".join(self._clean_cell(cell) for row in table for cell in row).lower()
            score = 0
            if "acceptance criteria" in flat:
                score += 5
            if "analytical procedure" in flat:
                score += 4
            elif "method" in flat:
                score += 3
            if re.search(r"\btest\b", flat):
                score += 3
            if "specification" in flat:
                score += 2
            if caption_hit:
                score += 1
            if score > best_score:
                best_score = score
                best_table = table
        return best_table if best_score >= 8 else []

    # ------------------------------------------------------------------
    # Orchestrate extraction
    # ------------------------------------------------------------------

    def _extract_s41_metadata_and_rows(
        self,
        content: ExtractedSectionContent,
    ) -> tuple[str, str, list[list[str]]]:
        """Return (standard, spec_ref, rows) in [test, acceptance, method] column order."""
        rows = self._extract_s41_source_rows_from_pdf_tables(content.source_pdf)
        if rows:
            # pdf_tables returns [test, method, acceptance]; template expects [test, acc, method]
            return "USP", "-----", [[t, a, m] for t, m, a in rows]

        raw = self._choose_s41_source_table(content)
        if raw:
            header = [self._norm(c) for c in (raw[0] if raw else [])]
            test_col = next((i for i, h in enumerate(header) if h == "test"), 0)
            acc_col = next((i for i, h in enumerate(header) if "acceptance" in h), 1)
            meth_col = next(
                (i for i, h in enumerate(header) if "procedure" in h or "method" in h), 2
            )
            data_rows = raw[1:]
            normalized = [
                [
                    r[test_col] if test_col < len(r) else "",
                    r[acc_col] if acc_col < len(r) else "",
                    r[meth_col] if meth_col < len(r) else "",
                ]
                for r in data_rows
            ]
            return "USP", "-----", normalized

        return "USP", "-----", []

    # ------------------------------------------------------------------
    # Narrative extraction (generic — no drug-specific patterns)
    # ------------------------------------------------------------------

    def _extract_clean_section_narrative(
        self,
        content: ExtractedSectionContent,
        refer_section: str,
        table_row_texts: set[str] | None = None,
    ) -> str:
        section_pat = re.compile(
            rf"^\s*{section_flexible_regex(refer_section)}(?:\s+.*)?$",
            re.IGNORECASE,
        )
        refer_pat = re.compile(
            rf"^\s*refer\s+section\s+{section_flexible_regex(refer_section)}\b",
            re.IGNORECASE,
        )
        page_pat = re.compile(r"^\s*(\d+\s+of\s+\d+|\d+)\s*$", re.IGNORECASE)
        header_pat = re.compile(
            r"^\s*(module\s*:|version\s*:|date\s*:|open\s+part|confidential)\b",
            re.IGNORECASE,
        )
        footer_pat = re.compile(r"^.{5,75}\s+\d{1,4}$")

        table_rows = table_row_texts or set()
        out_lines: list[str] = []
        for raw_line in (content.raw_text or "").splitlines():
            line = re.sub(r"\s+", " ", raw_line).strip()
            if not line:
                continue
            if section_pat.match(line) or refer_pat.match(line) or page_pat.match(line):
                continue
            if header_pat.match(line):
                continue
            if footer_pat.match(line) and len(line) < 80:
                continue
            if table_rows and self._norm(line) in table_rows:
                continue
            out_lines.append(line)

        paragraphs: list[str] = []
        current: list[str] = []
        for line in out_lines:
            current.append(line)
            if re.search(r"[\.;:]\s*$", line):
                paragraphs.append(" ".join(current).strip())
                current = []
        if current:
            paragraphs.append(" ".join(current).strip())

        deduped: list[str] = []
        seen: set[str] = set()
        for para in paragraphs:
            norm = self._norm(para)
            if not norm or norm in seen:
                continue
            deduped.append(para)
            seen.add(norm)

        return "\n".join(deduped).strip()

    # ------------------------------------------------------------------
    # S4.5 content insertion helper
    # ------------------------------------------------------------------

    @staticmethod
    def _select_first_existing_image(
        content: ExtractedSectionContent,
    ) -> Path | None:
        for img_path in content.image_paths:
            try:
                p = Path(img_path)
                if p.exists() and p.is_file():
                    return p
            except Exception:
                continue
        return None

    def _table_has_content(self, rows: list[list[str]]) -> bool:
        return any(self._clean_cell(cell) for row in rows for cell in row)

    def _normalize_table_rows(self, rows: list[list[str]]) -> list[list[str]]:
        if not rows:
            return []
        max_cols = max(len(r) for r in rows if r)
        if max_cols <= 0:
            return []
        return [
            [self._clean_cell(c) for c in row] + [""] * (max_cols - len(row))
            for row in rows
        ]

    def _insert_table_after(
        self,
        doc: _Document,
        paragraph: Paragraph,
        rows: list[list[str]],
    ):
        normalized = self._normalize_table_rows(rows)
        if not normalized:
            return None
        cols = len(normalized[0])
        table = doc.add_table(rows=len(normalized), cols=cols)
        for r_idx, row in enumerate(normalized):
            for c_idx in range(cols):
                table.cell(r_idx, c_idx).text = row[c_idx]
        paragraph._p.addnext(table._tbl)
        return table

    @staticmethod
    def _insert_paragraph_after_table(table, text: str = "") -> Paragraph:
        new_p = OxmlElement("w:p")
        table._tbl.addnext(new_p)
        para = Paragraph(new_p, table._parent)
        if text:
            para.add_run(text)
        return para

    def _populate_section_answer_with_content(
        self,
        doc: _Document,
        content: ExtractedSectionContent,
        refer_section: str,
        prompt_needle: str,
        start_idx: int,
        end_idx: int,
    ) -> str | None:
        idx = self._find_para_index_doc(doc, prompt_needle, start_idx, end_idx)
        if idx is None:
            return f"{refer_section}: target prompt not found in template"

        cursor = doc.paragraphs[idx]
        tables = [t for t in content.tables if self._table_has_content(t)]
        table_row_texts: set[str] = set()
        for table_rows in tables:
            for row in table_rows:
                joined = " ".join(self._clean_cell(cell) for cell in row if cell)
                if joined:
                    table_row_texts.add(self._norm(joined))

        narrative = self._extract_clean_section_narrative(
            content, refer_section, table_row_texts
        )
        first_image = self._select_first_existing_image(content)

        if not narrative and not tables and first_image is None:
            return f"{refer_section}: no usable image/text/table extracted"

        tables_inserted = False
        if narrative:
            for line in narrative.splitlines():
                cursor = self._insert_paragraph_after(cursor, line)
                if tables and not tables_inserted and re.search(
                    r"\btable\s*\d+\b", line, flags=re.IGNORECASE
                ):
                    for t_rows in tables:
                        tbl = self._insert_table_after(doc, cursor, t_rows)
                        if tbl is not None:
                            cursor = self._insert_paragraph_after_table(tbl)
                    tables_inserted = True

        if tables and not tables_inserted:
            for t_rows in tables:
                tbl = self._insert_table_after(doc, cursor, t_rows)
                if tbl is not None:
                    cursor = self._insert_paragraph_after_table(tbl)
            tables_inserted = True

        if first_image is not None:
            try:
                cursor = self._insert_paragraph_after(cursor)
                cursor.add_run().add_picture(str(first_image))
            except Exception as e:
                return (
                    f"{refer_section}: failed to insert image {first_image.name}: {e}"
                )

        return None

    # ------------------------------------------------------------------
    # Spec table helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _trim_table_rows(table, keep_rows: int) -> None:
        while len(table.rows) > keep_rows:
            row = table.rows[-1]._tr
            row.getparent().remove(row)

    def _append_rows(self, table, rows: list[list[str]]) -> None:
        for row_data in rows:
            row = table.add_row()
            for idx, value in enumerate(row_data[:3]):
                row.cells[idx].text = value

    # ------------------------------------------------------------------
    # Bold-question + plain-answer paragraph formatting
    # ------------------------------------------------------------------

    def _set_paragraph_text_with_suffix(
        self,
        doc: _Document,
        needle: str,
        suffix: str,
        start_idx: int,
        end_idx: int,
    ) -> None:
        idx = self._find_para_index_doc(doc, needle, start_idx, end_idx)
        if idx is None:
            return
        paragraph = doc.paragraphs[idx]
        base = re.sub(r"\s+", " ", (paragraph.text or "").strip())
        if not base:
            return
        for run in list(paragraph.runs):
            paragraph._p.remove(run._r)
        q_run = paragraph.add_run(base)
        q_run.bold = True
        q_run.italic = False
        a_run = paragraph.add_run(f" {suffix}".rstrip())
        a_run.bold = False
        a_run.italic = False

    # ------------------------------------------------------------------
    # Public entry point
    # ------------------------------------------------------------------

    def fill_s4_section(
        self,
        extracted: dict[str, ExtractedSectionContent],
        output_docx: Path,
    ) -> list[str]:
        doc = Document(self.template_docx)
        template_table_count = len(doc.tables)
        warnings: list[str] = []

        start_idx, end_idx = self._get_target_range(doc)

        s41_name = self._resolve_name_manufacturer_line(
            self.filled_reference_docx, "2.3.S.4.1 Specification"
        )
        s42_name = self._resolve_name_manufacturer_line(
            self.filled_reference_docx, "2.3.S.4.2 Analytical Procedures"
        )
        s45_name = self._resolve_name_manufacturer_line(
            self.filled_reference_docx, "2.3.S.4.5 Justification of Specification"
        )

        for heading in [
            "2.3.S.4 Control of the API",
            "2.3.S.4.1 Specification",
            "API specifications of the FPP manufacturer",
        ]:
            idx = self._find_para_index_doc(doc, heading, start_idx, end_idx)
            if idx is not None and s41_name:
                nxt = (
                    (doc.paragraphs[idx + 1].text or "").strip()
                    if idx + 1 < len(doc.paragraphs)
                    else ""
                )
                if not nxt.startswith("("):
                    self._insert_paragraph_after(doc.paragraphs[idx], s41_name)

        for heading, name_line in [
            ("2.3.S.4.2 Analytical Procedures", s42_name),
            ("2.3.S.4.5 Justification of Specification", s45_name),
        ]:
            idx = self._find_para_index_doc(doc, heading, start_idx, end_idx)
            if idx is not None and name_line:
                nxt = (
                    (doc.paragraphs[idx + 1].text or "").strip()
                    if idx + 1 < len(doc.paragraphs)
                    else ""
                )
                if not nxt.startswith("("):
                    self._insert_paragraph_after(doc.paragraphs[idx], name_line)

        start_idx, end_idx = self._get_target_range(doc)
        self._remove_paragraphs_matching(
            doc, ["refer section 3.2.s.4."], start_idx, end_idx
        )

        content = extracted["3.2.S.4.1"]
        if content.warning:
            warnings.append(f"3.2.S.4.1: {content.warning}")

        spec_table = self._find_s41_template_table(doc)
        if spec_table is None:
            warnings.append("2.3.S.4.1: target specification table not found in template")
        else:
            standard, spec_ref, rows = self._extract_s41_metadata_and_rows(content)
            spec_table.rows[0].cells[-1].text = standard or "USP"
            spec_table.rows[1].cells[-1].text = spec_ref or "-----"
            spec_table.rows[2].cells[0].text = "Test"
            spec_table.rows[2].cells[1].text = "Acceptance criteria"
            spec_table.rows[2].cells[2].text = "Analytical procedure\n(Type/Source/Version)"
            self._trim_table_rows(spec_table, keep_rows=3)
            self._append_rows(spec_table, rows)
            if not rows:
                warnings.append("2.3.S.4.1: no specification rows extracted from PDF")

        content_42 = extracted.get("3.2.S.4.2")
        if content_42 and content_42.warning:
            warnings.append(f"3.2.S.4.2: {content_42.warning}")
        self._set_paragraph_text_with_suffix(
            doc,
            "Summary of the analytical procedures",
            "Please refer to Module 3 Section 3.2.S.4.2",
            start_idx,
            end_idx,
        )

        content_45 = extracted.get("3.2.S.4.5")
        if content_45 and content_45.warning:
            warnings.append(f"3.2.S.4.5: {content_45.warning}")
        if content_45:
            warn = self._populate_section_answer_with_content(
                doc,
                content_45,
                "3.2.S.4.5",
                "Justification of the API specification",
                start_idx,
                end_idx,
            )
            if warn:
                warnings.append(warn)

        cleanup_stats = run_artifact_cleanup(
            doc, keep_first_n_tables=template_table_count
        )
        if any(cleanup_stats.values()):
            warnings.append(f"cleanup: {cleanup_stats}")

        output_docx.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_docx)
        return warnings
