"""builders/s3.py — S3DocxFiller: fills QOS section 2.3.S.3.1 Characterisation."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.document import Document as _Document
from docx.text.paragraph import Paragraph

from builders.base import _DocxHelper, run_artifact_cleanup
from config_loader import S2FillConfig, S3Config
from pdf_extractor import ExtractedSectionContent


class S3DocxFiller(_DocxHelper):
    SECTION_START = "2.3.S.3 Characterisation"
    SECTION_END = "2.3.S.4 Control of Drug Substance"

    _PAGE_STAMP_RE = re.compile(r"^(page\s*)?\d+\s+of\s+\d+$", re.IGNORECASE)

    def __init__(
        self,
        template_docx: Path,
        filled_reference_docx: Path | None = None,
        s2_fill_cfg: S2FillConfig | None = None,
        s3_cfg: S3Config | None = None,
    ) -> None:
        self.template_docx = template_docx
        self.filled_reference_docx = filled_reference_docx
        # Accept either the old s2_fill_cfg (backward compat) or the new s3_cfg
        self.cfg = s3_cfg or S3Config()
        # Also store s2_fill_cfg for restricted_phrase_keywords
        _s2 = s2_fill_cfg or S2FillConfig()
        self._restricted_phrase_keywords = _s2.restricted_phrase_keywords

    def _clean_block(self, text: str) -> str:
        out: list[str] = []
        for ln in text.splitlines():
            s = ln.strip()
            if not s:
                continue
            if self._PAGE_STAMP_RE.match(s):
                continue
            out.append(s)
        return "\n".join(out).strip()

    @staticmethod
    def _lines(text: str) -> list[str]:
        return [ln.strip() for ln in text.splitlines() if ln.strip()]

    @staticmethod
    def _normalize_s31_heading(doc: _Document, start_idx: int, end_idx: int) -> None:
        for i in range(start_idx, min(end_idx, len(doc.paragraphs))):
            p = doc.paragraphs[i]
            text = (p.text or "").strip()
            if "2.3.S.3.1" not in text:
                continue
            cleaned = re.sub(
                r"\s*refer\s+section\s+3\.2\.[sp]\.3\.1\s*$",
                "",
                text,
                flags=re.IGNORECASE,
            ).strip()
            if cleaned and cleaned != text:
                p.text = cleaned
            return

    def _parse_s31(self, text: str) -> dict[str, str]:
        lines = self._lines(self._clean_block(text))
        low_lines = [ln.lower() for ln in lines]

        start_idx = None
        for i, low in enumerate(low_lines):
            if any(k in low for k in self.cfg.s31_summary_start_keywords):
                start_idx = i
                break

        summary_lines: list[str] = []
        if start_idx is not None:
            for ln in lines[start_idx:]:
                low = ln.lower()
                if any(k in low for k in self.cfg.s31_summary_stop_keywords):
                    break
                if re.match(r"^3\.2\.[sp]\.3(\.\d+)*\b", ln, re.IGNORECASE):
                    if summary_lines:
                        break
                    continue
                summary_lines.append(ln)
                if len(summary_lines) >= self.cfg.s31_max_summary_lines:
                    break

        answer_a = "\n".join(summary_lines).strip()
        if not answer_a and lines:
            answer_a = lines[0]

        return {
            "a": answer_a,
            "b": self.cfg.s31_isomerism_default,
            "c": self.cfg.s31_polymorph_reference_default,
            "d": self.cfg.s31_particle_size_default,
            "e": self.cfg.s31_other_characteristics_default,
        }

    def _remove_refer_section_lines(self, doc: _Document, start_idx: int, end_idx: int) -> None:
        pat = re.compile(r"^refer\s+section\s+3\.2\.[sp]\.3(\.\d+)*", re.IGNORECASE)
        to_del = [
            i for i in range(start_idx, min(end_idx, len(doc.paragraphs)))
            if pat.match(self._norm(doc.paragraphs[i].text or ""))
        ]
        for i in reversed(to_del):
            if i < len(doc.paragraphs):
                self._delete_paragraph(doc.paragraphs[i])

    def fill_s3_section(
        self,
        extracted: dict[str, ExtractedSectionContent],
        output_docx: Path,
    ) -> list[str]:
        doc = Document(self.template_docx)
        template_table_count = len(doc.tables)
        warnings: list[str] = []

        start_idx, end_idx = self._get_target_range(doc)

        name_line = self._resolve_name_manufacturer_line(
            self.filled_reference_docx, "2.3.S.3 Characterisation"
        )
        idx = self._find_para_index_doc(doc, "2.3.S.3 Characterisation", start_idx, end_idx)
        if idx is not None and name_line:
            nxt = (doc.paragraphs[idx + 1].text or "").strip() if idx + 1 < len(doc.paragraphs) else ""
            if not nxt.startswith("("):
                self._insert_paragraph_after(doc.paragraphs[idx], name_line)

        if extracted["3.2.S.3.1"].warning:
            warnings.append(f"3.2.S.3.1: {extracted['3.2.S.3.1'].warning}")

        s31 = self._parse_s31(extracted["3.2.S.3.1"].raw_text)

        start_idx, end_idx = self._get_target_range(doc)
        self._remove_refer_section_lines(doc, start_idx, end_idx)
        self._normalize_s31_heading(doc, start_idx, end_idx)

        self._remove_paragraphs_matching(
            doc,
            ["<including identification of and data on the api lot used in bioavailability studies>"],
            start_idx, end_idx,
        )

        idx = self._find_para_index_doc(doc, "List of studies performed", start_idx, end_idx)
        if idx is not None and s31["a"].strip():
            self._insert_paragraph_after(doc.paragraphs[idx], s31["a"].strip())

        idx = self._find_para_index_doc(doc, "Discussion on the potential for isomerism", start_idx, end_idx)
        if idx is not None and s31["b"].strip():
            self._append_inline_value_after_colon(doc.paragraphs[idx], s31["b"].strip())

        idx = self._find_para_index_doc(
            doc, "Summary of studies performed to identify potential polymorphic forms", start_idx, end_idx
        )
        if idx is not None and s31["c"].strip():
            self._strip_angle_bracket_placeholder(doc.paragraphs[idx])
            if idx + 1 < len(doc.paragraphs) and not (doc.paragraphs[idx + 1].text or "").strip():
                doc.paragraphs[idx + 1].text = s31["c"].strip()
            else:
                self._insert_paragraph_after(doc.paragraphs[idx], s31["c"].strip())

        start_idx, end_idx = self._get_target_range(doc)

        idx = self._find_para_index_doc(
            doc, "Summary of studies performed to identify the particle size distribution of the API:",
            start_idx, end_idx,
        )
        if idx is not None and s31["d"].strip():
            self._strip_angle_bracket_placeholder(doc.paragraphs[idx])
            self._append_inline_value_after_colon(doc.paragraphs[idx], s31["d"].strip())

        idx = self._find_para_index_doc(doc, "Other characteristics:", start_idx, end_idx)
        if idx is not None and s31["e"].strip():
            self._append_inline_value_after_colon(doc.paragraphs[idx], s31["e"].strip())

        preserve_patterns = tuple(self._restricted_phrase_keywords)
        if name_line:
            preserve_patterns = preserve_patterns + (name_line,)
        cleanup_stats = run_artifact_cleanup(
            doc,
            keep_first_n_tables=template_table_count,
            preserve_repeated_patterns=preserve_patterns,
            preserve_phrases=(s31["c"],),
        )
        if any(cleanup_stats.values()):
            warnings.append(f"cleanup: {cleanup_stats}")

        output_docx.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_docx)
        return warnings
