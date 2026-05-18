"""builders/base.py — _DocxHelper base class and document cleanup utilities.

All section fillers inherit from _DocxHelper. Cleanup functions are
module-level so they can be called independently.
"""
from __future__ import annotations

import re
from pathlib import Path

import fitz
from docx.document import Document as _Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from pdf_extractor import ExtractedSectionContent  # noqa: F401 (re-exported)

_PAGE_NUM_RE = re.compile(r"^\d{1,4}$")
_PAGE_OF_RE = re.compile(r"^\d+\s+of\s+\d+\s*$", re.IGNORECASE)
_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ---------------------------------------------------------------------------
# Document-level cleanup (shared across all fillers)
# ---------------------------------------------------------------------------

def _safe_row_cells(row):
    try:
        return tuple(row.cells)
    except Exception:
        return ()


def _collapse_blank_paragraphs(doc: _Document) -> int:
    consecutive = 0
    to_remove = []
    for elem in list(doc.element.body):
        if elem.tag.split("}")[-1] == "p":
            txt = "".join(t.text or "" for t in elem.iter(f"{{{_NS}}}t")).strip()
            has_drawing = len(list(elem.iter(f"{{{_NS}}}drawing"))) > 0
            if not txt and not has_drawing:
                consecutive += 1
                if consecutive > 1:
                    to_remove.append(elem)
            else:
                consecutive = 0
        else:
            consecutive = 0
    for elem in to_remove:
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)
    return len(to_remove)


def _remove_repeated_header_paragraphs(
    doc: _Document,
    *,
    preserve_patterns: tuple[str, ...] = (),
) -> int:
    from collections import Counter
    freq = Counter(p.text.strip() for p in doc.paragraphs if p.text.strip())
    preserve = tuple(p.lower() for p in preserve_patterns if p)
    noise = {
        t
        for t, c in freq.items()
        if c >= 3 and len(t) < 120 and not any(p in t.lower() for p in preserve)
    }
    removed = 0
    for elem in list(doc.element.body):
        if elem.tag.split("}")[-1] != "p":
            continue
        text = "".join(t.text or "" for t in elem.iter(f"{{{_NS}}}t")).strip()
        if text in noise:
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)
                removed += 1
    return removed


def _remove_pdf_noise_paragraphs(
    doc: _Document,
    *,
    preserve_phrases: tuple[str, ...] = (),
) -> int:
    body = doc.element.body
    elements = list(body)

    def _wt(e) -> str:
        return "".join(t.text or "" for t in e.iter(f"{{{_NS}}}t")).strip()

    removed = 0
    preserve = tuple(p.lower() for p in preserve_phrases if p)
    for i, elem in enumerate(elements):
        if elem.tag.split("}")[-1] != "p":
            continue
        text = _wt(elem)
        if not text:
            continue
        text_lower = text.lower()
        if preserve and any(p in text_lower for p in preserve):
            continue
        drop = False
        if re.match(r"^.{10,75}\s+\d{1,4}$", text) and len(text) < 80:
            drop = True
        elif re.search(r"3\.2[\. ][A-Z0-9P]", text) and len(text) < 200:
            for j in range(max(0, i - 10), i):
                if re.search(r"2\.3\.[SP]", _wt(elements[j])):
                    drop = True
                    break
        elif _PAGE_NUM_RE.match(text) or _PAGE_OF_RE.match(text):
            drop = True
        if drop:
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)
                removed += 1
    return removed


def _remove_empty_visual_tables(doc: _Document, *, keep_first_n_tables: int) -> int:
    removed = 0

    def get_text(cell) -> str:
        return "".join(t.text or "" for t in cell._element.iter(f"{{{_NS}}}t")).strip()

    for idx, table in enumerate(list(doc.tables)):
        if idx < keep_first_n_tables:
            continue
        total_cells = text_cells = 0
        for row in table.rows:
            for cell in _safe_row_cells(row):
                total_cells += 1
                if get_text(cell):
                    text_cells += 1
        if total_cells == 0:
            continue
        if (total_cells > 4 and text_cells == 0) or (
            total_cells > 6 and (total_cells - text_cells) / total_cells > 0.8
        ):
            table._element.getparent().remove(table._element)
            removed += 1
    return removed


def _remove_low_content_injected_tables(doc: _Document, *, keep_first_n_tables: int) -> int:
    removed = 0

    def get_text(cell) -> str:
        return "".join(t.text or "" for t in cell._element.iter(f"{{{_NS}}}t")).strip()

    for idx, table in enumerate(list(doc.tables)):
        if idx < keep_first_n_tables:
            continue
        total_cells = text_cells = text_chars = 0
        for row in table.rows:
            for cell in _safe_row_cells(row):
                total_cells += 1
                text = get_text(cell)
                if text:
                    text_cells += 1
                    text_chars += len(text)
        if total_cells == 0:
            continue
        empty_ratio = (total_cells - text_cells) / total_cells
        if (
            (text_cells == 0 and total_cells >= 4)
            or (empty_ratio > 0.85 and text_chars < 80 and total_cells >= 6)
            or (text_cells <= 2 and total_cells >= 8 and text_chars < 100)
        ):
            table._element.getparent().remove(table._element)
            removed += 1
    return removed


def run_artifact_cleanup(
    doc: _Document,
    *,
    keep_first_n_tables: int,
    preserve_repeated_patterns: tuple[str, ...] = (),
    preserve_phrases: tuple[str, ...] = (),
) -> dict[str, int]:
    """Unified post-insertion cleanup. Returns per-step counters."""
    return {
        "repeated_headers_removed": _remove_repeated_header_paragraphs(
            doc, preserve_patterns=preserve_repeated_patterns
        ),
        "noise_paragraphs_removed": _remove_pdf_noise_paragraphs(
            doc, preserve_phrases=preserve_phrases
        ),
        "empty_tables_removed": _remove_empty_visual_tables(
            doc, keep_first_n_tables=keep_first_n_tables
        ),
        "low_content_tables_removed": _remove_low_content_injected_tables(
            doc, keep_first_n_tables=keep_first_n_tables
        ),
        "blank_paragraphs_collapsed": _collapse_blank_paragraphs(doc),
    }


# ---------------------------------------------------------------------------
# Shared base class for all section fillers
# ---------------------------------------------------------------------------

class _DocxHelper:
    """XML/paragraph utilities shared by every section filler."""

    SECTION_START: str = ""
    SECTION_END: str = ""

    # ------------------------------------------------------------------
    # Low-level XML
    # ------------------------------------------------------------------

    @staticmethod
    def _insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        para = Paragraph(new_p, paragraph._parent)
        if text:
            para.add_run(text)
        return para

    def _get_available_page_width(self, doc: _Document) -> int:
        try:
            section = doc.sections[0]
            available = section.page_width - section.left_margin - section.right_margin
            padding = int(914400 * 0.2)
            return max(available - 2 * padding, int(914400 * 4))
        except Exception:
            return int(914400 * 5.5)

    def _get_available_page_height(self, doc: _Document) -> int:
        try:
            section = doc.sections[0]
            available = section.page_height - section.top_margin - section.bottom_margin
            padding = int(914400 * 0.3)
            return max(available - 2 * padding, int(914400 * 3.5))
        except Exception:
            return int(914400 * 7.5)

    def _add_picture_autofit(self, run, image_path: Path | str, doc: _Document) -> None:
        """Insert picture scaled to fit within page width and height."""
        img_path = Path(image_path)
        max_width = self._get_available_page_width(doc)
        max_height = self._get_available_page_height(doc)
        target_width = max_width
        try:
            pix = fitz.Pixmap(str(img_path))
            w, h = max(1, pix.width), max(1, pix.height)
            aspect = h / w
            if aspect > 0:
                height_limited = int(max_height / aspect)
                if height_limited > 0:
                    target_width = min(max_width, height_limited)
        except Exception:
            pass
        run.add_picture(str(img_path), width=target_width)

    @staticmethod
    def _add_picture_autofit_bounds(
        run,
        image_path: Path | str,
        *,
        max_width_emu: int,
        max_height_emu: int,
    ) -> None:
        img_path = Path(image_path)
        target_width = int(max_width_emu)
        try:
            pix = fitz.Pixmap(str(img_path))
            w, h = max(1, pix.width), max(1, pix.height)
            aspect = h / w
            if aspect > 0:
                height_limited = int(max_height_emu / aspect)
                if height_limited > 0:
                    target_width = min(int(max_width_emu), height_limited)
        except Exception:
            pass
        run.add_picture(str(img_path), width=int(target_width))

    @staticmethod
    def _delete_paragraph(paragraph: Paragraph) -> None:
        elem = paragraph._element
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)

    @staticmethod
    def _safe_set_cell_text(cell, text: str) -> None:
        try:
            cell.text = text
        except Exception:
            return

    @staticmethod
    def _set_runs_style(
        paragraph: Paragraph,
        bold: bool | None = None,
        italic: bool | None = None,
    ) -> None:
        if not paragraph.runs:
            paragraph.add_run(paragraph.text or "")
        for run in paragraph.runs:
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic

    @staticmethod
    def _replace_paragraph_with_runs(
        paragraph: Paragraph,
        chunks: list[tuple[str, bool | None, bool | None]],
    ) -> None:
        for run in list(paragraph.runs):
            paragraph._p.remove(run._r)
        for text, bold, italic in chunks:
            r = paragraph.add_run(text)
            if bold is not None:
                r.bold = bold
            if italic is not None:
                r.italic = italic

    # ------------------------------------------------------------------
    # Paragraph search helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _norm(s: str) -> str:
        return re.sub(r"\s+", " ", s).strip().lower()

    @classmethod
    def _find_para_index(
        cls,
        paragraphs: list[Paragraph],
        needle: str,
        start_idx: int,
        end_idx: int,
        *,
        startswith: bool = False,
    ) -> int | None:
        n = cls._norm(needle)
        for i in range(start_idx, min(end_idx, len(paragraphs))):
            t = cls._norm(paragraphs[i].text or "")
            if startswith and t.startswith(n):
                return i
            if not startswith and n in t:
                return i
        return None

    @classmethod
    def _find_para_index_doc(
        cls,
        doc: _Document,
        needle: str,
        start_idx: int,
        end_idx: int,
        *,
        startswith: bool = False,
    ) -> int | None:
        return cls._find_para_index(
            doc.paragraphs, needle, start_idx, end_idx, startswith=startswith
        )

    # ------------------------------------------------------------------
    # Section range
    # ------------------------------------------------------------------

    def _get_target_range(self, doc: _Document) -> tuple[int, int]:
        start_idx = end_idx = None
        for i, p in enumerate(doc.paragraphs):
            text = (p.text or "").strip()
            if start_idx is None and self.SECTION_START.lower() in text.lower():
                start_idx = i
            if start_idx is not None and self.SECTION_END.lower() in text.lower():
                end_idx = i
                break
        if start_idx is None:
            raise ValueError(f"Section start '{self.SECTION_START}' not found in template")
        return start_idx, end_idx if end_idx is not None else len(doc.paragraphs)

    # ------------------------------------------------------------------
    # Common value insertion
    # ------------------------------------------------------------------

    def _add_answer_after(self, doc: _Document, anchor_idx: int, value: str) -> None:
        if not value.strip():
            return
        self._insert_paragraph_after(doc.paragraphs[anchor_idx], value)

    def _set_value_under_label(
        self, doc: _Document, label: str, value: str, start_idx: int, end_idx: int
    ) -> None:
        if not value.strip():
            return
        idx = self._find_para_index_doc(doc, label, start_idx, end_idx, startswith=True)
        if idx is None:
            return
        if idx + 1 < len(doc.paragraphs) and not (doc.paragraphs[idx + 1].text or "").strip():
            doc.paragraphs[idx + 1].text = value
        else:
            self._insert_paragraph_after(doc.paragraphs[idx], value)

    def _remove_paragraphs_matching(
        self,
        doc: _Document,
        patterns: list[str],
        start_idx: int,
        end_idx: int,
    ) -> None:
        to_delete = []
        for i in range(start_idx, min(end_idx, len(doc.paragraphs))):
            t = self._norm(doc.paragraphs[i].text or "")
            if any(t.startswith(self._norm(p)) for p in patterns):
                to_delete.append(i)
        for idx in reversed(to_delete):
            if idx < len(doc.paragraphs):
                self._delete_paragraph(doc.paragraphs[idx])

    # ------------------------------------------------------------------
    # Text-block helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _extract_block(text: str, start_pattern: str, end_patterns: list[str]) -> str:
        m = re.search(start_pattern, text, flags=re.IGNORECASE)
        if not m:
            return ""
        block = text[m.end():]
        end_pos = len(block)
        for ep in end_patterns:
            em = re.search(ep, block, flags=re.IGNORECASE)
            if em:
                end_pos = min(end_pos, em.start())
        return block[:end_pos].strip()

    @staticmethod
    def _first_meaningful_line(block: str) -> str:
        for ln in block.splitlines():
            c = ln.strip()
            if c:
                return c
        return ""

    @staticmethod
    def _clean_line(line: str) -> str:
        return re.sub(r"\s+", " ", line.strip())

    def _resolve_name_manufacturer_line(
        self, ref_docx: Path | None, section_heading: str
    ) -> str:
        override = getattr(self, "_name_mfr_line", "")
        if override:
            return override
        if not ref_docx or not ref_docx.exists():
            return ""
        try:
            from docx import Document
            ref = Document(ref_docx)
            for i, p in enumerate(ref.paragraphs):
                if section_heading.lower() in (p.text or "").lower():
                    for j in range(i + 1, min(i + 6, len(ref.paragraphs))):
                        cand = (ref.paragraphs[j].text or "").strip()
                        if cand.startswith("(") and "," in cand:
                            return cand
        except Exception:
            pass
        return ""

    @staticmethod
    def _append_inline_value_after_colon(paragraph: Paragraph, value: str) -> None:
        if not value:
            return
        text = (paragraph.text or "").rstrip()
        if not text:
            paragraph.text = value
            return
        if ":" in text:
            head, tail = text.rsplit(":", 1)
            if tail.strip():
                return
            paragraph.text = f"{head}: {value}"
            return
        if value.lower() not in text.lower():
            paragraph.text = f"{text} {value}".strip()

    @staticmethod
    def _strip_angle_bracket_placeholder(paragraph: Paragraph) -> None:
        text = (paragraph.text or "").strip()
        if not text:
            return
        paragraph.text = re.sub(r"\s*<[^>]*>\s*$", "", text)
