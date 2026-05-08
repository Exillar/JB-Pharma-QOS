"""builders/s2.py — S2DocxFiller: fills QOS section 2.3.S.2 Manufacture."""
from __future__ import annotations

import re
from pathlib import Path

import fitz
from docx import Document
from docx.document import Document as _Document
from docx.text.paragraph import Paragraph

from builders.base import _DocxHelper, run_artifact_cleanup
from config_loader import DiagramConfig, NoiseConfig, S2FillConfig
from pdf_extractor import ExtractedSectionContent


class S2DocxFiller(_DocxHelper):
    SECTION_START = "2.3.S.2 Manufacture"
    SECTION_END = "2.3.S.3 Characterisation"

    _PAGE_STAMP_RE = re.compile(r"^(page\s*)?\d+\s+of\s+\d+$", re.IGNORECASE)

    _PHARMA_RE = re.compile(
        r"\b(pharmaceutical|pharma|biotech|chemical|laboratory|laboratories)\b"
        r".*?\b(ltd|limited|inc|co\.?|corp|llc|plc|gmbh|ag|s\.a\.?)\b",
        re.IGNORECASE,
    )
    _ZIP_RE = re.compile(r"\d{4,}")

    _S22_KEYWORDS = {
        "flow": ("flow diagram", "flow chart", "synthesis process"),
        "brief": (
            "brief narrative",
            "description of the manufacturing process",
            "description of manufacturing process",
        ),
        "alternate": ("alternate processes", "alternative processes"),
        "reprocessing": ("reprocessing steps", "reprocessing step"),
    }

    def __init__(
        self,
        template_docx: Path,
        filled_reference_docx: Path | None = None,
        images_dir: Path | None = None,
        s2_fill_cfg: S2FillConfig | None = None,
        noise_cfg: NoiseConfig | None = None,
        diagram_cfg: DiagramConfig | None = None,
    ) -> None:
        self.template_docx = template_docx
        self.filled_reference_docx = filled_reference_docx
        self.images_dir = images_dir
        self.cfg = s2_fill_cfg or S2FillConfig()
        self.noise_cfg = noise_cfg or NoiseConfig()
        self.diagram_cfg = diagram_cfg or DiagramConfig()

    # ------------------------------------------------------------------
    # Text cleaning
    # ------------------------------------------------------------------

    def _clean_block(self, text: str) -> str:
        out: list[str] = []
        for ln in text.splitlines():
            s = ln.strip()
            if not s:
                continue
            if re.match(r"^3\.2\.[sp]\.\d+(\.\d+)*\b", s.lower()) and len(s) < 80:
                continue
            if self._PAGE_STAMP_RE.match(s):
                continue
            out.append(s)
        return "\n".join(out).strip()

    @staticmethod
    def _lines(text: str) -> list[str]:
        return [ln.strip() for ln in text.splitlines() if ln.strip()]

    # ------------------------------------------------------------------
    # S2.2 keyword extraction
    # ------------------------------------------------------------------

    @classmethod
    def _extract_s22_block(
        cls,
        lines: list[str],
        start_keywords: tuple[str, ...],
        stop_keywords: tuple[str, ...],
    ) -> str:
        start_idx = None
        for i, ln in enumerate(lines):
            low = ln.lower()
            if any(k in low for k in start_keywords):
                start_idx = i
                break
        if start_idx is None:
            return ""
        line = lines[start_idx]
        block: list[str] = []
        if ":" in line:
            tail = line.split(":", 1)[1].strip()
            if tail:
                block.append(tail)
        for ln in lines[start_idx + 1:]:
            low = ln.lower()
            if any(k in low for k in stop_keywords):
                break
            if re.match(r"^\(?[a-d]\)?\s*[\).:-]", low):
                if block:
                    break
            block.append(ln)
        return " ".join(s for s in block if s).strip()

    @staticmethod
    def _derive_s22_brief_narrative(lines: list[str]) -> str:
        heading_re = re.compile(
            r"^\s*3\s*[\.\-]\s*2\s*[\.\-]\s*[sp]\s*[\.\-]\s*\d+(?:\s*[\.\-]\s*\d+)*\b",
            re.IGNORECASE,
        )
        stage_re = re.compile(r"^(stage|step)\b", re.IGNORECASE)
        company_re = re.compile(r"\b(co\.?|ltd|limited|pharmaceutical|laboratories?)\b", re.IGNORECASE)
        candidates: list[str] = []
        for ln in lines:
            s = " ".join(ln.split())
            low = s.lower()
            if not s or heading_re.match(s):
                continue
            if low.startswith("3.2") and "description of manufacturing process" in low:
                continue
            if stage_re.match(low) or low.startswith("figure "):
                continue
            if "flow diagram" in low and len(s.split()) <= 8:
                continue
            if "chemical synthetical pathway" in low:
                continue
            if company_re.search(s) and len(s.split()) <= 10:
                continue
            if "  " in ln and len(s.split()) <= 5:
                continue
            if len(s) >= 45 and any(ch in s for ch in (".", ";", ":")):
                candidates.append(s)
        if candidates:
            return candidates[0]
        for ln in lines:
            s = " ".join(ln.split())
            if len(s) >= 55 and not heading_re.match(s):
                return s
        return ""

    # ------------------------------------------------------------------
    # S2.2 narrative image extraction
    # ------------------------------------------------------------------

    def _extract_s22_narrative_image(self, pdf_path: Path) -> Path | None:
        if self.images_dir is None:
            return None
        out_path = self.images_dir / "3_2_S_2_2_narrative.png"

        targets = (
            "brief narrative description",
            "description of the manufacturing process",
            "description of manufacturing process",
        )
        stop_targets = (
            "flow diagram",
            "alternate processes",
            "reprocessing steps",
            "control of materials",
        )

        def drawing_clip(page: fitz.Page) -> fitz.Rect | None:
            drawings = page.get_drawings()
            if not drawings:
                return None
            union = None
            page_rect = page.rect
            top_cut = page_rect.y0 + page_rect.height * self.diagram_cfg.header_crop_frac
            bottom_cut = page_rect.y1 - page_rect.height * self.diagram_cfg.footer_crop_frac
            for drawing in drawings:
                drect = drawing.get("rect")
                if drect is None:
                    continue
                if drect.width >= page_rect.width * 0.95 or drect.height >= page_rect.height * 0.95:
                    continue
                if drect.y1 <= top_cut or drect.y0 >= bottom_cut:
                    continue
                union = drect if union is None else union | drect
            if union is None:
                return None
            pad_x = max(4.0, union.width * 0.01)
            pad_y = max(4.0, union.height * 0.01)
            return (fitz.Rect(union.x0 - pad_x, union.y0 - pad_y, union.x1 + pad_x, union.y1 + pad_y)
                    & page.rect)

        def embedded_image_clip(page: fitz.Page) -> fitz.Rect | None:
            page_rect = page.rect
            union = None
            try:
                for img in page.get_images(full=True):
                    xref = img[0]
                    for rect in page.get_image_rects(xref):
                        if rect.width < 24 or rect.height < 24:
                            continue
                        if rect.y1 <= page_rect.y0 + page_rect.height * self.diagram_cfg.header_crop_frac:
                            continue
                        if rect.y0 >= page_rect.y1 - page_rect.height * self.diagram_cfg.footer_crop_frac:
                            continue
                        union = rect if union is None else union | rect
            except Exception:
                return None
            if union is None:
                return None
            pad_x = max(4.0, union.width * 0.02)
            pad_y = max(4.0, union.height * 0.02)
            return (fitz.Rect(union.x0 - pad_x, union.y0 - pad_y, union.x1 + pad_x, union.y1 + pad_y)
                    & page_rect)

        def preferred_visual_clip(page: fitz.Page, page_text_lower: str) -> fitz.Rect | None:
            clip = embedded_image_clip(page)
            if clip is None:
                clip = drawing_clip(page)
            if clip is None:
                return None
            keyword_ys: list[float] = []
            for kw in self.diagram_cfg.diagram_exclude_keywords:
                if kw not in page_text_lower:
                    continue
                try:
                    for r in page.search_for(kw):
                        keyword_ys.append(r.y1)
                except Exception:
                    continue
            if keyword_ys:
                top = max(clip.y0, min(keyword_ys) + 4.0)
                clip = fitz.Rect(clip.x0, top, clip.x1, clip.y1) & page.rect
            return clip if clip.height > 20 and clip.width > 20 else None

        def clip_visual_score(page: fitz.Page, clip: fitz.Rect) -> float:
            score = 0.0
            try:
                for img in page.get_images(full=True):
                    xref = img[0]
                    for rect in page.get_image_rects(xref):
                        inter = rect & clip
                        if not inter.is_empty:
                            score += max(1.0, inter.width * inter.height / 5000.0)
            except Exception:
                pass
            try:
                for d in page.get_drawings():
                    drect = d.get("rect")
                    if drect is None:
                        continue
                    inter = drect & clip
                    if not inter.is_empty:
                        score += 0.8
            except Exception:
                pass
            try:
                words = page.get_text("words", clip=clip)
                score -= min(20.0, len(words) / 25.0)
            except Exception:
                pass
            return score

        try:
            with fitz.open(pdf_path) as doc:
                best_page = best_clip = None
                best_score = -1.0
                for page in doc:
                    text = page.get_text("text", sort=True)
                    low = text.lower()
                    if not any(kw in low for kw in self.diagram_cfg.diagram_exclude_keywords):
                        continue
                    clip = preferred_visual_clip(page, low)
                    if clip is None:
                        continue
                    score = clip_visual_score(page, clip)
                    if score > best_score:
                        best_score = score
                        best_page = page
                        best_clip = clip

                if best_page is not None and best_clip is not None and best_score > 1.5:
                    scale = self.diagram_cfg.render_dpi_scale
                    pix = best_page.get_pixmap(
                        matrix=fitz.Matrix(scale, scale), clip=best_clip, alpha=False
                    )
                    pix.save(str(out_path))
                    return out_path

                # Strategy 2: scan for brief-narrative keyword page
                for page in doc:
                    text = page.get_text("text", sort=True)
                    low = text.lower()
                    if not any(kw in low for kw in targets):
                        continue
                    if any(kw in low for kw in stop_targets):
                        continue
                    clip = preferred_visual_clip(page, low)
                    if clip is None:
                        page_rect = page.rect
                        top_cut = page_rect.y0 + page_rect.height * self.diagram_cfg.header_crop_frac
                        bottom_cut = page_rect.y1 - page_rect.height * self.diagram_cfg.footer_crop_frac
                        clip = fitz.Rect(page_rect.x0, top_cut, page_rect.x1, bottom_cut)
                    scale = self.diagram_cfg.render_dpi_scale
                    pix = page.get_pixmap(
                        matrix=fitz.Matrix(scale, scale), clip=clip, alpha=False
                    )
                    pix.save(str(out_path))
                    return out_path
        except Exception:
            return None
        return None

    # ------------------------------------------------------------------
    # Parsers
    # ------------------------------------------------------------------

    def _parse_s21(self, text: str) -> dict[str, str]:
        lines = self._lines(self._clean_block(text))
        start = 0
        for i, ln in enumerate(lines):
            if any(kw in ln.lower() for kw in self.cfg.narrative_start_keywords):
                start = i
                break

        useful: list[str] = []
        for ln in lines[start:]:
            if any(kw in ln.lower() for kw in self.cfg.narrative_end_keywords):
                break
            useful.append(ln)

        _LEGAL_SUFFIX_RE = re.compile(
            r"\b(ltd\.?|limited|inc\.?|co\.?|corp\.?|llc|plc|gmbh|ag|s\.a\.?)\s*$",
            re.IGNORECASE,
        )
        mfr_name = mfr_idx = ""
        for i, ln in enumerate(useful):
            if self._PHARMA_RE.search(ln):
                mfr_name = ln
                mfr_idx = i
                break
        if not mfr_name:
            for i, ln in enumerate(useful):
                if _LEGAL_SUFFIX_RE.search(ln) and len(ln) > 6:
                    mfr_name = ln
                    mfr_idx = i
                    break

        mfr_addr = ""
        if mfr_idx != "":
            parts: list[str] = []
            stop_kw = ("manufactur", "responsib", "test", "packag", "gmp")
            addr_kw = ("china", "india", "japan", "germany", "france",
                       "street", "road", "avenue", "plot", "block", "zone",
                       "district", "province", "state", "city")
            for ln in useful[mfr_idx + 1: mfr_idx + 6]:
                low = ln.lower()
                if any(k in low for k in stop_kw):
                    break
                if re.search(r"\d", ln) or any(k in low for k in addr_kw):
                    parts.append(ln)
                    if self._ZIP_RE.search(ln):
                        break
            mfr_addr = " ".join(parts).strip()

        responsibility = ""
        for ln in useful:
            low = ln.lower()
            if "manufactur" in low and ("test" in low or "packag" in low):
                responsibility = self.cfg.manufacturer_table_responsibility_default
                break

        gmp_line = ""
        for ln in useful:
            if any(kw in ln.lower() for kw in self.cfg.gmp_keywords):
                gmp_line = self.cfg.gmp_found_sentence
                break
        gmp_line = gmp_line or self.cfg.gmp_fallback_sentence

        return {
            "mfr_name": mfr_name,
            "mfr_addr": mfr_addr,
            "responsibility": responsibility or self.cfg.manufacturer_table_responsibility_default,
            "gmp": gmp_line,
        }

    def _parse_s22(self, text: str) -> dict[str, str]:
        lines = self._lines(self._clean_block(text))
        stop_keys = (
            self._S22_KEYWORDS["brief"]
            + self._S22_KEYWORDS["alternate"]
            + self._S22_KEYWORDS["reprocessing"]
            + self._S22_KEYWORDS["flow"]
        )
        brief = self._extract_s22_block(lines, self._S22_KEYWORDS["brief"], stop_keys)
        if not brief:
            brief = self._derive_s22_brief_narrative(lines)
        alternate = self._extract_s22_block(lines, self._S22_KEYWORDS["alternate"], stop_keys)
        reprocessing = self._extract_s22_block(lines, self._S22_KEYWORDS["reprocessing"], stop_keys)
        return {
            "brief": brief,
            "alternate": alternate or self.cfg.alternate_processes_default,
            "reprocessing": reprocessing or self.cfg.reprocessing_steps_default,
        }

    def _restricted_phrase(self, text: str) -> str:
        kws = tuple(k.lower() for k in self.cfg.restricted_phrase_keywords)
        for ln in text.splitlines():
            if any(k in ln.lower() for k in kws):
                return ln.strip()
        return ""

    def _parse_s23(self, text: str) -> dict[str, str]:
        redacted = self._restricted_phrase(text)
        if not redacted:
            skip_pats = (
                re.compile(r"^3\.2\.[sp]\.2\.3\b", re.IGNORECASE),
                re.compile(r"^refer\s+section\s+3\.2\.[sp]\.2\.3\b", re.IGNORECASE),
                re.compile(r"^control\s+of\s+materials\b", re.IGNORECASE),
            )
            for ln in text.splitlines():
                s = " ".join(ln.split()).strip()
                if not s or any(p.search(s) for p in skip_pats):
                    continue
                redacted = s
                break
        return {
            "a": redacted, "b": redacted, "c": redacted, "d": redacted,
            "b_inline_default": self.cfg.s23_manufacturer_not_available_default,
        }

    def _parse_generic_restricted_section(self, section_ref: str, text: str) -> dict[str, str]:
        redacted = self._restricted_phrase(text)
        if not redacted:
            sec_tail = section_ref.split(".")[-1]
            try:
                sec_num = int(sec_tail)
                sec_re = re.compile(rf"^3\.2\.[sp]\.2\.{sec_num}\b", re.IGNORECASE)
            except Exception:
                sec_re = re.compile(r"^3\.2\.[sp]\.2\.[0-9]+\b", re.IGNORECASE)
            refer_re = re.compile(rf"^refer\s+section\s+{re.escape(section_ref)}", re.IGNORECASE)
            for ln in text.splitlines():
                s = " ".join(ln.split()).strip()
                if not s:
                    continue
                if sec_re.search(s) or refer_re.search(s) or re.search(r"control|controls|summary", s, re.IGNORECASE):
                    continue
                redacted = s
                break
        return {
            "a": redacted, "b": redacted, "c": redacted, "d": redacted,
            "b_inline_default": self.cfg.s23_manufacturer_not_available_default,
        }

    # ------------------------------------------------------------------
    # Table and heading helpers
    # ------------------------------------------------------------------

    def _fill_s21_table(self, doc: _Document, mfr_name: str, mfr_addr: str, responsibility: str) -> bool:
        for table in doc.tables:
            if len(table.rows) < 2 or len(table.columns) < 2:
                continue
            head = " ".join(
                (table.cell(0, c).text or "").strip().lower()
                for c in range(min(3, len(table.columns)))
            )
            if "name and address" in head and "responsibility" in head:
                addr_block = mfr_name
                if mfr_addr:
                    addr_block += f"\n\nAddress of Manufacturer:\n{mfr_addr}"
                table.cell(1, 0).text = addr_block.strip()
                table.cell(1, 1).text = responsibility
                if len(table.columns) > 2:
                    table.cell(1, 2).text = self.cfg.manufacturer_table_apprx_col
                return True
        return False

    @staticmethod
    def _normalize_s23_heading(doc: _Document, start_idx: int, end_idx: int) -> None:
        pat = re.compile(r"^(2\.3\.S\.2\.3\s+Control of Materials)\b", re.IGNORECASE)
        for i in range(start_idx, min(end_idx, len(doc.paragraphs))):
            p = doc.paragraphs[i]
            text = (p.text or "").strip()
            m = pat.match(text)
            if not m:
                continue
            canonical = m.group(1)
            if text != canonical:
                p.text = canonical
            return

    def _normalize_s23_table_header(self, doc: _Document) -> None:
        target = self.cfg.s23_table_first_header_default.strip()
        if not target:
            return
        for table in doc.tables:
            if not table.rows or len(table.columns) < 3:
                continue
            h0 = self._norm(table.cell(0, 0).text)
            h1 = self._norm(table.cell(0, 1).text)
            h2 = self._norm(table.cell(0, 2).text)
            if h0 in {"test parameter", "step / starting material"} and "test" in h1 and "acceptance criteria" in h2:
                table.cell(0, 0).text = target
                return

    def _normalize_s2_table_header(self, doc: _Document) -> None:
        self._normalize_s23_table_header(doc)

    def _normalize_s2_heading(self, doc: _Document, section_tail: str) -> None:
        pat = re.compile(rf"^(2\.3\.S\.2\.{re.escape(section_tail)}\s+.+)$", re.IGNORECASE)
        for p in doc.paragraphs:
            text = (p.text or "").strip()
            m = pat.match(text)
            if not m:
                continue
            canonical = m.group(1)
            if text != canonical:
                p.text = canonical
            return

    def _remove_refer_section_lines(self, doc: _Document, start_idx: int, end_idx: int) -> None:
        pat = re.compile(r"^refer\s+section\s+3\.2\.[sp]\.", re.IGNORECASE)
        to_del = [
            i for i in range(start_idx, min(end_idx, len(doc.paragraphs)))
            if pat.match(self._norm(doc.paragraphs[i].text or ""))
        ]
        for i in reversed(to_del):
            if i < len(doc.paragraphs):
                self._delete_paragraph(doc.paragraphs[i])

    # ------------------------------------------------------------------
    # Public entry point
    # ------------------------------------------------------------------

    def fill_s2_section(
        self,
        extracted: dict[str, ExtractedSectionContent],
        output_docx: Path,
    ) -> list[str]:
        doc = Document(self.template_docx)
        template_table_count = len(doc.tables)
        warnings: list[str] = []

        start_idx, end_idx = self._get_target_range(doc)

        name_line = self._resolve_name_manufacturer_line(
            self.filled_reference_docx, "2.3.S.2.1 Manufacturer"
        )
        for heading in [
            "2.3.S.2 Manufacture",
            "2.3.S.2.1 Manufacturer(s)",
            "2.3.S.2.2 Description of Manufacturing Process and Process Controls",
            "2.3.S.2.3 Control of Materials",
            "2.3.S.2.4 Controls of Critical Steps and Intermediates",
            "2.3.S.2.5 Process Validation and/or Evaluation",
            "2.3.S.2.6 Manufacturing Process Development",
        ]:
            idx = self._find_para_index_doc(doc, heading, start_idx, end_idx)
            if idx is not None and name_line:
                nxt = (doc.paragraphs[idx + 1].text or "").strip() if idx + 1 < len(doc.paragraphs) else ""
                if not nxt.startswith("("):
                    self._insert_paragraph_after(doc.paragraphs[idx], name_line)

        start_idx, end_idx = self._get_target_range(doc)
        self._remove_refer_section_lines(doc, start_idx, end_idx)

        for k in ["3.2.S.2.1", "3.2.S.2.2", "3.2.S.2.3", "3.2.S.2.4", "3.2.S.2.5", "3.2.S.2.6"]:
            if extracted[k].warning:
                warnings.append(f"{k}: {extracted[k].warning}")

        s21 = self._parse_s21(extracted["3.2.S.2.1"].raw_text)
        s22 = self._parse_s22(extracted["3.2.S.2.2"].raw_text)
        s23 = self._clean_block(extracted["3.2.S.2.3"].raw_text)
        s24 = self._clean_block(extracted["3.2.S.2.4"].raw_text)
        s25 = self._clean_block(extracted["3.2.S.2.5"].raw_text)
        s26 = self._clean_block(extracted["3.2.S.2.6"].raw_text)

        s23_data = self._parse_s23(s23)
        s24_data = self._parse_generic_restricted_section("3.2.S.2.4", s24)
        s25_data = self._parse_generic_restricted_section("3.2.S.2.5", s25)
        s26_data = self._parse_generic_restricted_section("3.2.S.2.6", s26)

        start_idx, end_idx = self._get_target_range(doc)
        self._normalize_s23_heading(doc, start_idx, end_idx)
        self._normalize_s2_heading(doc, "2.4")
        self._normalize_s2_heading(doc, "2.5")
        self._normalize_s2_heading(doc, "2.6")

        filled = self._fill_s21_table(doc, s21["mfr_name"], s21["mfr_addr"], s21["responsibility"])
        if not filled:
            warnings.append("3.2.S.2.1: manufacturer table not found in template")

        idx = self._find_para_index_doc(
            doc, "Manufacturing authorization for the production of API", start_idx, end_idx
        )
        if idx is not None:
            self._insert_paragraph_after(doc.paragraphs[idx], s21["gmp"])

        # S2.2 (a) flow diagrams
        idx = self._find_para_index_doc(doc, "Flow diagram of the synthesis process", start_idx, end_idx)
        if idx is not None and extracted["3.2.S.2.2"].image_paths:
            anchor = current = doc.paragraphs[idx]
            img_count = 0
            for img_path in extracted["3.2.S.2.2"].image_paths:
                try:
                    img_path_obj = Path(img_path) if not isinstance(img_path, Path) else img_path
                    if not img_path_obj.exists():
                        warnings.append(f"3.2.S.2.2: image not found: {img_path}")
                        continue
                    p = self._insert_paragraph_after(current, "")
                    self._add_picture_autofit(p.add_run(), img_path, doc)
                    current = p
                    img_count += 1
                except Exception as e:
                    warnings.append(f"3.2.S.2.2: failed to insert image {img_path}: {e}")
            if img_count == 0:
                warnings.append(f"3.2.S.2.2: {len(extracted['3.2.S.2.2'].image_paths)} images found but none inserted")
        elif idx is not None:
            warnings.append("3.2.S.2.2: no flow-diagram images extracted")

        # S2.2 (b) brief narrative
        idx = self._find_para_index_doc(
            doc, "Brief narrative description of the manufacturing process", start_idx, end_idx
        )
        if idx is not None:
            brief = s22["brief"].strip()
            current = doc.paragraphs[idx]
            narrative_img = self._extract_s22_narrative_image(extracted["3.2.S.2.2"].source_pdf)
            if narrative_img:
                img_path = Path(narrative_img) if not isinstance(narrative_img, Path) else narrative_img
                if img_path.exists():
                    try:
                        p = self._insert_paragraph_after(current, "")
                        self._add_picture_autofit(p.add_run(), img_path, doc)
                        current = p
                    except Exception as e:
                        warnings.append(f"3.2.S.2.2: failed to insert narrative image: {e}")
            if brief and len(brief) >= 30:
                self._insert_paragraph_after(current, brief)
            elif not narrative_img:
                warnings.append("3.2.S.2.2: narrative text/image not found")

        # S2.2 (c) alternate processes
        idx = self._find_para_index_doc(doc, "Alternate processes and explanation", start_idx, end_idx)
        if idx is not None:
            self._insert_paragraph_after(doc.paragraphs[idx], s22["alternate"])

        # S2.2 (d) reprocessing
        idx = self._find_para_index_doc(doc, "Reprocessing steps and justification", start_idx, end_idx)
        if idx is not None:
            self._insert_paragraph_after(doc.paragraphs[idx], s22["reprocessing"])

        # S2.3
        s23_label_map = [
            ("(a)\tName of starting material:", "a"),
            ("(b)\tName and manufacturing site address of starting material", "b"),
            ("Summary of the quality and controls of the starting materials", "c"),
            ("without risk of transmitting agents of animal spongiform", "d"),
        ]
        for label, key in s23_label_map:
            idx = self._find_para_index_doc(doc, label, start_idx, end_idx)
            if idx is None:
                continue
            if key == "b":
                self._append_inline_value_after_colon(
                    doc.paragraphs[idx], s23_data["b_inline_default"]
                )
            ans = s23_data[key].strip()
            if ans:
                self._insert_paragraph_after(doc.paragraphs[idx], ans)

        self._normalize_s23_table_header(doc)
        start_idx, end_idx = self._get_target_range(doc)

        # S2.4
        idx = self._find_para_index_doc(
            doc, "Summary of the controls performed at critical steps", start_idx, end_idx
        )
        if idx is not None:
            if s24_data.get("a"):
                self._insert_paragraph_after(doc.paragraphs[idx], s24_data["a"])
            self._normalize_s2_table_header(doc)
        start_idx, end_idx = self._get_target_range(doc)

        # S2.5
        idx = self._find_para_index_doc(
            doc, "Description of process validation and/or evaluation", start_idx, end_idx
        )
        if idx is not None and s25_data.get("a"):
            self._insert_paragraph_after(doc.paragraphs[idx], s25_data["a"])
        start_idx, end_idx = self._get_target_range(doc)

        # S2.6
        idx = self._find_para_index_doc(
            doc, "Description and discussion of the significant changes", start_idx, end_idx
        )
        if idx is not None and s26_data.get("a"):
            self._insert_paragraph_after(doc.paragraphs[idx], s26_data["a"])

        cleanup_stats = run_artifact_cleanup(
            doc,
            keep_first_n_tables=template_table_count,
            preserve_repeated_patterns=self.cfg.restricted_phrase_keywords,
        )
        if any(cleanup_stats.values()):
            warnings.append(f"cleanup: {cleanup_stats}")

        output_docx.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_docx)
        return warnings
